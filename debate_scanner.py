import requests, os, json, re, io, concurrent.futures, time
from flask import Blueprint, render_template, request, send_file, copy_current_request_context, jsonify
from flask_login import current_user, login_required
from datetime import datetime
from cache_models import CachedTranscript, CachedTWFYSearch, StakeholderOrg
try:
    import feedparser
except ImportError:
    feedparser = None

# Import Word Document libraries
try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None
    OxmlElement = None
    qn = None

def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0000FF'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

debate_scanner_bp = Blueprint('debates', __name__)


@debate_scanner_bp.app_template_filter('debate_url')
def debate_url_filter(url):
    """Return a full debate URL. Hansard API rows already carry full URLs;
    TWFY rows carry relative paths that need the domain prepended."""
    if not url:
        return ''
    if url.startswith('http'):
        return url
    return f"https://www.theyworkforyou.com{url}"

TWFY_API_KEY = os.environ.get("TWFY_API_KEY")
CLAUDE_API_KEY = os.environ.get("CLAUDE_API_KEY")


def _claude_fallback(prompt, max_tokens=2000):
    """Call Claude API with the same prompt Gemini received.
    Silent fallback — returns response text or None, never raises."""
    if not CLAUDE_API_KEY:
        return None
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=CLAUDE_API_KEY)
        msg = client.messages.create(
            model="claude-haiku-4-5-20251001",
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text
    except Exception:
        return None
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
TWFY_API_URL = "https://www.theyworkforyou.com/api/getDebates"
TWFY_WRANS_URL = "https://www.theyworkforyou.com/api/getWrans"
TWFY_WMS_URL = "https://www.theyworkforyou.com/api/getWMS"
HANSARD_API_BASE = "https://hansard-api.parliament.uk"
MINISTER_CACHE_FILE = os.path.join(os.path.dirname(__file__), 'minister_cache.json')
MINISTER_CACHE_TTL = 30 * 24 * 3600  # 30 days — reshuffles are infrequent
MINISTER_CACHE_VERSION = 2  # increment when display_name logic changes to bust stale files

# Ministers whose names the TWFY getLords/getMPs search endpoint fails to match.
# Verified person_ids from direct debate records — these are seeded at cache-write time.
KNOWN_TWFY_IDS = {
    'Baroness Smith of Malvern': {'person_id': '10549', 'is_lord': True, 'lookup_failed': False},
    # Parliamentary Under-Secretary of State for Education (DfE) — TWFY person_id verified 2026-04
    'Josh MacAlister': {'person_id': '26321', 'is_lord': False, 'lookup_failed': False},
}

DEPARTMENTS_TWFY = [
    "All Departments", "Department for Education", "Department of Health and Social Care",
    "HM Treasury", "Home Office", "Ministry of Defence", "Ministry of Justice",
    "Department for Science, Innovation and Technology", "Cabinet Office"
]

DEPT_KEYWORDS = {
    "Department for Education": '("Education" OR "Schools" OR "Skills" OR "Childcare" OR "Universities" OR "SEND")',
    "Department of Health and Social Care": '("Health" OR "NHS" OR "Social Care")',
    "HM Treasury": '("Treasury" OR "Tax" OR "Economy" OR "Spending")',
    "Home Office": '("Home Office" OR "Police" OR "Immigration")',
    "Ministry of Defence": '("Defence" OR "Military" OR "Armed Forces")',
    "Ministry of Justice": '("Justice" OR "Prisons" OR "Courts")',
    "Department for Science, Innovation and Technology": '("Science" OR "Technology" OR "Innovation")',
    "Cabinet Office": '("Cabinet Office" OR "Civil Service")'
}

PARLIAMENT_WQ_API = "https://questions-statements-api.parliament.uk/api/writtenquestions/questions"
PARLIAMENT_DEPT_IDS = {
    "Department for Education": 60, "Department of Health and Social Care": 17,
    "HM Treasury": 14, "Home Office": 1, "Ministry of Defence": 11,
    "Ministry of Justice": 54, "Department for Science, Innovation and Technology": 216,
    "Cabinet Office": 53,
}
PARTY_COLOURS_RESEARCH = {
    'Labour': '#E4003B', 'Conservative': '#0087DC', 'Liberal Democrat': '#FAA61A',
    'Scottish National Party': '#FDF38E', 'Green Party': '#02A95B', 'Reform UK': '#12B6CF',
    'Plaid Cymru': '#005B54', 'Democratic Unionist Party': '#D46A4C',
}

_MODEL_CACHE = {}
_VERIFY_CACHE = {}

def get_working_model(api_key):
    if api_key in _MODEL_CACHE:
        return _MODEL_CACHE[api_key]
    result = 'models/gemini-2.5-flash-lite'
    try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}"
        resp = requests.get(url, timeout=5)
        if resp.status_code == 200:
            available = [m['name'] for m in resp.json().get('models', [])
                         if 'generateContent' in m.get('supportedGenerationMethods', [])]
            for prefix in ['models/gemini-2.5-flash-lite', 'models/gemini-2.5-flash',
                           'models/gemini-flash-latest']:
                match = next((m for m in available if m.startswith(prefix)), None)
                if match:
                    result = match
                    break
            else:
                if available:
                    result = available[0]
    except Exception:
        pass
    _MODEL_CACHE[api_key] = result
    return result

def get_twfy_date_range(start_str, end_str):
    def parse_date(d_str):
        if not d_str: return ""
        if re.match(r'^\d{4}-\d{2}-\d{2}$', d_str): return d_str.replace('-', '')
        if '/' in d_str or '-' in d_str:
            parts = d_str.replace('/', '-').split('-')
            if len(parts) == 3 and len(parts[2]) == 4:
                return f"{parts[2]}{parts[1].zfill(2)}{parts[0].zfill(2)}"
        return d_str.replace('-', '').replace('/', '')

    s = parse_date(start_str)
    e = parse_date(end_str)

    if s and e: return f"{s}..{e}"
    elif s: return f"{s}..{datetime.now().strftime('%Y%m%d')}"
    elif e: return f"19000101..{e}"
    return ""

def _get_user_pref():
    """Return the current user's UserPreference or None."""
    try:
        from flask_app import UserPreference
        if current_user.is_authenticated:
            return UserPreference.query.filter_by(user_id=current_user.id).first()
    except Exception:
        pass
    return None

def get_debate_type(title, source=None):
    t = title.lower()
    if source == 'wms': return '📜 Ministerial Statement'
    if source == 'westminsterhall': return '🏛️ Westminster Hall'
    if 'written answers' in t or 'written answer' in t: return '✍️ Written Answer'
    if 'urgent question' in t: return '❗ Urgent Question'
    if 'oral answers' in t or 'question time' in t: return '🗣️ Oral Question'
    if 'prime minister' in t and 'question' in t: return '🗣️ Oral Question'
    if source == 'commons' and any(t.startswith(p + ':') for p in (
        'education', 'health', 'treasury', 'home office', 'defence', 'justice',
        'transport', 'environment', 'science', 'work and pensions', 'cabinet office',
        'foreign', 'culture', 'housing', 'digital', 'northern ireland', 'business',
    )): return '🗣️ Oral Question'
    if 'statement' in t: return '📜 Ministerial Statement'
    if ('statutory instrument' in t or 'affirmative' in t or 'delegated legislation' in t
            or ('draft' in t and ('regulation' in t or 'order' in t))):
        return '⚖️ Statutory Instrument'
    if 'bill' in t or 'reading' in t or 'amendment' in t: return '⚖️ Legislation'
    if 'motion' in t: return '📝 Motion'
    return '💬 General Debate'

def clean_body_text(text):
    if not text: return ""
    text = re.sub(r'<[^>]+>', ' ', text)
    return re.sub(r'\s+', ' ', text).strip()


def _parse_ai_json(text):
    """Robustly parse JSON from an AI response.
    Handles markdown fences, trailing commas, and other common AI formatting issues."""
    if not text:
        return None
    # Strip markdown fences
    text = re.sub(r'```json\s*', '', text)
    text = re.sub(r'```\s*', '', text)
    text = text.strip()
    # Try direct parse first
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    # Remove trailing commas before } or ] (valid JS, invalid JSON)
    cleaned = re.sub(r',\s*([}\]])', r'\1', text)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass
    # Try to extract just the outermost JSON object or array
    m = re.search(r'(\{.*\}|\[.*\])', text, re.DOTALL)
    if m:
        try:
            return json.loads(m.group())
        except json.JSONDecodeError:
            cleaned2 = re.sub(r',\s*([}\]])', r'\1', m.group())
            try:
                return json.loads(cleaned2)
            except json.JSONDecodeError:
                pass
    return None

def get_source_label(source):
    return {'commons': 'Commons', 'westminsterhall': 'Westminster Hall',
            'lords': 'Lords', 'wrans': 'Written Answer',
            'wms': 'Ministerial Statement'}.get(source, source.title())

# TWFY returns a mix of abbreviated and full party names — normalise to canonical
# full names so set membership checks work consistently everywhere in the pipeline.
_TWFY_PARTY_NORM = {
    'snp':          'Scottish National Party',
    'ld':           'Liberal Democrat',
    'lib dem':      'Liberal Democrat',
    'con':          'Conservative',
    'lab':          'Labour',
    'green':        'Green Party',
    'dup':          'Democratic Unionist Party',
    'uup':          'Ulster Unionist Party',
    'sdlp':         'Social Democratic and Labour Party',
    'pc':           'Plaid Cymru',
    'reform':       'Reform UK',
    'ind':          'Independent',
    'alba':         'Alba Party',
    'sf':           'Sinn Féin',
    'alliance':     'Alliance Party of Northern Ireland',
    'crossbench':   'Crossbench',
}

def _normalise_party(raw):
    """Normalise TWFY party strings to canonical full names."""
    if not raw:
        return ''
    return _TWFY_PARTY_NORM.get(raw.strip().lower(), raw.strip())

def fetch_twfy_topic(search, source_type, date_range, num=150):
    """Fetch rows from TWFY for a topic search. Returns normalised list or [] on failure.
    Results are cached for 6h (date-filtered) or 24h (open) to reduce API usage."""
    cache_key_query = f"{search} {date_range}".strip()
    ttl = 6 if date_range else 24
    try:
        cached = CachedTWFYSearch.get(cache_key_query, source_type, ttl_hours=ttl)
        if cached is not None:
            return cached
    except Exception:
        pass
    try:
        if source_type == 'wrans':
            api_url = TWFY_WRANS_URL
        elif source_type == 'wms':
            api_url = TWFY_WMS_URL
        else:
            api_url = TWFY_API_URL
        # TWFY date range syntax: SEARCHTERM YYYYMMDD..YYYYMMDD in search string.
        # TWFY cannot parse (A OR B) with a date range appended.
        # Strip outer parens from any leading OR group — handles both the plain case
        # '("a" OR "b")' and the narrow-keyword case '("a" OR "b") AND "narrow"'.
        if date_range:
            clean = search.strip()
            clean = re.sub(r'^\(([^)]+)\)', r'\1', clean)
            query = f"{clean} {date_range}"
        else:
            query = search
        params = {'key': TWFY_API_KEY, 'search': query, 'order': 'r', 'num': num, 'output': 'json'}
        if source_type not in ('wrans', 'wms'):
            params['type'] = source_type
        resp = requests.get(api_url, params=params, timeout=15)
        if resp.status_code != 200:
            return [{'_error': f"TWFY {source_type} HTTP {resp.status_code}"}]
        data = resp.json()
        if isinstance(data, dict) and 'error' in data:
            return [{'_error': f"TWFY {source_type}: {data['error']}"}]
        rows = data.get('rows', [])
        if rows and source_type == 'commons':
            import logging as _kl
            _kl.warning(f"[twfy_row_keys] sample raw keys={list(rows[0].keys())} section_id={rows[0].get('section_id')} subsection_id={rows[0].get('subsection_id')} gid={rows[0].get('gid')}")
        results = []
        for r in rows:
            body_raw = r.get('body', '')
            body_text = clean_body_text(body_raw)
            debate_title = re.sub(r'<[^>]+>', '', r.get('parent', {}).get('body', '') or '')
            if source_type == 'wrans':
                debate_title = re.sub(r'<[^>]+>', '', body_raw)[:80]
            elif source_type == 'wms':
                # WMS parent body contains the actual section title (e.g. "Students: Loans")
                # Fall back to body text only if parent is missing
                parent_body = re.sub(r'<[^>]+>', '', r.get('parent', {}).get('body', '') or '')
                debate_title = parent_body if parent_body else re.sub(r'<[^>]+>', '', body_raw)[:80]
            if source_type == 'wrans':
                dtype = 'Written Answer'
            elif source_type == 'wms':
                dtype = 'Ministerial Statement'
            else:
                dtype = get_debate_type(debate_title, source=source_type)
            results.append({
                'listurl': r.get('listurl', ''),
                'body_clean': body_text[:500],
                'body_export': body_text[:3000],
                'body_word_count': len(body_text.split()),
                'speaker_name': (r.get('speaker') or {}).get('name', 'Unknown'),
                'speaker_party': _normalise_party((r.get('speaker') or {}).get('party', '')),
                'hdate': r.get('hdate', ''),
                'debate_title': debate_title,
                'source': source_type,
                'source_label': get_source_label(source_type),
                'relevance': r.get('relevance', 0),
                'debate_type': dtype,
            })
        CachedTWFYSearch.store(cache_key_query, source_type, results)
        return results
    except Exception as e:
        return [{'_error': f"TWFY {source_type} exception: {type(e).__name__}: {e}"}]

def deduplicate_by_listurl(rows):
    seen = set()
    out = []
    for r in rows:
        key = r.get('listurl', '')
        if not key:
            out.append(r)  # no URL — always keep, never dedup
        elif key not in seen:
            seen.add(key)
            out.append(r)
    return out

def format_briefing_as_text(briefing_dict, topic):
    """Converts structured AI briefing dict to a markdown string for Word export."""
    lines = [f"## PARLIAMENTARY BRIEFING: {topic.upper()}\n"]
    lines.append(f"## 1. TOPIC SUMMARY\n{briefing_dict.get('topic_summary', '')}\n")
    lines.append(f"## 2. GOVERNMENT POSITION\n{briefing_dict.get('government_position', '')}\n")
    lines.append(f"## 3. OPPOSITION POSITION\n{briefing_dict.get('opposition_position', '')}\n")

    govt_speakers = briefing_dict.get('government_speakers', [])
    if govt_speakers:
        lines.append("## 4. GOVERNMENT SPEAKERS")
        for s in govt_speakers:
            lines.append(f"- {s.get('name', '')} ({s.get('role', '')}): {s.get('stance', '')}")
        lines.append("")

    non_govt_speakers = briefing_dict.get('non_government_speakers', [])
    if non_govt_speakers:
        lines.append("## 5. NON-GOVERNMENT SPEAKERS (Opposition / Backbench)")
        for s in non_govt_speakers:
            lines.append(f"- {s.get('name', '')} ({s.get('role_or_party', '')}): {s.get('stance', '')}")
        lines.append("")

    questions = briefing_dict.get('key_questions', [])
    if questions:
        lines.append("## 6. KEY OPPOSITION QUESTIONS")
        for q in questions:
            lines.append(f"- {q.get('speaker', '')} ({q.get('role_or_party', '')}, {q.get('date', '')}): {q.get('question', '')}")
        lines.append("")

    anticipated = briefing_dict.get('anticipated_questions', [])
    if anticipated:
        lines.append("## 7. ANTICIPATED QUESTIONS (AI-GENERATED)")
        for q in anticipated:
            lines.append(f"- {q.get('question', '')}")
            if q.get('rationale'):
                lines.append(f"  [Why: {q.get('rationale', '')}]")
        lines.append("")

    lines.append(f"## 8. NEXT STEPS\n{briefing_dict.get('next_steps', '')}\n")
    lines.append(f"## 9. COVERAGE NOTE\n{briefing_dict.get('coverage_note', '')}\n")
    return "\n".join(lines)

def expand_search_query(topic, api_key):
    """Use Gemini to expand a topic into related parliamentary search terms."""
    # Policy vocabulary is stable — cache for 7 days to skip Gemini on repeat searches
    try:
        cached = CachedTWFYSearch.get(topic, '_expand', ttl_hours=168)
        if cached is not None:
            return cached[0] if cached else f'"{topic}"'
    except Exception:
        pass
    try:
        model_path = get_working_model(api_key)
        ai_url = f"https://generativelanguage.googleapis.com/v1beta/{model_path}:generateContent?key={api_key}"
        prompt = (
            f"You are a UK Parliamentary researcher. A user wants to search Hansard for: \"{topic}\"\n"
            "Give up to 5 additional search phrases that refer to THE SAME SPECIFIC POLICY OR TOPIC "
            "using vocabulary MPs and Lords actually use in Parliament.\n"
            "STRICT RULES:\n"
            "- Only include terms that directly refer to this specific policy — not broader related topics\n"
            "- Include the policy acronym if one exists (e.g. 'LLE', 'SEND', 'HTQ', 'ICR')\n"
            "- Include any previous official name for the policy if it was renamed\n"
            "- Include the specific legislation name if relevant (e.g. 'Skills and Post-16 Education Act')\n"
            "- Do NOT include generic financial, legal, or policy terms that could match unrelated debates\n"
            "- Do NOT include synonyms that are broader than the original topic\n"
            "Return ONLY a JSON array of strings, no explanation, no markdown.\n"
            "Examples:\n"
            "'student loan repayments' → [\"repayment threshold\", \"Plan 2 loans\", \"income-contingent repayment\", \"ICR\", \"loan write-off\"]\n"
            "'lifelong learning entitlement' → [\"LLE\", \"lifelong loan entitlement\", \"Skills and Post-16 Education Act\", \"modular learning\"]"
        )
        payload = {"contents": [{"parts": [{"text": prompt}]}],
                   "generationConfig": {"responseMimeType": "application/json"}}
        resp = requests.post(ai_url, json=payload, timeout=15)
        raw = None
        if resp.status_code == 200:
            raw = resp.json()['candidates'][0]['content']['parts'][0]['text']
        else:
            raw = _claude_fallback(prompt, max_tokens=300)
        if raw:
            terms = _parse_ai_json(raw)
            if isinstance(terms, list) and terms:
                # Only include the original topic phrase if it's short enough to
                # appear verbatim in debates. Long policy descriptions (> 4 words)
                # are never said word-for-word and return 0 results from TWFY Xapian.
                if len(topic.split()) <= 4:
                    all_terms = [topic] + [str(t) for t in terms[:6]]
                else:
                    all_terms = [str(t) for t in terms[:6]]
                if not all_terms:
                    return f'"{topic}"'
                result = '(' + ' OR '.join(f'"{t}"' for t in all_terms) + ')'
                try:
                    CachedTWFYSearch.store(topic, '_expand', [result])
                except Exception:
                    pass
                return result
    except Exception:
        pass
    return f'"{topic}"'


_SOURCE_GID_PREFIX = {
    'commons':        'uk.org.publicwhip/debate/',
    'westminsterhall':'uk.org.publicwhip/westminhall/',
    'lords':          'uk.org.publicwhip/lords/',
    'wms':            'uk.org.publicwhip/wms/',
}

def _listurl_to_parent_gid(listurl, source):
    """Convert a TWFY listurl to the full TWFY speech GID for use with gid= parameter.
    Uses the actual speech position from the listurl — constructing a fake .0 parent
    causes 404 because that position may not exist in TWFY's database.
    Returns the full uk.org.publicwhip/PREFIX/DATE.SECTION.POSITION GID.
    Deduplication in fetch_all_debate_sessions uses just the section portion as the key.
    Returns None if listurl cannot be parsed, is a Hansard URL, or source is not expandable."""
    if not listurl or source not in _SOURCE_GID_PREFIX:
        return None
    if listurl.startswith('http'):
        return None  # Hansard API full URLs cannot expand via TWFY GID path
    try:
        m = re.search(r'\?id=([^&#]+)', listurl)
        if not m:
            return None
        id_val = m.group(1)                   # e.g. "2026-01-15.123.4" (actual speech position)
        if '.' not in id_val:
            return None
        return f"{_SOURCE_GID_PREFIX[source]}{id_val}"  # full GID with actual position
    except Exception:
        return None


def fetch_full_debate_session(parent_gid, source):
    """Fetch ALL speeches from a TWFY debate section via its parent GID.
    Returns normalised speech list (same schema as fetch_twfy_topic). relevance=0.
    Cached with 30-day TTL — published debate transcripts never change."""
    if not TWFY_API_KEY:
        return []
    cached = CachedTWFYSearch.get(parent_gid, f'session_{source}', ttl_hours=720)
    if cached is not None:
        return cached
    try:
        import logging as _sl
        api_url = TWFY_WMS_URL if source == 'wms' else TWFY_API_URL
        params = {'key': TWFY_API_KEY, 'gid': parent_gid, 'output': 'json'}
        if source != 'wms':
            params['type'] = source
        resp = requests.get(api_url, params=params, timeout=10)
        _sl.warning(f"[session_fetch] gid={parent_gid!r} type={params.get('type')!r} status={resp.status_code}")
        if resp.status_code != 200:
            return []
        rjson = resp.json()
        rows = rjson.get('rows', [])
        _sl.warning(f"[session_fetch] gid={parent_gid!r} rows_returned={len(rows)} response_keys={list(rjson.keys())} error={rjson.get('error','')!r}")
        if rows:
            sample = rows[0]
            _sl.warning(f"[session_fetch] sample_row keys={list(sample.keys())} speaker={sample.get('speaker')} hdate={sample.get('hdate')} listurl={sample.get('listurl','')[:80]!r}")
        results = []
        for r in rows:
            body_raw = r.get('body', '')
            body_text = clean_body_text(body_raw)
            debate_title = re.sub(r'<[^>]+>', '', r.get('parent', {}).get('body', '') or '')
            results.append({
                'listurl': r.get('listurl', ''),
                'body_clean': body_text[:500],
                'body_export': body_text[:3000],
                'body_word_count': len(body_text.split()),
                'speaker_name': (r.get('speaker') or {}).get('name', 'Unknown'),
                'speaker_party': _normalise_party((r.get('speaker') or {}).get('party', '')),
                'hdate': r.get('hdate', ''),
                'debate_title': debate_title,
                'source': source,
                'source_label': get_source_label(source),
                'relevance': 0,
                'debate_type': get_debate_type(debate_title, source=source),
                'from_session_fetch': True,
            })
        if results:
            CachedTWFYSearch.store(parent_gid, f'session_{source}', results)
        return results
    except Exception as e:
        import logging
        logging.warning(f"[session_expand] {parent_gid} {source}: {type(e).__name__}: {e}")
        return []


def fetch_all_debate_sessions(matched_rows, max_debates=15):
    """For each unique debate in matched_rows, fetch all speeches via TWFY GID lookup.
    This guarantees ministers appear even when their responses don't contain search keywords.
    Skips wrans (written answers have no multi-speaker session to expand).
    Routes by data source: rows with debate_section_ext_id use Hansard API;
    rows without (TWFY keyword search) fall back to TWFY GID expansion."""
    seen_sections = set()
    hansard_pairs = []   # (ext_id, source) — Hansard API expansion
    gid_source_pairs = []  # (gid, source) — TWFY GID expansion (fallback)

    for r in matched_rows:
        source = r.get('source', '')
        if source == 'wrans':
            continue
        title = r.get('debate_title', '').lower()
        if 'written answers' in title or 'written answer' in title:
            continue

        ext_id = r.get('debate_section_ext_id', '')
        if ext_id:
            if ext_id not in seen_sections:
                seen_sections.add(ext_id)
                hansard_pairs.append((ext_id, source))
        else:
            gid = _listurl_to_parent_gid(r.get('listurl', ''), source)
            if gid:
                last_dot = gid.rfind('.')
                section_key = gid[:last_dot] if last_dot != -1 else gid
                if section_key not in seen_sections:
                    seen_sections.add(section_key)
                    gid_source_pairs.append((gid, source))

        if len(hansard_pairs) + len(gid_source_pairs) >= max_debates:
            break

    import logging as _slog
    _slog.warning(f"[session_expand] {len(matched_rows)} rows in → "
                  f"{len(hansard_pairs)} hansard + {len(gid_source_pairs)} twfy sessions"
                  f" | hansard_sample={hansard_pairs[:2]} | twfy_sample={gid_source_pairs[:2]}")

    if not hansard_pairs and not gid_source_pairs:
        return []

    extra = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=8) as ex:
        futs = {}
        for ext_id, src in hansard_pairs:
            futs[ex.submit(copy_current_request_context(fetch_full_hansard_session), ext_id, src)] = ext_id
        for gid, src in gid_source_pairs:
            futs[ex.submit(copy_current_request_context(fetch_full_debate_session), gid, src)] = gid
        for f in concurrent.futures.as_completed(futs):
            try:
                extra.extend(f.result())
            except Exception:
                pass

    _slog.warning(f"[session_expand] {len(extra)} total speeches fetched from "
                  f"{len(hansard_pairs) + len(gid_source_pairs)} sessions")
    return extra


def _group_by_debate(rows):
    """Group speeches by debate (date + title). Within each group, ministers first.
    Returns list of (debate_key_dict, [rows]) sorted: minister-debates first, then most recent."""
    from collections import defaultdict
    groups = defaultdict(list)
    for r in rows:
        key = (r.get('hdate', ''), r.get('debate_title', ''))
        groups[key].append(r)
    for key in groups:
        groups[key].sort(key=lambda x: (not x.get('is_minister', False), -x.get('relevance', 0)))
    # Minister-containing debates first (1 > 0), then most recent date first (reverse string sort)
    group_list = sorted(
        groups.items(),
        key=lambda kv: (1 if any(r.get('is_minister') for r in kv[1]) else 0, kv[0][0]),
        reverse=True
    )
    result = []
    for k, v in group_list:
        matched = sum(1 for r in v if r.get('relevance', 0) > 0)
        if matched >= 3:
            rel_level = 'high'
        elif matched >= 1:
            rel_level = 'medium'
        else:
            rel_level = 'low'
        result.append({
            'date': k[0], 'title': k[1], 'speeches': v,
            'has_minister': any(r.get('is_minister') for r in v),
            'source_label': v[0].get('source_label', '') if v else '',
            'source': v[0].get('source', '') if v else '',
            'matched_count': matched,
            'relevance_level': rel_level,
        })
    return result


def _classify_group(grp):
    """Classify a debate group into a section.
    Uses title patterns first (definitive), then word-count heuristic.
    A group where no speech exceeds 300 words is likely Oral Questions —
    the minister's prepared answer is ~150 words, follow-ups shorter still.
    A 10-minute debate speech runs ~1300 words."""
    source = grp.get('source', '')
    title = grp.get('title', '').lower()
    speeches = grp.get('speeches', [])

    # Normalise dashes so em-dash/en-dash variants match consistently
    title_norm = re.sub(r'[–—]', '-', title)

    if source == 'wms':
        return 'statement'
    # Written answers must never reach _group_by_debate — filtered out upstream.
    # If any slip through, return 'wq' as a sentinel so they don't appear in Debates.
    if 'written answers' in title or 'written answer' in title:
        return 'wq'
    if (re.search(r'\bbill\b', title) or 'second reading' in title or 'third reading' in title
            or 'committee stage' in title or 'report stage' in title
            or 'lords amendments' in title or 'royal assent' in title):
        return 'legislation'
    if 'urgent question' in title:
        return 'urgent'
    # Commons oral statements: titled "Statement on X" (source=commons, not wms)
    if source == 'commons' and 'statement' in title:
        return 'statement'
    if 'oral answers' in title or 'question time' in title:
        return 'oral'
    if 'prime minister' in title and 'question' in title:
        return 'oral'
    # Lords oral questions — normalised title so em-dash/hyphen variants both match
    if source == 'lords' and (title_norm.rstrip().endswith('- question') or
                              '- oral question' in title_norm or
                              '- question' in title_norm):
        return 'oral'
    # Commons oral questions: "Department: Topic" format
    _OQ_DEPT_PREFIXES = (
        'education', 'health', 'treasury', 'home office', 'defence', 'justice',
        'transport', 'environment', 'science', 'work and pensions', 'cabinet office',
        'foreign', 'culture', 'housing', 'digital', 'northern ireland', 'business',
        'wales', 'scotland', 'attorney general', 'leader of the house',
        'women and equalities', 'energy security', 'levelling up',
        'international trade', 'international development',
    )
    if source == 'commons' and any(title.startswith(p + ':') for p in _OQ_DEPT_PREFIXES):
        return 'oral'
    if 'topical question' in title:
        return 'oral'
    # Word-count heuristic: short speeches are likely Oral PQ follow-ups.
    # Applied to both commons and lords — threshold 400 words.
    if speeches:
        max_words = max((r.get('body_word_count', 0) for r in speeches), default=0)
        if source in ('commons', 'lords') and 0 < max_words < 400:
            return 'oral'
    return 'debate'


def _fetch_topic_wqs(topic, start_date, end_date, selected_depts, limit=100):
    """Fetch WQs matching topic from Parliament API. Returns (list_of_dicts, total_count).
    Does NOT pass answeringBodies to the API — that filter causes severe timeouts (>30s).
    Instead fetches without dept filter and filters client-side by answeringBodyName."""
    import logging
    try:
        # No answeringBodies param — server-side dept filter causes 30s+ timeouts.
        # Client-side filter is applied below instead.
        params = {'searchTerm': topic, 'take': limit, 'skip': 0,
                  'expandMember': 'false'}
        if start_date:
            params['tabledWhenFrom'] = start_date
        if end_date:
            params['tabledWhenTo'] = end_date
        resp = requests.get(PARLIAMENT_WQ_API, params=params, timeout=30)
        if resp.status_code != 200:
            logging.warning(f"WQ fetch failed: {resp.status_code}")
            return [], 0
        data = resp.json()
        total = data.get('totalResults', 0)
        # Client-side dept filter — match answeringBodyName against selected dept names
        dept_names = set(selected_depts)
        results = []
        seen_uins = set()
        for item in data.get('results', []):
            val = item.get('value', {})
            uin = str(val.get('uin', ''))
            if uin and uin in seen_uins:
                continue  # deduplicate — API can return same WQ if search matches both Q and A
            if uin:
                seen_uins.add(uin)
            # Client-side dept filter
            answering_body = val.get('answeringBodyName', '')
            if dept_names and answering_body not in dept_names:
                continue
            raw_date = (val.get('dateTabled') or '').split('T')[0]
            date_ans = (val.get('dateAnswered') or '').split('T')[0]
            answer_raw = val.get('answerText') or ''
            answer_clean = re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', answer_raw)).strip()
            # Use cleaned answer text (not raw HTML) to determine answered status;
            # raw HTML can contain only tags/whitespace, giving a false positive.
            is_answered = bool(answer_clean or date_ans)
            is_withdrawn = bool(val.get('isWithdrawn'))
            is_holding = (is_answered and not is_withdrawn and
                          'i will write' in answer_clean.lower()[:200])
            q_text = re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', val.get('questionText') or '')).strip()
            asking = val.get('askingMember') or {}
            party = asking.get('party', '')
            membership = asking.get('latestHouseMembership') or {}
            house = 'Lords' if membership.get('house') == 2 else 'Commons'
            constituency = membership.get('membershipFrom', '')
            role = 'Life Peer' if house == 'Lords' else (f'MP for {constituency}' if constituency else 'MP')
            minister = val.get('answeringMember') or {}
            results.append({
                'uin': uin,
                'dept': answering_body,
                'question_text': q_text,
                'answer_text': answer_clean[:1500],
                'answering_minister': minister.get('name', ''),
                'asking_mp': asking.get('name') or asking.get('listAs', ''),
                'role': role,
                'party': party,
                'party_colour': PARTY_COLOURS_RESEARCH.get(party, '#888'),
                'date_tabled': raw_date,
                'date_answered': date_ans,
                'is_answered': is_answered,
                'is_holding': is_holding,
                'is_withdrawn': is_withdrawn,
                'heading': val.get('heading', ''),
                'url': f"https://questions-statements.parliament.uk/written-questions/detail/{raw_date}/{uin}",
            })
        return results, total
    except Exception as e:
        import logging
        logging.error(f"WQ fetch exception: {type(e).__name__}: {e}")
        return [], 0


def _display_name(title):
    """Strip honorifics/post-nominals from a GOV.UK title for clean display.
    Preserves Lords titles (Baroness/Lord) so TWFY lookups can find peers.
    'The Rt Hon Baroness Smith of Malvern' → 'Baroness Smith of Malvern'
    'The Rt Hon Josh MacAlister MP' → 'Josh MacAlister'"""
    name = title.strip()
    # Strip Rt Hon / Right Hon only (not Lords titles)
    for prefix in ['The Rt Hon ', 'The Right Hon ', 'Rt Hon ', 'Right Hon ', 'The ']:
        if name.startswith(prefix):
            name = name[len(prefix):]
            break
    # Strip commons-only honorifics (Dame, Sir, Dr, Mr etc) but NOT Baroness/Lord
    for prefix in ['Dame ', 'Sir ', 'Dr ', 'Mr ', 'Mrs ', 'Ms ', 'Miss ']:
        if name.startswith(prefix):
            name = name[len(prefix):]
            break
    name = re.sub(r'\s+(MP|OBE|CBE|MBE|PC|QC|KC|DBE|KBE)(\b.*)?$', '', name)
    return name.strip()


def _fetch_govuk_dept_ministers(dept_base_path, dept_name):
    """Fetch ordered_ministers for one GOV.UK department page.
    Returns list of (normalised_name, role_string, display_name, dept_name) tuples."""
    results = []
    try:
        r = requests.get(f'https://www.gov.uk/api/content{dept_base_path}', timeout=8)
        if r.status_code == 200:
            for minister in r.json().get('links', {}).get('ordered_ministers', []):
                title = minister.get('title', '').strip()
                norm = _normalise_name(title)
                display = _display_name(title)
                if norm and display:
                    results.append((norm, f'Minister ({dept_name})', display, dept_name))
    except Exception:
        pass
    return results


def get_minister_list():
    """Fetch ALL current government ministers from GOV.UK content API.
    Returns dict with 'by_norm' {normalised_name: role} keyed by _normalise_name().
    File-cached for 24 hours. Includes junior ministers, not just Cabinet.
    Falls back to Parliament GovernmentPosts if GOV.UK is unavailable."""
    cache = {}
    try:
        if os.path.exists(MINISTER_CACHE_FILE):
            with open(MINISTER_CACHE_FILE) as f:
                cache = json.load(f)
        if (cache.get('_ts') and time.time() - cache['_ts'] < MINISTER_CACHE_TTL
                and cache.get('_source') == 'govuk'
                and cache.get('_version') == MINISTER_CACHE_VERSION
                and cache.get('by_norm') and cache.get('by_dept')):
            return cache
    except Exception:
        pass

    by_norm = {}   # normalised_name → role
    by_id = {}     # parliament_member_id → role
    by_dept = {}   # dept_name → [{"name": display_name}]

    try:
        # Step 1: GOV.UK ministers page — Cabinet + dept links
        r = requests.get('https://www.gov.uk/api/content/government/ministers', timeout=10)
        if r.status_code == 200:
            links = r.json().get('links', {})

            # Cabinet ministers (direct list on the page)
            for section in ['ordered_cabinet_ministers', 'ordered_also_attends_cabinet']:
                for m in links.get(section, []):
                    norm = _normalise_name(m.get('title', ''))
                    if norm:
                        by_norm[norm] = 'Cabinet Minister'

            # All departments — fetch in parallel
            depts = links.get('ordered_ministerial_departments', [])
            with concurrent.futures.ThreadPoolExecutor(max_workers=8) as ex:
                futures = {
                    ex.submit(_fetch_govuk_dept_ministers, d['base_path'], d.get('title', 'Government')): d
                    for d in depts if d.get('base_path')
                }
                for future in concurrent.futures.as_completed(futures):
                    for norm, role, display, dept in future.result():
                        if norm not in by_norm:
                            by_norm[norm] = role
                        by_dept.setdefault(dept, [])
                        if display and not any(m['name'] == display for m in by_dept[dept]):
                            by_dept[dept].append({'name': display})
            for dept in by_dept:
                by_dept[dept].sort(key=lambda m: m['name'])
    except Exception:
        pass

    # Fallback / supplement: Parliament GovernmentPosts (Cabinet only — populates by_id)
    try:
        r = requests.get(
            'https://members-api.parliament.uk/api/Posts/GovernmentPosts', timeout=10
        )
        if r.status_code == 200:
            for post in r.json():
                v = post.get('value', {})
                role = v.get('name', '')
                for holder in v.get('postHolders', []):
                    if holder.get('endDate'):
                        continue  # skip past holders
                    m = (holder.get('member') or {}).get('value', {})
                    name = m.get('nameDisplayAs', '')
                    member_id = m.get('id')
                    norm = _normalise_name(name)
                    if norm and norm not in by_norm:
                        by_norm[norm] = role
                    if member_id and role:
                        by_id[str(member_id)] = role
    except Exception:
        pass

    # Preserve twfy_ids from the old cache — TWFY person IDs are permanent identifiers
    # assigned to a person for life. No need to re-resolve them after a minister list refresh.
    old_twfy_ids = cache.get('twfy_ids', {})
    # Seed known IDs for ministers whose TWFY name search fails (newer Lords peers etc.)
    # Use unconditional assignment so verified IDs always override stale/failed cache entries.
    for name, entry in KNOWN_TWFY_IDS.items():
        old_twfy_ids[name] = entry
    data = {'_ts': time.time(), '_source': 'govuk', '_version': MINISTER_CACHE_VERSION,
            'by_norm': by_norm, 'by_id': by_id, 'by_dept': by_dept, 'twfy_ids': old_twfy_ids}
    try:
        with open(MINISTER_CACHE_FILE, 'w') as f:
            json.dump(data, f)
    except Exception:
        pass

    return data


def _normalise_name(name):
    """Strip honorifics, titles, and post-nominal letters for name matching.
    Works on both TWFY speaker names and GOV.UK person titles.
    Applies prefix stripping in two passes to handle 'The Rt Hon Baroness ...'."""
    n = name.lower().strip()
    # Strip trailing post-nominals: MP, OBE, CBE, MBE, PC, QC, KC, etc.
    n = re.sub(r'\s+(obe|cbe|mbe|pc|qc|kc|dbe|kbe|mp|obe mp|cbe mp|mbe mp)\s*$', '', n)
    # Two passes of prefix stripping handles 'the rt hon baroness' → 'baroness' → ''
    prefixes = ['the rt hon ', 'the right hon ', 'rt hon ', 'right hon ',
                'the baroness ', 'baroness ', 'the lord ', 'lord ',
                'dame ', 'sir ', 'dr ', 'mr ', 'mrs ', 'ms ', 'miss ']
    for _ in range(2):
        for prefix in prefixes:
            if n.startswith(prefix):
                n = n[len(prefix):]
                break
    return n.strip()


def verify_government_speaker(name):
    """Confirm a speaker holds a current government post.
    Returns dict with confirmed (bool) and role (str).

    Three-tier approach:
    1. Normalised name match against GOV.UK minister cache (covers all ministers)
    2. Parliament Members Search → check each result ID against Cabinet id cache
    3. Parliament Members Biography → governmentPosts with no endDate (gets precise role)
    """
    norm_name = _normalise_name(name)
    if norm_name and norm_name in _VERIFY_CACHE:
        return _VERIFY_CACHE[norm_name]

    minister_data = get_minister_list()
    by_norm = minister_data.get('by_norm', {})
    by_id = minister_data.get('by_id', {})

    # 1. Normalised match against GOV.UK full minister list
    if norm_name and norm_name in by_norm:
        result = {'confirmed': True, 'role': by_norm[norm_name]}
        if norm_name:
            _VERIFY_CACHE[norm_name] = result
        return result

    # 2. Parliament Members Search → ID → Cabinet id cache
    try:
        s_resp = requests.get(
            'https://members-api.parliament.uk/api/Members/Search',
            params={'Name': name, 'IsCurrentMember': 'true', 'take': 5},
            timeout=5
        )
        if s_resp.status_code == 200:
            items = s_resp.json().get('items', [])
            for item in items:
                member_id = str(item['value']['id'])
                if member_id in by_id:
                    result = {'confirmed': True, 'role': by_id[member_id]}
                    if norm_name:
                        _VERIFY_CACHE[norm_name] = result
                    return result

            # 3. Biography endpoint — has full government post history with endDate
            if items:
                member_id = str(items[0]['value']['id'])
                bio_resp = requests.get(
                    f'https://members-api.parliament.uk/api/Members/{member_id}/Biography',
                    timeout=5
                )
                if bio_resp.status_code == 200:
                    govt_posts = bio_resp.json().get('value', {}).get('governmentPosts', [])
                    current_posts = [p for p in govt_posts if not p.get('endDate')]
                    if current_posts:
                        result = {'confirmed': True, 'role': current_posts[0].get('name', '')}
                        if norm_name:
                            _VERIFY_CACHE[norm_name] = result
                        return result
    except Exception:
        pass

    result = {'confirmed': False, 'role': ''}
    if norm_name:
        _VERIFY_CACHE[norm_name] = result
    return result


def lookup_twfy_person(name):
    """Look up TWFY person_id by name. Checks MPs first, then Lords.
    Tries multiple name variants to handle titles like 'Baroness Smith of Malvern'.
    Returns (person_id, matched_name, is_lord) or (None, None, False)."""
    # Build variants: original → stripped display → normalised title-case
    variants = [name]
    display = _display_name(name)
    if display and display != name:
        variants.append(display)
    norm_title = _normalise_name(name).title()
    if norm_title and norm_title not in variants:
        variants.append(norm_title)

    # If the name has no Lords title prefix, also try with common prefixes.
    # This handles "Smith of Malvern" → "Baroness Smith of Malvern" for TWFY getLords.
    lords_prefixes = ['Baroness ', 'Lord ', 'Baron ']
    has_lords_prefix = any(name.lower().startswith(p.lower()) for p in lords_prefixes + ['The '])
    if not has_lords_prefix:
        for prefix in lords_prefixes:
            candidate = prefix + name
            if candidate not in variants:
                variants.append(candidate)

    for variant in variants:
        # For variants with a Lords title prefix or a peerage 'of' pattern (e.g.
        # "Smith of Malvern", "Baroness Smith of Malvern"), check getLords FIRST
        # and skip getMPs entirely — getMPs search on "Smith" would return the
        # wrong person (any MP with that surname) before Lords is tried.
        has_lords_prefix_v = any(variant.lower().startswith(p.lower()) for p in lords_prefixes)
        has_peerage_of = ' of ' in variant.lower()
        if has_lords_prefix_v or has_peerage_of:
            endpoint_pairs = [('getLords', True)]
        else:
            endpoint_pairs = [('getMPs', False), ('getLords', True)]
        for endpoint, is_lord in endpoint_pairs:
            try:
                resp = requests.get(
                    'https://www.theyworkforyou.com/api/' + endpoint,
                    params={'key': TWFY_API_KEY, 'search': variant, 'output': 'json'},
                    timeout=10
                )
                if resp.status_code == 200:
                    data = resp.json()
                    if isinstance(data, list) and data:
                        return data[0].get('person_id'), data[0].get('name'), is_lord
            except Exception:
                pass

    # Fallback for Lords: getLords name search fails for newer peers (created ~2023+)
    # whose records aren't indexed for name search.
    # Strategy: fetch that peer's recent Lords speeches via a broad search, then
    # scan speaker fields for a name normalisation match.
    # We try multiple search terms to maximise the chance of finding a recent speech.
    is_lords_name = any(name.lower().startswith(p.lower()) for p in ['baroness ', 'lord ', 'baron ']) \
                    or ' of ' in name.lower()
    if is_lords_name and TWFY_API_KEY:
        # Build candidate search terms: stripped display name, family name, territorial name
        display = _display_name(name)   # e.g. "Baroness Smith of Malvern"
        parts = display.split()
        # ["Baroness", "Smith", "of", "Malvern"] → try "Smith", "Malvern", "Smith of Malvern"
        search_terms = []
        if ' of ' in display.lower():
            before_of = display.split(' of ')[0].split()
            after_of  = display.split(' of ')[1].split()
            if before_of: search_terms.append(before_of[-1])       # family name: "Smith"
            if after_of:  search_terms.append(after_of[0])         # territorial: "Malvern"
        else:
            if len(parts) > 1: search_terms.append(parts[-1])
        for search_term in search_terms:
            try:
                resp = requests.get(
                    TWFY_API_URL,
                    params={'key': TWFY_API_KEY, 'search': search_term, 'type': 'lords',
                            'num': 100, 'order': 'd', 'output': 'json'},
                    timeout=10
                )
                if resp.status_code == 200:
                    for row in resp.json().get('rows', []):
                        spk = row.get('speaker') or {}
                        spk_name = spk.get('name', '')
                        spk_pid = spk.get('person_id')
                        if spk_pid and _normalise_name(spk_name) == _normalise_name(display):
                            return str(spk_pid), spk_name, True
            except Exception:
                pass

    return None, None, False


def fetch_twfy_minister_topic(person_id, topic, date_range, sources, num=50, is_lord=False):
    """Fetch speeches for a specific TWFY person_id, filtered by topic + date.
    Returns rows in the same normalised schema as fetch_twfy_topic() so they
    can be merged and deduped with keyword-search results.
    Results cached for 6h to reduce TWFY API quota usage."""
    # Lords ministers only speak in Lords (+ wms), not Commons/WestminsterHall.
    # Filtering here saves ~2 API calls per Lords minister per search.
    if is_lord:
        sources = [s for s in sources if s in ('lords', 'wms')]

    rows = []
    for source in sources:
        cache_key_query = f"minister:{person_id}:{topic} {date_range}".strip()
        cached = CachedTWFYSearch.get(cache_key_query, source, ttl_hours=6)
        if cached is not None:
            rows.extend(cached)
            continue
        api_url = TWFY_WMS_URL if source == 'wms' else TWFY_API_URL
        query = f"{topic} {date_range}".strip() if topic else date_range.strip()
        params = {
            'key': TWFY_API_KEY, 'person': str(person_id),
            'order': 'd', 'num': num, 'output': 'json'
        }
        if query:
            params['search'] = query
        if source != 'wms':
            params['type'] = source
        try:
            resp = requests.get(api_url, params=params, timeout=15)
            if resp.status_code != 200:
                continue
            source_rows = []
            for r in resp.json().get('rows', []):
                body_raw = r.get('body', '')
                body_text = clean_body_text(body_raw)
                debate_title = re.sub(r'<[^>]+>', '', r.get('parent', {}).get('body', '') or '')
                if source == 'wms':
                    parent_body = re.sub(r'<[^>]+>', '', r.get('parent', {}).get('body', '') or '')
                    debate_title = parent_body if parent_body else re.sub(r'<[^>]+>', '', body_raw)[:80]
                elif source == 'wrans':
                    debate_title = re.sub(r'<[^>]+>', '', body_raw)[:80]
                dtype = get_debate_type(debate_title, source=source)
                source_rows.append({
                    'listurl': r.get('listurl', ''),
                    'body_clean': body_text[:500],
                    'body_export': body_text[:3000],
                    'body_word_count': len(body_text.split()),
                    'speaker_name': (r.get('speaker') or {}).get('name', 'Unknown'),
                    'speaker_party': _normalise_party((r.get('speaker') or {}).get('party', '')),
                    'hdate': r.get('hdate', ''),
                    'debate_title': debate_title,
                    'source': source,
                    'source_label': get_source_label(source),
                    'relevance': r.get('relevance', 0),
                    'debate_type': dtype,
                })
            # TWFY doesn't reliably filter by date when combined with person= param.
            # Apply Python-level filter before caching to prevent stale speeches leaking in.
            if date_range and '..' in date_range:
                parts = date_range.split('..')
                d_from = parts[0].strip()  # YYYYMMDD
                d_to   = parts[1].strip()
                # Convert YYYYMMDD → YYYY-MM-DD for comparison with hdate
                def _fmt(d): return f"{d[:4]}-{d[4:6]}-{d[6:]}" if len(d) == 8 else d
                d_from_iso = _fmt(d_from) if d_from else ''
                d_to_iso   = _fmt(d_to)   if d_to   else ''
                source_rows = [r for r in source_rows
                               if (not d_from_iso or r.get('hdate', '') >= d_from_iso)
                               and (not d_to_iso or r.get('hdate', '') <= d_to_iso)]
            CachedTWFYSearch.store(cache_key_query, source, source_rows)
            rows.extend(source_rows)
        except Exception:
            pass
    return rows


# --- Hansard API helpers (Phase 1: minister search) ---

_HANSARD_PARTY_ABBREV = {
    'lab': 'Labour', 'con': 'Conservative', 'ld': 'Liberal Democrat',
    'lib dem': 'Liberal Democrat', 'snp': 'Scottish National Party',
    'green': 'Green Party', 'reform': 'Reform UK', 'pc': 'Plaid Cymru',
    'dup': 'Democratic Unionist Party', 'ind': 'Independent',
    'crossbench': 'Crossbench',
}


def _parse_hansard_party(attributed_to):
    """Extract full party name from AttributedTo string.
    Handles: 'Lord Sikka (Lab)', 'Baroness Barran (Con)',
    'Baroness in Waiting/Government Whip (Baroness Blake of Leeds) (Lab)'
    Party is always the last (...) group."""
    m = re.search(r'\(([^)]+)\)\s*$', attributed_to or '')
    if m:
        abbrev = m.group(1).strip().lower()
        return _HANSARD_PARTY_ABBREV.get(abbrev, m.group(1).strip())
    return ''


def _extract_attributed_name(attributed_to):
    """Extract clean speaker name from AttributedTo string.
    Handles: 'Lord Sikka (Lab)' → 'Lord Sikka'
             'Baroness Barran (Con)' → 'Baroness Barran'
             'Baroness in Waiting/Government Whip (Baroness Blake of Leeds) (Lab)' → 'Baroness Blake of Leeds'
             'Lord Sikka' (no party) → 'Lord Sikka'"""
    if not attributed_to:
        return ''
    parens = re.findall(r'\(([^)]+)\)', attributed_to)
    if len(parens) >= 2:
        # Role (Name) (Party) format — name is the second-to-last group
        return parens[-2].strip()
    elif len(parens) == 1:
        # Name (Party) format — name is everything before the opening '('
        return attributed_to[:attributed_to.rfind('(')].strip()
    return attributed_to.strip()


def _hansard_section_to_source(house, section):
    """Map Hansard API House + Section fields to TWFY-compatible source codes."""
    section_lower = (section or '').lower()
    if 'written statement' in section_lower:
        return 'wms'
    if 'written answer' in section_lower:
        return 'wrans'
    if 'westminster hall' in section_lower:
        return 'westminsterhall'
    if (house or '').lower() == 'lords':
        return 'lords'
    return 'commons'


def fetch_hansard_minister_topic(parliament_id, topic, date_range, sources, num=50, is_lord=False):
    """Fetch minister speeches via Hansard API using Parliament member ID directly.
    Replaces fetch_twfy_minister_topic() when SEARCH_BACKEND=hansard.
    Returns same normalised row schema — zero changes to grouping or display code."""
    if not parliament_id:
        return []

    # Parse YYYYMMDD..YYYYMMDD → ISO dates for the Hansard API
    start_date_api = ''
    end_date_api = ''
    if date_range and '..' in date_range:
        parts = date_range.split('..')
        def _fmt(d): return f"{d[:4]}-{d[4:6]}-{d[6:]}" if len(d) == 8 else d
        start_date_api = _fmt(parts[0].strip()) if parts[0].strip() else ''
        end_date_api = _fmt(parts[1].strip()) if parts[1].strip() else ''

    import logging as _ml
    cache_key_query = f"hansard:minister:{parliament_id}:{topic} {date_range}".strip()
    cached = CachedTWFYSearch.get(cache_key_query, 'hansard_minister', ttl_hours=6)
    if cached is not None:
        _ml.warning(f"[minister_search] parliament_id={parliament_id} from_cache=True rows={len(cached)}")
        return cached

    params = {'queryParameters.memberId': parliament_id, 'take': num, 'skip': 0}
    if is_lord:
        params['queryParameters.house'] = 'Lords'
    if topic:
        params['queryParameters.searchTerm'] = topic
    if start_date_api:
        params['queryParameters.startDate'] = start_date_api
    if end_date_api:
        params['queryParameters.endDate'] = end_date_api

    try:
        import logging as _ml
        resp = requests.get(f"{HANSARD_API_BASE}/search.json", params=params, timeout=25)
        if resp.status_code != 200:
            _ml.warning(f"[minister_search] parliament_id={parliament_id} status={resp.status_code}")
            return []
        data = resp.json()
        results = data.get('Contributions', [])
        _ml.warning(f"[minister_search] parliament_id={parliament_id} contributions={len(results)} total={data.get('TotalContributions','?')}")

        rows = []
        for r in results:
            house = r.get('House', 'Commons')
            section = r.get('Section', '')
            source = _hansard_section_to_source(house, section)
            if source not in sources:
                continue

            body_raw = r.get('ContributionTextFull', '') or ''
            body_text = clean_body_text(body_raw)
            debate_title = r.get('DebateSection', '') or ''

            sitting_date = r.get('SittingDate', '')
            hdate = sitting_date[:10] if sitting_date else ''

            ext_id = r.get('DebateSectionExtId', '') or ''
            house_path = 'lords' if house.lower() == 'lords' else 'commons'
            listurl = (f"https://hansard.parliament.uk/{house_path}/{hdate}/debates/{ext_id}/"
                       if ext_id and hdate else '')

            dtype = get_debate_type(debate_title, source=source)
            rows.append({
                'listurl': listurl,
                'body_clean': body_text[:500],
                'body_export': body_text[:3000],
                'body_word_count': len(body_text.split()),
                'speaker_name': r.get('MemberName', '') or '',
                'speaker_party': _parse_hansard_party(r.get('AttributedTo', '')),
                'hdate': hdate,
                'debate_title': debate_title,
                'source': source,
                'source_label': get_source_label(source),
                'relevance': r.get('Rank', 0),
                'debate_type': dtype,
                'debate_section_ext_id': ext_id,
            })

        with_ext = sum(1 for r in rows if r.get('debate_section_ext_id'))
        _ml.warning(f"[minister_search] parliament_id={parliament_id} rows={len(rows)} with_ext_id={with_ext} sources={set(r['source'] for r in rows)}")
        if rows:  # never cache empty results — empty means API miss, not "no speeches"
            CachedTWFYSearch.store(cache_key_query, 'hansard_minister', rows)
        return rows
    except Exception as e:
        import logging
        logging.warning(f"[minister_search] parliament_id={parliament_id} exception={type(e).__name__}: {e}")
        return []


def fetch_full_hansard_session(ext_id, source):
    """Fetch ALL speeches from a Hansard debate section via DebateSectionExtId.
    Used instead of fetch_full_debate_session() for rows that came via the Hansard API.
    Returns normalised speech list with same schema as fetch_twfy_topic. relevance=0."""
    if not ext_id:
        return []
    cache_key = f"hansard:session:{ext_id}"
    cached = CachedTWFYSearch.get(cache_key, f'session_{source}', ttl_hours=720)
    if cached is not None:
        return cached
    try:
        import logging as _hl
        resp = requests.get(f"{HANSARD_API_BASE}/debates/debate/{ext_id}.json", timeout=15)
        _hl.warning(f"[hansard_session] ext_id={ext_id!r} status={resp.status_code}")
        if resp.status_code != 200:
            return []
        data = resp.json()

        # Flatten Items from top-level debate and all nested ChildDebates
        items = []

        def _collect(node):
            for item in node.get('Items', []):
                items.append(item)
            for child in node.get('ChildDebates', []):
                _collect(child)

        _collect(data)
        _hl.warning(f"[hansard_session] ext_id={ext_id!r} items={len(items)}")
        if items:
            s = items[0]
            _hl.warning(f"[hansard_session] sample keys={list(s.keys())[:12]} attributed={s.get('AttributedTo','')!r}")

        # Metadata lives in Overview, not at the root level
        overview = data.get('Overview') or {}
        debate_title = overview.get('Title', '') or ''
        raw_date = overview.get('Date', '') or ''
        hdate = raw_date[:10] if raw_date else ''
        house = overview.get('House', 'Commons') or 'Commons'
        house_path = 'lords' if 'lord' in house.lower() else 'commons'
        base_url = f"https://hansard.parliament.uk/{house_path}/{hdate}/debates/{ext_id}/"

        results = []
        for item in items:
            # Skip procedural items (no speaker attached)
            if not item.get('MemberId') and not item.get('AttributedTo'):
                continue
            body_raw = item.get('Value', '') or ''
            if not body_raw:
                continue
            body_text = clean_body_text(body_raw)
            if not body_text.strip():
                continue

            attributed_to = item.get('AttributedTo', '') or ''
            member_name = _extract_attributed_name(attributed_to)
            item_date = item.get('SittingDate', '') or raw_date
            item_hdate = item_date[:10] if item_date else hdate

            results.append({
                'listurl': base_url,
                'body_clean': body_text[:500],
                'body_export': body_text[:3000],
                'body_word_count': len(body_text.split()),
                'speaker_name': member_name,
                'speaker_party': _parse_hansard_party(attributed_to),
                'hdate': item_hdate,
                'debate_title': debate_title,
                'source': source,
                'source_label': get_source_label(source),
                'relevance': 0,
                'debate_type': get_debate_type(debate_title, source=source),
                'from_session_fetch': True,
                'debate_section_ext_id': ext_id,
            })

        if results:
            CachedTWFYSearch.store(cache_key, f'session_{source}', results)
        return results
    except Exception as e:
        import logging
        logging.warning(f"[hansard_session] {ext_id} {source}: {type(e).__name__}: {e}")
        return []


def _resolve_parliament_id(display_name):
    """Look up the Parliament member ID for a minister by name.
    Returns (parliament_id, house) or (None, None)."""
    try:
        resp = requests.get(
            'https://members-api.parliament.uk/api/Members/Search',
            params={'Name': display_name, 'IsCurrentMember': 'true', 'take': 3},
            timeout=5
        )
        if resp.status_code == 200:
            items = resp.json().get('items', [])
            if items:
                v = items[0]['value']
                house_num = (v.get('latestHouseMembership') or {}).get('house', 1)
                house = 'Lords' if house_num == 2 else 'Commons'
                return v['id'], house
    except Exception:
        pass
    return None, None


def lookup_twfy_person_by_parliament_id(parliament_id, display_name, house):
    """DB-first TWFY person ID lookup.
    1. Check MemberLink DB — instant, zero API calls if already resolved.
    2. If not resolved and not marked failed, run lookup_twfy_person() and store result.
    Returns (twfy_person_id, is_lord) or (None, False).
    """
    from cache_models import MemberLink
    row = MemberLink.get_by_parliament_id(parliament_id)
    if row:
        if row.twfy_person_id:
            return row.twfy_person_id, row.house == 'Lords'
        if row.lookup_failed:
            return None, house == 'Lords'
        # Row exists but unresolved — fall through to lookup

    person_id, matched_name, is_lord = lookup_twfy_person(display_name)
    method = 'twfy_name_search' if person_id else 'failed'
    MemberLink.upsert(
        parliament_id=parliament_id,
        display_name=display_name,
        house=house,
        twfy_person_id=person_id,
        twfy_name=matched_name,
        resolution_method=method,
        lookup_failed=(person_id is None),
    )
    return person_id, is_lord


def seed_all_minister_links(app):
    """Pre-populate MemberLink with all current government ministers.
    Runs once in a background thread at startup. Skips if already well-seeded
    (> 40 resolved entries means a previous run completed successfully).
    This converts MemberLink from lazy per-search cache to a pre-built reference
    table — minister-led searches become instant with zero quota cost after first seed.
    """
    import threading

    def _run():
        from cache_models import MemberLink
        with app.app_context():
            try:
                resolved_count = MemberLink.query.filter(
                    MemberLink.twfy_person_id.isnot(None)
                ).count()
                if resolved_count > 40:
                    return  # already well-seeded — skip
                minister_data = get_minister_list()
                # by_dept: {dept_name: [{"name": display_name}, ...]}
                by_dept = minister_data.get('by_dept', {})
                # Flatten to unique display names across all departments
                seen = set()
                all_names = []
                for dept_ministers in by_dept.values():
                    for m in dept_ministers:
                        name = m.get('name', '').strip()
                        if name and name not in seen:
                            seen.add(name)
                            all_names.append(name)
                if not all_names:
                    return
                seeded = 0
                for display_name in all_names:
                    if not display_name:
                        continue
                    # Skip if already resolved or marked failed
                    existing = MemberLink.query.filter(
                        MemberLink.display_name == display_name
                    ).first()
                    if existing and (existing.twfy_person_id or existing.lookup_failed):
                        continue
                    # Resolve Parliament ID then TWFY person ID
                    parliament_id, house = _resolve_parliament_id(display_name)
                    if not parliament_id:
                        continue
                    lookup_twfy_person_by_parliament_id(parliament_id, display_name, house)
                    seeded += 1
                    time.sleep(0.3)  # gentle rate limiting — ~3 calls/sec
            except Exception as e:
                print(f'[seed_minister_links] error: {e}')

    t = threading.Thread(target=_run, daemon=True)
    t.start()


def get_dept_minister_twfy_ids(dept_name, minister_data):
    """Resolve all ministers for a department to TWFY person IDs.
    Returns list of {person_id, name, role, is_lord} dicts.

    Lookup chain (cheapest first):
    1. MemberLink DB — zero TWFY API calls if already resolved
    2. JSON twfy_ids_cache — legacy fast path during transition
    3. Parliament Members/Search → lookup_twfy_person_by_parliament_id
    """
    from cache_models import MemberLink
    ministers = minister_data.get('by_dept', {}).get(dept_name, [])
    twfy_ids_cache = minister_data.setdefault('twfy_ids', {})
    results = []
    json_updated = False

    for m in ministers:
        display = m.get('display_name') or m.get('name', '')
        if not display:
            continue

        # 1. DB lookup by display name (via MemberLink index on display_name would be ideal,
        #    but we use parliament_id as key so we need it first — check JSON cache for it)
        cached_entry = twfy_ids_cache.get(display, {})

        # Fast path: JSON cache already has the TWFY ID
        if cached_entry.get('person_id'):
            results.append({
                'person_id': cached_entry['person_id'],
                'name': display,
                'role': m.get('role', ''),
                'is_lord': cached_entry.get('is_lord', False),
            })
            continue

        # Skip if previously marked as failed in JSON cache
        if cached_entry.get('lookup_failed'):
            continue

        # Resolve Parliament ID → DB lookup → TWFY lookup
        parliament_id, house = _resolve_parliament_id(display)
        if parliament_id:
            person_id, is_lord = lookup_twfy_person_by_parliament_id(
                parliament_id, display, house)
        else:
            # Parliament ID lookup failed — fall back to name-only TWFY lookup
            person_id, matched_name, is_lord = lookup_twfy_person(display)

        # Write result back to JSON cache so next search is instant
        twfy_ids_cache[display] = {
            'person_id': person_id, 'is_lord': is_lord,
            'lookup_failed': person_id is None
        }
        json_updated = True

        if person_id:
            results.append({
                'person_id': person_id, 'name': display,
                'role': m.get('role', ''), 'is_lord': is_lord
            })
        else:
            import logging
            logging.warning(f"TWFY person ID lookup failed for minister: {display!r}")

    if json_updated:
        try:
            with open(MINISTER_CACHE_FILE, 'w') as f:
                json.dump(minister_data, f)
        except Exception:
            pass
    return results


def fetch_minister_debates(person_id, topic, date_range, house_filter='all'):
    """Fetch debates where a minister spoke, deduped by unique debate section."""
    sources = []
    if house_filter in ('all', 'commons'):
        sources.extend(['commons', 'westminsterhall'])
    if house_filter in ('all', 'lords'):
        sources.append('lords')

    all_rows = []
    for source in sources:
        try:
            params = {
                'key': TWFY_API_KEY, 'person': person_id,
                'type': source, 'order': 'd', 'num': 100, 'output': 'json'
            }
            query = f"{topic} {date_range}".strip() if topic else date_range.strip()
            if query:
                params['search'] = query
            resp = requests.get(TWFY_API_URL, params=params, timeout=15)
            if resp.status_code == 200:
                all_rows.extend(resp.json().get('rows', []))
        except Exception:
            pass

    unique_debates = {}
    for r in all_rows:
        parent_body = re.sub(r'<[^>]+>', '', (r.get('parent') or {}).get('body', '') or '')
        date_val = r.get('hdate', '')
        key = f"{parent_body}_{date_val}"
        url = 'https://www.theyworkforyou.com' + r.get('listurl', '')
        if key not in unique_debates:
            unique_debates[key] = {
                'title': parent_body, 'date': date_val, 'url': url,
                'type': get_debate_type(parent_body), 'contributions': 1
            }
        else:
            unique_debates[key]['contributions'] += 1

    result = list(unique_debates.values())
    result.sort(key=lambda x: x['date'], reverse=True)
    return result


# ==========================================
# ROUTE 0a: EXPAND TOPIC TERMS API (used by preview-and-edit UI)
# ==========================================
@debate_scanner_bp.route('/api/expand_topic', methods=['POST'])
def api_expand_topic():
    from flask import jsonify
    topic = (request.json or {}).get('topic', '').strip()
    if not topic:
        return jsonify({'error': 'No topic provided'}), 400
    if not GEMINI_API_KEY:
        return jsonify({'error': 'AI not configured'}), 503
    expanded = expand_search_query(topic, GEMINI_API_KEY)
    # Parse the OR terms back out for display: ("a" OR "b" OR "c") → ["a", "b", "c"]
    terms = re.findall(r'"([^"]+)"', expanded)
    return jsonify({'expanded': expanded, 'terms': terms})


# ==========================================
# ROUTE 0: MINISTERS BY DEPARTMENT API
# ==========================================
@debate_scanner_bp.route('/api/ministers_by_dept')
def api_ministers_by_dept():
    from flask import jsonify
    minister_data = get_minister_list()
    by_dept = minister_data.get('by_dept', {})
    dept = request.args.get('dept', '').strip()
    if dept:
        return jsonify(by_dept.get(dept, []))
    return jsonify(sorted(by_dept.keys()))


# ==========================================
# ROUTE 1: SCAN AND GROUP BY THEME (STEP 1)
# ==========================================
@debate_scanner_bp.route('/debates', methods=['GET', 'POST'])
def scan_debates():
    grouped_debates = {}
    error_message = None

    start_date = ""
    end_date = ""
    selected_dept = "All Departments"
    selected_house = "all"
    content_type = "exclude_bills"

    # Handle stakeholder tab GET — show search form + org directory
    if request.method == 'GET' and request.args.get('mode') == 'stakeholder':
        return render_template('debate_scanner.html',
                               mode='stakeholder',
                               stakeholder_topic=request.args.get('stakeholder_topic', ''),
                               stakeholder_org=None,
                               stakeholder_hansard=[], stakeholder_news=[],
                               stakeholder_publications=[], stakeholder_social=[],
                               stakeholder_briefing=None,
                               all_orgs=StakeholderOrg.all_active(),
                               orgs_by_category=StakeholderOrg.by_category(),
                               grouped_debates={}, error_message=None,
                               start_date='', end_date='',
                               topic='', topic_rows=[], topic_briefing=None,
                               topic_briefing_as_text='', house_filter='all',
                               selected_depts=[], wq_rows=[], oral_grouped=[],
                               urgent_grouped=[], statement_grouped=[],
                               debate_grouped=[], legislation_grouped=[],
                               debug_query='', user_pref=_get_user_pref(),
                               is_post=False,
                               departments=DEPARTMENTS_TWFY, selected_dept='All Departments',
                               selected_house='all', content_type='exclude_bills')

    if request.method == 'POST':
        action = request.form.get('action', 'search')
        
        start_date = request.form.get('start_date', '').strip()
        end_date = request.form.get('end_date', '').strip()
        selected_dept = request.form.get('department', 'All Departments').strip()
        selected_house = request.form.get('house', 'all')
        content_type = request.form.get('content_type', 'exclude_bills')

        if action == 'search':
            try:
                date_query = get_twfy_date_range(start_date, end_date)
                
                search_term = DEPT_KEYWORDS.get(selected_dept, '("Oral questions" OR "Debate")')
                if selected_house == 'all':
                    houses_to_search = ['commons', 'westminsterhall', 'lords']
                elif selected_house == 'commons':
                    houses_to_search = ['commons', 'westminsterhall']
                else:
                    houses_to_search = [selected_house]
                all_rows = []

                for h in houses_to_search:
                    query = f"{search_term} {date_query}".strip()

                    params = {
                        'key': TWFY_API_KEY,
                        'search': query,
                        'type': h,
                        'output': 'json',
                        'num': 1000 
                    }
                    
                    resp = requests.get(TWFY_API_URL, params=params, timeout=15)
                    if resp.status_code == 200:
                        all_rows.extend(resp.json().get('rows', []))

                if all_rows:
                    unique_debates = {}
                    
                    for r in all_rows:
                        raw_title = r.get('parent', {}).get('body', 'General Debate/Question')
                        title = re.sub(r'<[^>]+>', '', raw_title)
                        date_val = r.get('hdate', '')
                        
                        url = "https://www.theyworkforyou.com" + r.get('listurl', '')
                        speaker = r.get('speaker', {}).get('name', 'Unknown')
                        
                        if start_date and date_val < start_date: continue
                        if end_date and date_val > end_date: continue
                        
                        is_bill_or_amendment = bool(re.search(r'\b(bill|amendment)\b', title, re.IGNORECASE))
                        if content_type == 'exclude_bills' and is_bill_or_amendment: continue
                        if content_type == 'only_bills' and not is_bill_or_amendment: continue

                        key = f"{title}_{date_val}"
                        if key not in unique_debates:
                            unique_debates[key] = {
                                'id': key,
                                'title': title,
                                'date': date_val,
                                'minister': speaker,
                                'contributions': 1,
                                'url': url,
                                'type': get_debate_type(title)
                            }
                        else:
                            unique_debates[key]['contributions'] += 1
                            
                    debate_list = list(unique_debates.values())
                    debate_list.sort(key=lambda x: x['date'], reverse=True)
                    
                    if debate_list and GEMINI_API_KEY:
                        try:
                            MAX_AI_ITEMS = 120
                            process_list = debate_list[:MAX_AI_ITEMS]
                            leftover_list = debate_list[MAX_AI_ITEMS:]

                            titles_only = [{"id": d['id'], "title": d['title'], "minister": d['minister']} for d in process_list]
                            
                            prompt = (
                                f"You are an expert UK political analyst. The user only wants debates related to the '{selected_dept}'. "
                                "1. Group the genuinely relevant debates into 3 to 5 clear policy themes. "
                                "2. Group irrelevant/false positive debates into 'Irrelevant (Other Departments)'. "
                                "Return ONLY a pure JSON object where the keys are the 'Theme Names' and the values are arrays of the 'id' strings."
                            )
                            
                            model_path = get_working_model(GEMINI_API_KEY)
                            ai_url = f"https://generativelanguage.googleapis.com/v1beta/{model_path}:generateContent?key={GEMINI_API_KEY}"
                            payload = {
                                "contents": [{"parts": [{"text": prompt + "\n\nData: " + json.dumps(titles_only)}]}],
                                "generationConfig": {"responseMimeType": "application/json"}
                            }
                            ai_resp = requests.post(ai_url, json=payload, timeout=60)
                            
                            if ai_resp.status_code == 200:
                                ai_text = ai_resp.json().get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '{}')
                                clean_text = ai_text.replace('```json', '').replace('```', '').strip()
                                theme_mapping = json.loads(clean_text)

                                mapped_ids = set()
                                for theme, ids in theme_mapping.items():
                                    if "Irrelevant" in theme: continue
                                        
                                    theme_debates = [d for d in process_list if d['id'] in ids]
                                    if theme_debates:
                                        grouped_debates[theme] = theme_debates
                                        mapped_ids.update(ids)
                                
                                unmapped = [d for d in process_list if d['id'] not in mapped_ids]
                                if unmapped: grouped_debates["General / Uncategorized"] = unmapped
                                if leftover_list: grouped_debates[f"Older / Extra Sessions ({len(leftover_list)} Uncategorized)"] = leftover_list

                            else:
                                grouped_debates[f"All Debates (AI Error {ai_resp.status_code})"] = debate_list
                        except Exception as e:
                            grouped_debates[f"All Debates (Failsafe Triggered)"] = debate_list
                    else:
                        grouped_debates["All Debates"] = debate_list
                else:
                    error_message = "No debates found for this period/department."

            except Exception as e:
                error_message = f"Internal Error: {str(e)}"

    return render_template('debate_scanner.html',
                           mode='dept',
                           grouped_debates=grouped_debates,
                           error_message=error_message,
                           start_date=start_date, end_date=end_date,
                           departments=DEPARTMENTS_TWFY, selected_dept=selected_dept,
                           selected_house=selected_house, content_type=content_type,
                           is_post=(request.method == 'POST'),
                           # Other tab defaults
                           topic='', topic_rows=[], topic_briefing=None,
                           topic_briefing_as_text='', house_filter='all',
                           stakeholder_topic='')


# ==========================================
# ROUTE 2: TOPIC SEARCH (NEW — PARALLEL API CALLS)
# ==========================================
@debate_scanner_bp.route('/debates_topic', methods=['GET', 'POST'])
def debates_topic():
    topic_rows = []
    oral_rows = []
    statement_rows = []
    debate_rows = []
    urgent_rows = []
    oral_grouped = []
    urgent_grouped = []
    debate_grouped = []
    legislation_grouped = []
    legislation_rows = []
    statement_grouped = []
    wq_rows = []
    wq_total = 0
    topic_briefing = None
    topic_briefing_as_text = ""
    opp_speaker_links = {}
    error_message = None
    topic = ""
    narrow_keyword = ""
    start_date = ""
    end_date = ""
    house_filter = "all"
    selected_depts = []
    debug_query = ""
    user_pref = _get_user_pref()

    if request.method == 'POST':
        topic = request.form.get('topic', '').strip()
        narrow_keyword = request.form.get('narrow_keyword', '').strip()
        start_date = request.form.get('start_date', '').strip()
        end_date = request.form.get('end_date', '').strip()
        house_filter = request.form.get('house_filter', 'all')
        selected_depts = request.form.getlist('dept_filter')
        # Pre-expanded query from the preview-and-edit UI — skip Gemini expansion if provided
        preset_expanded = request.form.get('expanded_query', '').strip()

        if not topic:
            error_message = "Please enter a topic to search."
        elif not TWFY_API_KEY:
            error_message = "TWFY API key is not configured."
        else:
            date_range = get_twfy_date_range(start_date, end_date)

            if house_filter == 'lords_only':
                sources = ['lords', 'wms']
            elif house_filter == 'commons_only':
                sources = ['commons', 'westminsterhall', 'wms']
            else:
                sources = ['commons', 'westminsterhall', 'lords', 'wms']

            # Use pre-expanded terms from preview UI if provided, otherwise expand with AI
            if preset_expanded:
                expanded = preset_expanded
            else:
                expanded = expand_search_query(topic, GEMINI_API_KEY) if GEMINI_API_KEY else f'"{topic}"'
            if narrow_keyword:
                search_query = f'{expanded} AND "{narrow_keyword}"'
            else:
                search_query = expanded
            if date_range:
                clean_sq = search_query.strip()
                if clean_sq.startswith('(') and clean_sq.endswith(')'):
                    clean_sq = clean_sq[1:-1]
                debug_query = f"{clean_sq} {date_range}"
            else:
                debug_query = search_query

            # Resolve department ministers to TWFY person IDs for minister-led search.
            # When a dept is selected, we search each minister's debate history directly
            # so their sessions appear even when their speeches don't contain the keywords.
            minister_people = []
            if selected_depts:
                m_data_for_ids = get_minister_list()
                for dept in selected_depts:
                    minister_people.extend(get_dept_minister_twfy_ids(dept, m_data_for_ids))

            # Run TWFY debate search, Parliament WQ API, and minister-led searches in parallel
            all_rows = []

            def _do_wq_fetch():
                return _fetch_topic_wqs(topic, start_date, end_date, selected_depts)

            source_counts = {}
            use_hansard_minister = os.environ.get('SEARCH_BACKEND', '').lower() == 'hansard'
            with concurrent.futures.ThreadPoolExecutor(max_workers=16) as executor:
                twfy_futs = {executor.submit(copy_current_request_context(fetch_twfy_topic), search_query, src, date_range): src for src in sources}
                wq_fut = executor.submit(copy_current_request_context(_do_wq_fetch))
                # Fan-out: one future per (minister, source) so all 24 calls run in parallel
                # Lords ministers only speak in Lords/WMS; Commons ministers only in Commons/WH/WMS
                minister_futs = {}
                for mp in minister_people:
                    for src in sources:
                        if mp.get('is_lord') and src not in ('lords', 'wms'):
                            continue
                        if not mp.get('is_lord') and src == 'lords':
                            continue
                        if use_hansard_minister and mp.get('parliament_id'):
                            fut = executor.submit(
                                copy_current_request_context(fetch_hansard_minister_topic),
                                mp['parliament_id'], expanded, date_range, [src],
                                is_lord=mp.get('is_lord', False))
                        else:
                            fut = executor.submit(
                                copy_current_request_context(fetch_twfy_minister_topic),
                                mp['person_id'], expanded, date_range, [src],
                                is_lord=mp.get('is_lord', False))
                        minister_futs[fut] = mp
                all_futs = list(twfy_futs.keys()) + [wq_fut] + list(minister_futs.keys())
                for future in concurrent.futures.as_completed(all_futs):
                    if future is wq_fut:
                        try:
                            wq_rows, wq_total = future.result()
                        except Exception:
                            pass
                    elif future in twfy_futs:
                        src = twfy_futs[future]
                        try:
                            rows = future.result()
                            errors = [r for r in rows if r.get('_error')]
                            good = [r for r in rows if not r.get('_error')]
                            source_counts[src] = f"{len(good)} results" + (f" | ERR: {errors[0]['_error']}" if errors else "")
                            all_rows.extend(rows)
                        except Exception as e:
                            source_counts[src] = f"exception: {e}"
                    else:
                        try:
                            all_rows.extend(future.result())
                        except Exception:
                            pass

            # Extract any TWFY error markers before dedup
            twfy_errors = [r['_error'] for r in all_rows if r.get('_error')]
            all_rows = [r for r in all_rows if not r.get('_error')]
            source_debug = ' | '.join(f"{s}: {c}" for s, c in sorted(source_counts.items()))
            if source_debug:
                debug_query += f" | [{source_debug}]"
            if twfy_errors:
                debug_query += ' | ERRORS: ' + '; '.join(set(twfy_errors))

            topic_rows = deduplicate_by_listurl(all_rows)
            import logging as _tlog
            twfy_count = sum(1 for r in topic_rows if not r.get('listurl', '').startswith('http'))
            hansard_count = sum(1 for r in topic_rows if r.get('listurl', '').startswith('http'))
            _tlog.warning(f"[topic_rows] {len(topic_rows)} total after dedup: {twfy_count} TWFY, {hansard_count} Hansard | source_counts={source_counts}")

            # Apply date filter BEFORE session expansion so we only expand in-range sessions.
            # TWFY keyword search returns historical results by relevance even with a date
            # range in the search string — the Python filter is the reliable enforcement.
            if start_date or end_date:
                import logging as _flog
                before_filter = len(topic_rows)
                topic_rows = [r for r in topic_rows
                              if (not start_date or r.get('hdate', '') >= start_date)
                              and (not end_date or r.get('hdate', '') <= end_date)]
                _flog.warning(f"[date_filter] start={start_date!r} end={end_date!r} before={before_filter} after={len(topic_rows)}")

            # Debates-first: expand each matched debate to include ALL speeches.
            # This ensures ministerial responses appear even when their speeches don't
            # contain the search keywords (minister: "supporting graduates", not "loan repayments").
            if topic_rows:
                session_speeches = fetch_all_debate_sessions(topic_rows, max_debates=25)
                if session_speeches:
                    topic_rows = deduplicate_by_listurl(topic_rows + session_speeches)
            # Re-apply date filter — session expansion fetches all speeches in a session,
            # which can occasionally include adjacent-day edge cases.
            if start_date or end_date:
                topic_rows = [r for r in topic_rows
                              if (not start_date or r.get('hdate', '') >= start_date)
                              and (not end_date or r.get('hdate', '') <= end_date)]

            # Flag ministerial speakers and sort them to the top
            minister_data = get_minister_list()
            by_norm = minister_data.get('by_norm', {})
            for row in topic_rows:
                spk = row.get('speaker_name', '')
                norm_spk = _normalise_name(spk)
                role = by_norm.get(norm_spk) if norm_spk else None
                row['is_minister'] = bool(role)
                row['minister_role'] = role or ''

            if not topic_rows:
                error_message = f"No parliamentary contributions found for '{topic}'. Try a broader search term or wider date range."
            elif GEMINI_API_KEY:
                try:
                    # Build a balanced AI payload:
                    # - Ministers first (capped at 15, highest relevance)
                    # - Formal opposition parties next — ordered first within non-minister rows
                    #   so they're not crowded out by Labour backbenchers or empty-party rows
                    # - All other non-government voices: Labour backbenchers, Crossbench peers,
                    #   Independents, unknown party — included via _GOVERNMENT_PARTIES exclusion
                    #   (not in _OPPOSITION_PARTIES is intentionally broader)
                    _GOVERNMENT_PARTIES = {'Labour'}
                    _OPPOSITION_PARTIES = {
                        'Conservative', 'Liberal Democrat', 'Scottish National Party',
                        'Reform UK', 'Plaid Cymru', 'Green Party', 'Democratic Unionist Party',
                        'Ulster Unionist Party', 'Social Democratic and Labour Party',
                        'Sinn Féin', 'Alliance Party of Northern Ireland', 'Alba Party',
                    }
                    import logging as _dlog2
                    _dlog2.warning(f"[briefing_diag] topic_rows_at_briefing={len(topic_rows)} minister_flags={sum(1 for r in topic_rows if r.get('is_minister'))}")
                    minister_rows = sorted(
                        [r for r in topic_rows if r.get('is_minister')],
                        key=lambda x: -x.get('relevance', 0)
                    )[:15]
                    non_minister_rows = [r for r in topic_rows if not r.get('is_minister')]
                    # opp_rows: formal opposition parties — prioritised in payload ordering
                    opp_rows = sorted(
                        [r for r in non_minister_rows if r.get('speaker_party', '') in _OPPOSITION_PARTIES],
                        key=lambda x: -x.get('relevance', 0)
                    )
                    # other_rows: Labour backbenchers, Crossbench peers, Independents, unknown
                    other_rows = sorted(
                        [r for r in non_minister_rows if r.get('speaker_party', '') not in _OPPOSITION_PARTIES],
                        key=lambda x: -x.get('relevance', 0)
                    )
                    if not opp_rows and non_minister_rows:
                        import logging as _log
                        _log.warning(
                            f"[briefing] opp_rows empty despite {len(non_minister_rows)} non-minister rows"
                            f" — sample parties: {list({r.get('speaker_party','') for r in non_minister_rows[:10]})}"
                        )
                    balanced = minister_rows + opp_rows[:10] + other_rows[:10]
                    # Build name → URL lookup for opposition/backbench speakers so the
                    # template can render a "View Speech" link alongside "View PQs".
                    opp_speaker_links = {}
                    for r in opp_rows + other_rows:
                        name = r.get('speaker_name', '')
                        url = r.get('listurl', '')
                        if name and url and name not in opp_speaker_links:
                            # TWFY listurls are relative — prepend domain
                            if url and not url.startswith('http'):
                                url = 'https://www.theyworkforyou.com' + url
                            opp_speaker_links[name] = url
                    import logging as _dlog
                    _dlog.warning(
                        f"[briefing_diag] non_minister={len(non_minister_rows)} "
                        f"opp={len(opp_rows)} other={len(other_rows)} balanced={len(balanced)}"
                    )
                    _dlog.warning(
                        f"[briefing_diag] opp sample: {[(r.get('speaker_name',''),r.get('speaker_party','')) for r in opp_rows[:5]]}"
                    )
                    _dlog.warning(
                        f"[briefing_diag] other sample: {[(r.get('speaker_name',''),r.get('speaker_party','')) for r in other_rows[:5]]}"
                    )
                    _dlog.warning(
                        f"[briefing_diag] payload parties: {[r.get('speaker_party','') for r in balanced]}"
                    )
                    ai_payload = [
                        {'listurl': r['listurl'], 'speaker': r['speaker_name'],
                         'party': r['speaker_party'], 'date': r['hdate'],
                         'source': r['source_label'], 'text': r['body_clean'][:250],
                         'is_minister': r.get('is_minister', False)}
                        for r in balanced
                    ]
                    dept_context = (
                        f" The user is specifically interested in contributions related to the "
                        f"{', '.join(selected_depts)} {'portfolio' if len(selected_depts) == 1 else 'portfolios'}."
                        if selected_depts else ""
                    )
                    prompt = (
                        f"You are a senior UK civil servant writing a parliamentary briefing on: \"{topic}\"."
                        f"{dept_context}\n\n"
                        "Return ONLY a valid JSON object (no markdown fences) with these exact keys:\n\n"
                        "\"topic_summary\": 2–3 sentence overview of the parliamentary debate on this topic, "
                        "covering both government and opposition positions at a high level.\n\n"
                        "\"government_position\": Summarise the government's stated position on this topic "
                        "as expressed in Parliament — what ministers have said, any commitments made, "
                        "and the policy direction."
                        + (f" Focus specifically on {', '.join(selected_depts)} ministers." if selected_depts else "")
                        + "\n\n"
                        "\"opposition_position\": Summarise the positions of ALL opposition parties "
                        "(Conservative, Liberal Democrat, SNP, Reform UK, Plaid Cymru, and any others). "
                        "Give each party's distinct stance where they differ. "
                        "Do NOT include Labour backbenchers here — Labour is the governing party.\n\n"
                        "\"government_speakers\": Array of {\"name\", \"role\", \"stance\"} — "
                        "only ministers, Secretaries of State, and PPSs who spoke on this topic"
                        + (f" in the context of {', '.join(selected_depts)}" if selected_depts else "")
                        + ". Include Lords ministers (e.g. Baroness X). Up to 8 entries.\n\n"
                        "\"non_government_speakers\": Array of {\"name\", \"role_or_party\", \"stance\"} — "
                        "Include ALL non-government voices present in the DATA: opposition MPs asking "
                        "questions in Oral Questions sessions, shadow ministers, crossbench peers, "
                        "and any other non-Labour/non-government speaker. "
                        "Order: shadow frontbenchers / opposition spokespeople first (most senior first), "
                        "then opposition backbenchers, then crossbenchers. "
                        "IMPORTANT: In Oral Questions, Conservative or other opposition MPs asking "
                        "supplementary questions ARE non-government speakers — always include them. "
                        "Return empty array ONLY if the DATA genuinely contains zero non-government speakers. "
                        "Up to 10 entries.\n\n"
                        "\"key_questions\": Array of {\"speaker\", \"role_or_party\", \"date\", \"source\", \"listurl\", \"question\"} — "
                        "The most significant questions raised by OPPOSITION and non-government speakers about this topic. "
                        "These should be actual interrogative challenges, probing questions, or lines of attack that a minister "
                        "would need to answer — not statements or positions. "
                        "Focus on: questions about policy gaps, accountability questions, questions challenging government figures or commitments, "
                        "questions that reveal opposition priorities or likely future lines of attack. "
                        "ONLY include questions from opposition frontbenchers, shadow ministers, opposition backbenchers, crossbenchers, "
                        "or Lords not in government. Do NOT include questions asked by government ministers or PPSs. "
                        "\"question\" should be a concise paraphrase of the actual question asked (1-2 sentences). "
                        "Use the exact listurl from the matching DATA entry. Up to 6 questions.\n\n"
                        "\"anticipated_questions\": Array of {\"question\", \"rationale\"} — "
                        "Up to 5 HARDER questions that have NOT yet been asked in Parliament but could plausibly come up "
                        "based on the debates, gaps in government answers, and opposition lines of attack identified above. "
                        "These should be the uncomfortable questions — ones that probe weaknesses in the government position, "
                        "unresolved policy tensions, or areas where the evidence base is thin. "
                        "\"rationale\" should be a brief (1 sentence) explanation of why this question is likely to arise. "
                        "These are for ministerial preparation — make them specific and genuinely challenging.\n\n"
                        "\"next_steps\": Any upcoming parliamentary business or announced policy milestones.\n\n"
                        "\"coverage_note\": Brief note on the date range and sources covered.\n\n"
                        f"DATA: {json.dumps(ai_payload)}"
                    )
                    model_path = get_working_model(GEMINI_API_KEY)
                    ai_url = f"https://generativelanguage.googleapis.com/v1beta/{model_path}:generateContent?key={GEMINI_API_KEY}"
                    payload = {
                        "contents": [{"parts": [{"text": prompt}]}],
                        "generationConfig": {"responseMimeType": "application/json"}
                    }
                    ai_resp = requests.post(ai_url, json=payload, timeout=90)
                    if ai_resp.status_code == 503:
                        time.sleep(3)
                        ai_resp = requests.post(ai_url, json=payload, timeout=90)
                    if ai_resp.status_code == 200:
                        raw_text = ai_resp.json().get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '{}')
                    else:
                        raw_text = _claude_fallback(prompt, max_tokens=3000)
                    _dlog.warning(f"[briefing_raw] non_govt_speakers snippet: {raw_text[raw_text.find('non_government'):raw_text.find('non_government')+300] if 'non_government' in (raw_text or '') else 'KEY NOT FOUND'}")
                    if raw_text:
                        topic_briefing = _parse_ai_json(raw_text)

                        # Verify government speakers against Parliament Members API (parallel)
                        govt_speakers = topic_briefing.get('government_speakers', [])
                        if govt_speakers:
                            with concurrent.futures.ThreadPoolExecutor(max_workers=5) as vex:
                                vfutures = {vex.submit(copy_current_request_context(verify_government_speaker), s.get('name', '')): s for s in govt_speakers}
                                for vf in concurrent.futures.as_completed(vfutures):
                                    spk = vfutures[vf]
                                    try:
                                        v = vf.result()
                                        spk['verified'] = v['confirmed']
                                        spk['confirmed_role'] = v['role']
                                    except Exception:
                                        spk['verified'] = False
                                        spk['confirmed_role'] = ''

                        topic_briefing_as_text = format_briefing_as_text(topic_briefing, topic)
                    else:
                        debug_query += f" | AI HTTP {ai_resp.status_code} (Gemini+Claude both failed)"
                        topic_briefing = None
                except Exception as e:
                    import logging
                    logging.error(f"AI briefing failed: {type(e).__name__}: {e}")
                    debug_query += f" | AI error: {type(e).__name__}: {str(e)[:120]}"
                    topic_briefing = None

            # Split TWFY rows into display sections, then group debates by session
            # Group ALL non-statement rows first, then classify each group.
            # Group-level classification is more reliable than per-row:
            # a group where no speech exceeds 300 words is an Oral PQ session
            # (prepared answer ~150w, follow-ups shorter) not a debate.
            # Genuine WMS: source=wms AND title doesn't look like a written answer
            statement_rows = [r for r in topic_rows
                              if r.get('source') == 'wms'
                              and 'written answer' not in r.get('debate_title', '').lower()]
            # wrans + any row whose title contains "written answer" (any source) → WQ section
            def _is_written_answer(r):
                return (r.get('source') == 'wrans'
                        or 'written answer' in r.get('debate_title', '').lower())

            # TWFY wrans records hold the minister's ANSWER text, not the question.
            # speaker_name is the answering minister, body_clean is the answer body.
            # Displaying them as WQ cards would show the minister as the "asker" and
            # the answer as the "question" — misleading. The Parliament API WQ results
            # (from _fetch_topic_wqs) are richer and accurate, so TWFY wrans are skipped.
            # wq_rows stays as-is (Parliament API only).
            wq_total = len(wq_rows)
            non_statement_rows = [r for r in topic_rows
                                  if r not in statement_rows
                                  and not _is_written_answer(r)]
            all_grouped = _group_by_debate(non_statement_rows)

            oral_grouped, urgent_grouped, debate_grouped, legislation_grouped = [], [], [], []
            for grp in all_grouped:
                section = _classify_group(grp)
                if section == 'oral':
                    oral_grouped.append(grp)
                elif section == 'urgent':
                    urgent_grouped.append(grp)
                elif section == 'legislation':
                    legislation_grouped.append(grp)
                elif section == 'wq':
                    pass  # written answer slipped through upstream filter — discard
                else:
                    debate_grouped.append(grp)

            # Group WMS by session so they render like debates, not flat speech snippets
            statement_grouped = _group_by_debate(statement_rows)

            # Filter WMS and Oral Questions to selected departments only.
            # WMS titles: "Written Ministerial Statements — [Dept]: [Topic]"
            # Oral titles: "Oral Answers to Questions — [Dept]" or "[Dept] Questions"
            if selected_depts:
                def _dept_in_title(title, depts):
                    title_lower = title.lower()
                    for d in depts:
                        # Full name match (e.g. "Department for Education")
                        if d.lower() in title_lower:
                            return True
                        # Key words only — strip generic words to get e.g. "Education", "Treasury"
                        keywords = [w for w in d.split()
                                    if w.lower() not in ('department', 'for', 'of', 'the',
                                                         'hm', 'ministry', 'and', 'office')
                                    and len(w) > 3]
                        if any(kw.lower() in title_lower for kw in keywords):
                            return True
                    return False
                statement_grouped = [g for g in statement_grouped
                                     if _dept_in_title(g['title'], selected_depts)]
                # Oral questions: only filter sessions that explicitly name a
                # DIFFERENT department (e.g. "Oral Answers to Questions — Treasury").
                # Topic-specific oral questions ("Plan 2 Student Loans: Repayment
                # Terms — Question") have no dept in the title and must not be removed.
                def _oral_from_other_dept(grp_title, depts):
                    t = grp_title.lower()
                    if 'oral answers to questions' not in t:
                        return False  # not a dept oral questions session → keep
                    return not _dept_in_title(grp_title, depts)
                oral_grouped = [g for g in oral_grouped
                                if not _oral_from_other_dept(g['title'], selected_depts)]

            # Sort WQs newest-first
            wq_rows.sort(key=lambda q: q.get('date_tabled', ''), reverse=True)

            # Flat row lists for JS download variables
            oral_rows = [r for grp in oral_grouped for r in grp['speeches']]
            urgent_rows = [r for grp in urgent_grouped for r in grp['speeches']]
            debate_rows = [r for grp in debate_grouped for r in grp['speeches']]
            legislation_rows = [r for grp in legislation_grouped for r in grp['speeches']]

    return render_template('debate_scanner.html',
                           mode='topic',
                           topic=topic, topic_rows=topic_rows,
                           oral_rows=oral_rows, statement_rows=statement_rows,
                           urgent_rows=urgent_rows,
                           debate_rows=debate_rows,
                           oral_grouped=oral_grouped, debate_grouped=debate_grouped,
                           statement_grouped=statement_grouped,
                           urgent_grouped=urgent_grouped,
                           legislation_grouped=legislation_grouped,
                           legislation_rows=legislation_rows,
                           wq_rows=wq_rows, wq_total=wq_total,
                           topic_briefing=topic_briefing,
                           topic_briefing_as_text=topic_briefing_as_text,
                           opp_speaker_links=opp_speaker_links,
                           start_date=start_date, end_date=end_date,
                           house_filter=house_filter,
                           narrow_keyword=narrow_keyword,
                           selected_depts=selected_depts,
                           debug_query=debug_query,
                           user_pref=user_pref,
                           error_message=error_message,
                           grouped_debates={}, departments=DEPARTMENTS_TWFY,
                           selected_dept="All Departments", selected_house="all",
                           content_type="exclude_bills", is_post=True,
                           stakeholder_topic='')


# ==========================================
# ROUTE 3: FETCH TRANSCRIPTS AND ANALYSE (DEPT SCAN)
# ==========================================
@debate_scanner_bp.route('/debates_analyze', methods=['POST'])
def analyze_selected():
    selected = request.form.getlist('selected_debates')
    if not selected:
        return "Please go back and select at least one debate.", 400

    full_transcript_text = ""
    
    try:
        for item in selected:
            parts = item.split('||')
            url = parts[0]
            title = parts[1] if len(parts) > 1 else "Unknown Title"
            date_val = parts[2] if len(parts) > 2 else ""

            if url and url.startswith('http'):
                # Check cache first — transcripts never change once published
                cached = CachedTranscript.get(url)
                if cached:
                    full_transcript_text += f"### DEBATE: {cached.title} ({cached.date})\n\n{cached.transcript_text}\n\n"
                    continue

                headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
                resp = requests.get(url, headers=headers, timeout=20)

                if resp.status_code == 200:
                    html = resp.text
                    html = re.sub(r'<head.*?>.*?</head>', '', html, flags=re.DOTALL|re.IGNORECASE)
                    html = re.sub(r'<script.*?>.*?</script>', '', html, flags=re.DOTALL|re.IGNORECASE)
                    html = re.sub(r'<style.*?>.*?</style>', '', html, flags=re.DOTALL|re.IGNORECASE)
                    html = re.sub(r'<nav.*?>.*?</nav>', '', html, flags=re.DOTALL|re.IGNORECASE)
                    html = re.sub(r'<header.*?>.*?</header>', '', html, flags=re.DOTALL|re.IGNORECASE)
                    html = re.sub(r'<footer.*?>.*?</footer>', '', html, flags=re.DOTALL|re.IGNORECASE)

                    clean_text = re.sub(r'<[^>]+>', ' ', html)
                    clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                    clean_text = clean_text[:60000]

                    # Store in cache for next time
                    try:
                        CachedTranscript.store(url=url, title=title, date=date_val,
                                               house='', transcript_text=clean_text)
                    except Exception:
                        pass

                    session_text = f"### DEBATE: {title} ({date_val})\n\n{clean_text}\n\n"
                    full_transcript_text += session_text
            
        if not full_transcript_text:
            return "No content could be extracted from the URL. Please try another session.", 400

        if GEMINI_API_KEY:
            model_path = get_working_model(GEMINI_API_KEY)
            # THE FIX: Added Section 6 for Active Stakeholders
            prompt = (
                "You are a senior UK Parliamentary Researcher. Analyze the following Hansard transcript "
                "and provide a high-level parliamentary briefing.\n\n"
                "Ignore any remaining website navigation text (like 'Home', 'Members', etc.) and focus only on the debate content.\n"
                "Structure your response exactly as follows (use Markdown formatting):\n"
                "## 1. EXECUTIVE SUMMARY\n(A 3-sentence overview of the main policy direction or debate focus)\n\n"
                "## 2. MINISTERIAL COMMITMENTS\n(Bullet points of any concrete promises, funding, or deadlines mentioned by the government)\n\n"
                "## 3. OPPOSITION CRITIQUE\n(Key challenges raised by Shadow Ministers or other parties)\n\n"
                "## 4. MOOD OF THE HOUSE\n(Is the sentiment generally supportive, hostile, or concerned?)\n\n"
                "## 5. KEY QUESTIONS & RESPONSES\n(Extract 3 to 5 of the most pressing questions asked by Members, paired with the Minister's direct response.)\n\n"
                "## 6. ACTIVE STAKEHOLDERS (NON-MINISTERIAL)\n(Provide a bulleted list of the backbench MPs, Opposition leads, or Peers who spoke in these specific sessions. Exclude the responding Government Minister. Briefly summarize their stance or the specific angle they focused on.)\n\n"
                f"TRANSCRIPTS:\n{full_transcript_text}"
            )
            
            ai_url = f"https://generativelanguage.googleapis.com/v1beta/{model_path}:generateContent?key={GEMINI_API_KEY}"
            ai_resp = requests.post(ai_url, json={"contents": [{"parts": [{"text": prompt}]}]}, timeout=60)
            
            if ai_resp.status_code == 200:
                briefing_content = ai_resp.json().get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
            else:
                briefing_content = f"AI Analysis failed. (Error {ai_resp.status_code})"

            return render_template('debate_briefing.html', briefing=briefing_content, selected_count=len(selected))

    except Exception as e:
        return f"Error during analysis: {str(e)}", 500


# ==========================================
# ROUTE 4: MINISTERIAL RECORD — DEBATE LIST
# ==========================================
@debate_scanner_bp.route('/debates_minister', methods=['GET', 'POST'])
def debates_minister():
    minister_debates = []
    minister_info = None
    error_message = None
    minister_name = ""
    topic = ""
    start_date = ""
    end_date = ""
    house_filter = "all"

    if request.method == 'POST':
        minister_name = request.form.get('minister_name', '').strip()
        topic = request.form.get('topic', '').strip()
        start_date = request.form.get('start_date', '').strip()
        end_date = request.form.get('end_date', '').strip()
        house_filter = request.form.get('house_filter', 'all')

        if not minister_name:
            error_message = "Please enter a minister's name."
        elif not TWFY_API_KEY:
            error_message = "TWFY API key is not configured."
        else:
            person_id, matched_name, is_lord = lookup_twfy_person(minister_name)
            if not person_id:
                error_message = f"Could not find '{minister_name}' on TheyWorkForYou. Try their full name."
            else:
                minister_info = {'name': matched_name, 'person_id': person_id, 'is_lord': is_lord}
                date_range = get_twfy_date_range(start_date, end_date)
                minister_debates = fetch_minister_debates(person_id, topic, date_range, house_filter)
                if not minister_debates:
                    error_message = f"No debates found for {matched_name}. Try a broader topic or date range."

    return render_template('debate_scanner.html',
                           mode='minister',
                           minister_name=minister_name, minister_info=minister_info,
                           minister_debates=minister_debates,
                           topic=topic, start_date=start_date, end_date=end_date,
                           house_filter=house_filter, error_message=error_message,
                           grouped_debates={}, departments=DEPARTMENTS_TWFY,
                           selected_dept="All Departments", selected_house="all",
                           content_type="exclude_bills", is_post=(request.method == 'POST'),
                           topic_rows=[], topic_briefing=None, topic_briefing_as_text='',
                           stakeholder_topic='')


# ==========================================
# ROUTE 5: MINISTERIAL RECORD — ANALYSE SELECTED
# ==========================================
@debate_scanner_bp.route('/debates_minister_analyze', methods=['POST'])
def debates_minister_analyze():
    selected = request.form.getlist('selected_debates')
    minister_name = request.form.get('minister_name', 'the Minister').strip()

    if not selected:
        return "Please select at least one debate.", 400

    full_transcript_text = ""
    for item in selected:
        parts = item.split('||')
        url = parts[0]
        title = parts[1] if len(parts) > 1 else "Unknown"
        date_val = parts[2] if len(parts) > 2 else ""

        cached = CachedTranscript.get(url)
        if cached:
            full_transcript_text += f"### {cached.title} ({cached.date})\n\n{cached.transcript_text}\n\n"
            continue

        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            resp = requests.get(url, headers=headers, timeout=20)
            if resp.status_code == 200:
                html = resp.text
                for tag in ['head', 'script', 'style', 'nav', 'header', 'footer']:
                    html = re.sub(rf'<{tag}.*?>.*?</{tag}>', '', html, flags=re.DOTALL | re.IGNORECASE)
                clean_text = re.sub(r'<[^>]+>', ' ', html)
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()[:60000]
                try:
                    CachedTranscript.store(url=url, title=title, date=date_val, house='', transcript_text=clean_text)
                except Exception:
                    pass
                full_transcript_text += f"### {title} ({date_val})\n\n{clean_text}\n\n"
        except Exception:
            pass

    if not full_transcript_text:
        return "Could not extract content from the selected debates.", 400

    if not GEMINI_API_KEY:
        return "Gemini API key is not configured.", 500

    prompt = (
        f"You are a senior UK Parliamentary Researcher. Analyze the following Hansard transcripts "
        f"focusing specifically on contributions by {minister_name}.\n\n"
        "Ignore any website navigation text. Focus only on debate content.\n\n"
        "IMPORTANT: In UK oral questions, the minister gives an OPENING ANSWER to the lead question "
        "before supplementary questions follow. Treat this opening answer as a speech — include it "
        "in full in Section 1 under the first exchange, as it is the minister's prepared position.\n\n"
        "Structure your response in Markdown exactly as follows:\n\n"
        f"## MINISTERIAL RECORD: {minister_name.upper()}\n\n"
        "## 1. SPEECHES & OPENING STATEMENTS\n"
        f"Include the full text of any opening statements, prepared answers, or standalone speeches "
        f"by {minister_name}. In oral questions sessions, this is the minister's first substantive "
        "answer to each topic before supplementaries begin. Quote the full text, not a summary.\n\n"
        "## 2. SUPPLEMENTARY Q&A EXCHANGES\n"
        "After the opening statement, list each follow-up (supplementary) exchange:\n"
        "**Q [{questioner name} ({party})]:** The question\n"
        f"**A [{minister_name}]:** The minister's response\n\n"
        "## 3. COMMUNICATION STYLE ANALYSIS\n"
        "Tone and register, how the minister handles difficult or hostile questions, "
        "recurring phrases or framing, approach to statistics and evidence.\n\n"
        "## 4. KEY POLICY POSITIONS STATED\n"
        f"Bullet points of the specific positions and commitments {minister_name} has expressed.\n\n"
        f"TRANSCRIPTS:\n{full_transcript_text}"
    )

    model_path = get_working_model(GEMINI_API_KEY)
    ai_url = f"https://generativelanguage.googleapis.com/v1beta/{model_path}:generateContent?key={GEMINI_API_KEY}"
    ai_resp = requests.post(ai_url, json={"contents": [{"parts": [{"text": prompt}]}]}, timeout=120)

    if ai_resp.status_code == 200:
        briefing_content = ai_resp.json().get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '')
    else:
        briefing_content = f"AI Analysis failed (Error {ai_resp.status_code})."

    return render_template('debate_briefing.html', briefing=briefing_content, selected_count=len(selected))


# ==========================================
# ROUTE 6: EXPORT BRIEFING TO WORD DOC
# ==========================================
@debate_scanner_bp.route('/download_debate_briefing', methods=['POST'])
def download_debate_briefing():
    if not Document:
        return "Word library missing.", 500

    briefing_text = request.form.get('briefing_text', '')
    if not briefing_text: return "No text found.", 400

    doc = Document()
    doc.add_heading('Parliamentary Intelligence Briefing', 0)
    
    for line in briefing_text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith('##'):
            doc.add_heading(line.replace('#', '').strip(), level=2)
        elif line.startswith('*') or line.startswith('-'):
            doc.add_paragraph(line.lstrip('*- ').strip(), style='List Bullet')
        else:
            doc.add_paragraph(line.replace('**', ''))

    mem_doc = io.BytesIO()
    doc.save(mem_doc)
    mem_doc.seek(0)

    return send_file(mem_doc, as_attachment=True, download_name=f"Briefing_{datetime.now().strftime('%Y%m%d')}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# ==========================================
# ROUTE 7: CUSTOM EXPORT — PICK & CHOOSE WORD DOC
# ==========================================
@debate_scanner_bp.route('/export_custom_briefing', methods=['POST'])
def export_custom_briefing():
    if not Document:
        return "Word library missing.", 500

    topic = request.form.get('topic', 'Parliamentary Research').strip()
    include_briefing = request.form.get('include_briefing', '') == 'true'
    briefing_text = request.form.get('briefing_text', '')
    items_json = request.form.get('items', '[]')
    try:
        items = json.loads(items_json)
    except Exception:
        items = []

    contributions = [i for i in items if i.get('type') in ('contribution', 'speech')]
    pq_items = [i for i in items if i.get('type') == 'pq']

    doc = Document()

    # ── Title block ─────────────────────────────────────────────
    h = doc.add_heading('Parliamentary Research Brief', 0)
    h.alignment = 1  # centre
    sub = doc.add_paragraph(f'Topic: {topic}')
    sub.alignment = 1
    sub.runs[0].bold = True
    date_p = doc.add_paragraph(f'Generated: {datetime.now().strftime("%d %B %Y")}')
    date_p.alignment = 1
    doc.add_paragraph()

    # ── AI Briefing ─────────────────────────────────────────────
    if include_briefing and briefing_text:
        doc.add_heading('AI-Generated Briefing', 1)
        for line in briefing_text.split('\n'):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue
            if line.startswith('## '):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith('# '):
                doc.add_heading(line[2:].strip(), level=2)
            elif line.startswith('* ') or line.startswith('- '):
                doc.add_paragraph(line[2:].strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.replace('**', ''))
        doc.add_paragraph()

    # ── Parliamentary Contributions ──────────────────────────────
    if contributions:
        doc.add_heading(f'Selected Parliamentary Contributions ({len(contributions)})', 1)
        for item in contributions:
            d = item.get('data', {})
            # Header row
            p = doc.add_paragraph()
            p.add_run(f"{d.get('hdate', '')}").bold = True
            p.add_run(f"  ·  {d.get('source_label', '')}  ·  {d.get('debate_type', '')}")
            # Speaker
            p2 = doc.add_paragraph()
            p2.add_run(f"{d.get('speaker_name', '')}").bold = True
            party = d.get('speaker_party', '')
            if party:
                p2.add_run(f"  ({party})")
            if d.get('is_minister'):
                p2.add_run("  ✓ Minister").font.color.rgb = None  # keep default
            # Debate title
            title = d.get('debate_title', '')
            if title:
                tp = doc.add_paragraph(title)
                tp.runs[0].italic = True
            # Body text
            body = d.get('body_clean', '')
            if body:
                doc.add_paragraph(f'"{body[:600]}{"…" if len(body) > 600 else ""}"')
            # Link
            url = d.get('listurl', '')
            if url:
                if not url.startswith('http'):
                    url = 'https://www.theyworkforyou.com' + url
                lp = doc.add_paragraph()
                lp.add_run('Source: ').bold = True
                _add_hyperlink(lp, url, url)
            doc.add_paragraph('─' * 60)

    # ── Written Questions ────────────────────────────────────────
    if pq_items:
        # Group by speaker
        by_speaker = {}
        for item in pq_items:
            spk = item.get('speaker', 'Unknown')
            by_speaker.setdefault(spk, []).append(item.get('data', {}))

        doc.add_heading(f'Written Questions ({len(pq_items)} questions across {len(by_speaker)} member(s))', 1)
        for speaker, pqs in by_speaker.items():
            doc.add_heading(f'{speaker}  ({len(pqs)} questions)', 2)
            for pq in pqs:
                # Status badge
                status = pq.get('status', '')
                dept = pq.get('dept', '')
                date = pq.get('date', '')
                p = doc.add_paragraph()
                p.add_run(f"{date}  ·  {dept}  ·  ").bold = False
                sr = p.add_run(status)
                sr.bold = True
                # Question text
                doc.add_paragraph(f'"{pq.get("text", "")}"')
                # Link
                link = pq.get('link', '')
                if link:
                    lp = doc.add_paragraph()
                    lp.add_run('Parliament link: ').bold = True
                    _add_hyperlink(lp, link, link)
                doc.add_paragraph()

    mem_doc = io.BytesIO()
    doc.save(mem_doc)
    mem_doc.seek(0)

    safe_topic = re.sub(r'[^\w\s-]', '', topic)[:40].strip()
    filename = f"Brief - {safe_topic} - {datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(mem_doc, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# ==========================================
# STAKEHOLDER RESEARCH — HELPERS
# ==========================================

def fetch_org_rss(rss_url, topic, limit=15):
    """Fetch RSS feed and return items loosely matching topic. Returns [] on failure."""
    if not feedparser or not rss_url:
        return []
    try:
        feed = feedparser.parse(rss_url)
        topic_words = set(re.findall(r'\b\w{4,}\b', topic.lower()))
        results = []
        for entry in feed.entries[:40]:
            title = entry.get('title', '')
            summary = entry.get('summary', '') or entry.get('description', '')
            text = (title + ' ' + summary).lower()
            # Include if any topic word appears in the entry
            if not topic_words or any(w in text for w in topic_words):
                published = ''
                if hasattr(entry, 'published_parsed') and entry.published_parsed:
                    try:
                        published = datetime(*entry.published_parsed[:3]).strftime('%Y-%m-%d')
                    except Exception:
                        pass
                results.append({
                    'title': title,
                    'link': entry.get('link', ''),
                    'published': published,
                    'summary': re.sub(r'<[^>]+>', ' ', summary)[:300].strip(),
                    'source_type': 'publication',
                })
            if len(results) >= limit:
                break
        return results
    except Exception:
        return []


def fetch_org_bluesky(handle, topic, limit=10):
    """Fetch recent Bluesky posts from a handle and filter loosely by topic."""
    try:
        from atproto import Client as BskyClient
        BSKY_HANDLE = os.environ.get('BSKY_HANDLE')
        BSKY_PASSWORD = os.environ.get('BSKY_PASSWORD')
        if not BSKY_HANDLE or not BSKY_PASSWORD:
            return []
        bsky = BskyClient()
        bsky.login(BSKY_HANDLE, BSKY_PASSWORD)
        clean_handle = handle.lstrip('@')
        feed = bsky.get_author_feed(actor=clean_handle, limit=25)
        topic_words = set(re.findall(r'\b\w{4,}\b', topic.lower()))
        results = []
        for item in feed.feed:
            post = item.post
            text = post.record.text if hasattr(post.record, 'text') else ''
            if not topic_words or any(w in text.lower() for w in topic_words):
                post_id = post.uri.split('/')[-1]
                results.append({
                    'title': text[:120],
                    'link': f'https://bsky.app/profile/{clean_handle}/post/{post_id}',
                    'published': str(post.indexed_at)[:10] if hasattr(post, 'indexed_at') else '',
                    'summary': text[:300],
                    'source_type': 'social',
                })
            if len(results) >= limit:
                break
        return results
    except Exception:
        return []


def fetch_stakeholder_news(org_names, topic, start_date='', end_date='', limit=30):
    """Search News API for media coverage mentioning stakeholder orgs on a topic."""
    NEWS_API_KEY = os.environ.get('NEWS_API_KEY')
    if not NEWS_API_KEY:
        return []
    try:
        from newsapi import NewsApiClient
        newsapi = NewsApiClient(api_key=NEWS_API_KEY)
        # Build query: topic AND (org1 OR org2 OR ...)
        org_clause = ' OR '.join(f'"{n}"' for n in org_names[:5])  # API has query length limits
        query = f'({topic}) AND ({org_clause})'
        kwargs = {'q': query, 'language': 'en', 'sort_by': 'relevancy', 'page_size': limit}
        if start_date:
            kwargs['from_param'] = start_date
        if end_date:
            kwargs['to'] = end_date
        resp = newsapi.get_everything(**kwargs)
        results = []
        for art in resp.get('articles', []):
            results.append({
                'title': art.get('title', ''),
                'link': art.get('url', ''),
                'published': (art.get('publishedAt') or '')[:10],
                'summary': art.get('description') or art.get('title', ''),
                'outlet': art.get('source', {}).get('name', ''),
                'source_type': 'news',
            })
        return results
    except Exception:
        return []


def generate_stakeholder_briefing(topic, org=None, hansard_rows=None, news=None,
                                  publications=None, social=None):
    """Generate an AI summary of stakeholder positions on a topic."""
    if not GEMINI_API_KEY:
        return None
    hansard_rows = hansard_rows or []
    news = news or []
    publications = publications or []
    social = social or []

    org_context = f'Organisation: {org.name}\n' if org else 'Multiple stakeholder organisations\n'

    # Build evidence summary (capped to avoid 503)
    evidence_lines = []
    for r in hansard_rows[:8]:
        evidence_lines.append(f'[Hansard] {r.get("speaker_name","")} ({r.get("hdate","")}): '
                               f'{clean_body_text(r.get("body",""))[:200]}')
    for r in news[:6]:
        evidence_lines.append(f'[News/{r.get("outlet","")}] {r.get("title","")} ({r.get("published","")}): '
                               f'{r.get("summary","")[:200]}')
    for r in publications[:6]:
        evidence_lines.append(f'[Publication] {r.get("title","")} ({r.get("published","")}): '
                               f'{r.get("summary","")[:200]}')
    for r in social[:4]:
        evidence_lines.append(f'[Social] {r.get("summary","")[:150]}')

    if not evidence_lines:
        return None

    evidence_text = '\n'.join(evidence_lines)
    prompt = (
        f'You are a UK civil service parliamentary researcher. Analyse the following evidence '
        f'about what stakeholders are saying on the topic: "{topic}".\n\n'
        f'{org_context}\n'
        f'EVIDENCE:\n{evidence_text}\n\n'
        f'Return a JSON object with these fields:\n'
        f'"summary": 2-3 sentence overview of the stakeholder landscape on this topic.\n'
        f'"stated_positions": Array of {{\"org\", \"position\", \"source_type\", \"source_date\"}} — '
        f'one entry per organisation visible in the evidence, summarising their stated position in 1-2 sentences. '
        f'Include the type of source (Hansard/News/Publication/Social) and date.\n'
        f'"key_asks": Array of strings — the main policy asks or demands from stakeholders (up to 5).\n'
        f'"coverage_note": One sentence on the breadth and recency of evidence found.\n'
        f'Return ONLY valid JSON, no markdown fences.'
    )

    try:
        model_path = get_working_model(GEMINI_API_KEY)
        ai_url = (f'https://generativelanguage.googleapis.com/v1beta/{model_path}'
                  f':generateContent?key={GEMINI_API_KEY}')
        payload = {'contents': [{'parts': [{'text': prompt}]}],
                   'generationConfig': {'temperature': 0.2, 'maxOutputTokens': 800}}
        resp = requests.post(ai_url, json=payload, timeout=25)
        if resp.status_code == 503:
            time.sleep(3)
            resp = requests.post(ai_url, json=payload, timeout=25)
        if resp.status_code == 200:
            raw = resp.json()['candidates'][0]['content']['parts'][0]['text']
        else:
            raw = _claude_fallback(prompt, max_tokens=800)
        if raw:
            return _parse_ai_json(raw)
    except Exception:
        pass
    return None


def _discover_rss(website):
    """Try to auto-discover an RSS feed URL from a website domain."""
    if not website:
        return None
    base = website.rstrip('/')
    if not base.startswith('http'):
        base = 'https://' + base
    candidates = [
        base + '/feed', base + '/rss', base + '/rss.xml',
        base + '/feed.xml', base + '/news/rss', base + '/news/feed',
        base + '/blog/feed', base + '/publications/rss',
    ]
    for url in candidates:
        try:
            r = requests.get(url, timeout=6, allow_redirects=True,
                             headers={'User-Agent': 'Mozilla/5.0'})
            ct = r.headers.get('Content-Type', '')
            if r.status_code == 200 and ('xml' in ct or 'rss' in ct or 'atom' in ct
                                          or r.text.strip().startswith('<')):
                return url
        except Exception:
            pass
    return None


# ==========================================
# ROUTE 9: STAKEHOLDER SEARCH
# ==========================================
@debate_scanner_bp.route('/stakeholder_search', methods=['POST'])
def stakeholder_search():
    topic = request.form.get('stakeholder_topic', '').strip()
    org_id = request.form.get('stakeholder_org_id', '').strip()
    start_date = request.form.get('stakeholder_start', '')
    end_date = request.form.get('stakeholder_end', '')

    all_orgs = StakeholderOrg.all_active()
    orgs_by_category = StakeholderOrg.by_category()

    if not topic:
        return render_template('debate_scanner.html', mode='stakeholder',
                               stakeholder_topic='', stakeholder_org=None,
                               all_orgs=all_orgs, orgs_by_category=orgs_by_category,
                               stakeholder_hansard=[], stakeholder_news=[],
                               stakeholder_publications=[], stakeholder_social=[],
                               stakeholder_briefing=None,
                               # other tab defaults
                               topic='', topic_rows=[], topic_briefing=None,
                               topic_briefing_as_text='', house_filter='all',
                               selected_depts=[], start_date='', end_date='',
                               wq_rows=[], oral_grouped=[], urgent_grouped=[],
                               statement_grouped=[], debate_grouped=[],
                               legislation_grouped=[], debug_query='',
                               user_pref=None, is_post=False)

    # Resolve selected org (if any)
    org = None
    if org_id:
        try:
            org = StakeholderOrg.query.get(int(org_id))
        except Exception:
            pass

    # Determine which orgs to search
    search_orgs = [org] if org else all_orgs
    org_names = [o.hansard_search_name or o.short_name or o.name for o in search_orgs]

    date_range = ''
    if start_date and end_date:
        date_range = f'{start_date.replace("-","")}..{end_date.replace("-","")}'
    elif start_date:
        date_range = f'{start_date.replace("-","")}..{datetime.now().strftime("%Y%m%d")}'

    # Build Hansard search: topic + org name(s)
    if org:
        hansard_query = f'"{topic}" AND "{org.hansard_search_name or org.name}"'
    else:
        hansard_query = f'"{topic}"'
        if date_range:
            hansard_query += f' {date_range}'

    # Parallel fetch
    hansard_rows = []
    news_items = []
    publications = []
    social_posts = []

    def _do_hansard():
        rows = []
        for src in ['commons', 'lords', 'westminsterhall']:
            rows.extend(fetch_twfy_topic(hansard_query, src, date_range))
        return deduplicate_by_listurl(rows)

    def _do_news():
        return fetch_stakeholder_news(org_names, topic, start_date, end_date)

    def _do_rss():
        results = []
        rss_orgs = [o for o in search_orgs if o.rss_url]
        for o in rss_orgs[:10]:  # cap parallel RSS fetches
            items = fetch_org_rss(o.rss_url, topic)
            for item in items:
                item['org_name'] = o.name
            results.extend(items)
        return sorted(results, key=lambda x: x.get('published', ''), reverse=True)

    def _do_bluesky():
        results = []
        bsky_orgs = [o for o in search_orgs if o.bsky_handle]
        for o in bsky_orgs[:5]:
            items = fetch_org_bluesky(o.bsky_handle, topic)
            for item in items:
                item['org_name'] = o.name
            results.extend(items)
        return sorted(results, key=lambda x: x.get('published', ''), reverse=True)

    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as executor:
        fut_hansard = executor.submit(copy_current_request_context(_do_hansard))
        fut_news = executor.submit(copy_current_request_context(_do_news))
        fut_rss = executor.submit(copy_current_request_context(_do_rss))
        fut_bsky = executor.submit(copy_current_request_context(_do_bluesky))
        try:
            hansard_rows = fut_hansard.result(timeout=30)
        except Exception:
            hansard_rows = []
        try:
            news_items = fut_news.result(timeout=20)
        except Exception:
            news_items = []
        try:
            publications = fut_rss.result(timeout=20)
        except Exception:
            publications = []
        try:
            social_posts = fut_bsky.result(timeout=20)
        except Exception:
            social_posts = []

    # Apply date filter to Hansard rows
    if start_date or end_date:
        hansard_rows = [r for r in hansard_rows
                        if (not start_date or r.get('hdate', '') >= start_date)
                        and (not end_date or r.get('hdate', '') <= end_date)]

    # Flag ministers in Hansard results
    minister_data = get_minister_list()
    by_norm = minister_data.get('by_norm', {})
    for row in hansard_rows:
        spk = row.get('speaker_name', '')
        norm_spk = _normalise_name(spk)
        role = by_norm.get(norm_spk) if norm_spk else None
        row['is_minister'] = bool(role)
        row['minister_role'] = role or ''
        row['body_clean'] = clean_body_text(row.get('body', ''))

    # AI briefing
    briefing = None
    if hansard_rows or news_items or publications or social_posts:
        try:
            briefing = generate_stakeholder_briefing(
                topic, org=org,
                hansard_rows=hansard_rows,
                news=news_items,
                publications=publications,
                social=social_posts,
            )
        except Exception:
            pass

    user_pref = _get_user_pref()

    return render_template('debate_scanner.html',
        mode='stakeholder',
        stakeholder_topic=topic,
        stakeholder_org=org,
        stakeholder_hansard=hansard_rows,
        stakeholder_news=news_items,
        stakeholder_publications=publications,
        stakeholder_social=social_posts,
        stakeholder_briefing=briefing,
        all_orgs=all_orgs,
        orgs_by_category=orgs_by_category,
        # other tab defaults (required by template)
        topic=topic, topic_rows=[], topic_briefing=None,
        topic_briefing_as_text='', house_filter='all',
        selected_depts=[], start_date=start_date, end_date=end_date,
        wq_rows=[], oral_grouped=[], urgent_grouped=[],
        statement_grouped=[], debate_grouped=[],
        legislation_grouped=[], debug_query='',
        user_pref=user_pref, is_post=True)


# ==========================================
# ROUTE 10: ADD STAKEHOLDER ORG
# ==========================================
@debate_scanner_bp.route('/stakeholder_add', methods=['POST'])
@login_required
def stakeholder_add():
    name = request.form.get('org_name', '').strip()
    website = request.form.get('org_website', '').strip().lstrip('https://').lstrip('http://').rstrip('/')
    department = request.form.get('org_department', '').strip()

    if not name or not website:
        return jsonify({'error': 'Name and website required'}), 400

    # Map department to category
    dept_to_category = {
        'Department for Education': 'Education',
        'Department of Health and Social Care': 'Health & Social Care',
        'HM Treasury': 'Economics & Finance',
        'Home Office': 'Home Affairs',
        'Ministry of Defence': 'Defence',
        'Ministry of Justice': 'Justice',
        'Department for Science, Innovation and Technology': 'Science & Technology',
        'Cabinet Office': 'Government & Public Administration',
    }
    category = dept_to_category.get(department, department or 'Other')

    from extensions import db as _db

    # Check for duplicate
    existing = StakeholderOrg.query.filter(
        StakeholderOrg.name.ilike(name)
    ).first()
    if existing:
        return jsonify({'error': f'"{name}" is already in the directory'}), 409

    # Attempt RSS auto-discovery (best-effort)
    rss_url = None
    try:
        rss_url = _discover_rss(website)
    except Exception:
        pass

    org = StakeholderOrg(
        name=name, category=category,
        website=website, rss_url=rss_url,
        hansard_search_name=name, active=True,
    )
    try:
        _db.session.add(org)
        _db.session.commit()
    except Exception:
        _db.session.rollback()
        return jsonify({'error': 'Could not save organisation'}), 500

    return jsonify({
        'id': org.id, 'name': org.name, 'category': org.category,
        'website': org.website, 'rss_found': bool(rss_url),
    })


# ==========================================
# ROUTE 8: SECTIONED RESEARCH BRIEF DOWNLOAD
# ==========================================
@debate_scanner_bp.route('/download_research_brief', methods=['POST'])
def download_research_brief():
    if not Document:
        return "Word library missing.", 500

    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    topic = request.form.get('topic', 'Parliamentary Research').strip()
    mode = request.form.get('mode', '')
    briefing_text = request.form.get('briefing_text', '')
    sections_json = request.form.get('sections_json', '[]')
    depts_json = request.form.get('depts_json', '[]')
    briefing_json = request.form.get('briefing_json', '')
    start_date = request.form.get('start_date', '')
    end_date = request.form.get('end_date', '')
    try:
        sections = json.loads(sections_json)
    except Exception:
        sections = []
    try:
        depts = json.loads(depts_json)
    except Exception:
        depts = []
    try:
        briefing_struct = json.loads(briefing_json) if briefing_json else None
        if isinstance(briefing_struct, dict) and not briefing_struct:
            briefing_struct = None  # treat empty {} as no briefing
    except Exception:
        briefing_struct = None

    # ── Stakeholder brief mode ──────────────────────────────────────
    if mode == 'stakeholder':
        def _parse(field):
            try:
                return json.loads(request.form.get(field, '[]'))
            except Exception:
                return []

        sk_news    = _parse('stakeholder_news')
        sk_pubs    = _parse('stakeholder_publications')
        sk_social  = _parse('stakeholder_social')
        sk_hansard = _parse('stakeholder_hansard')

        doc = Document()
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        def _clr(run, hex_str):
            r, g, b = int(hex_str[0:2],16), int(hex_str[2:4],16), int(hex_str[4:6],16)
            run.font.color.rgb = RGBColor(r, g, b)

        # Cover heading
        title_p = doc.add_heading(f'Stakeholder Research Brief', 0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = doc.add_paragraph(f'Topic: {topic}   ·   Generated {datetime.now().strftime("%d %b %Y")}')
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub.runs[0].italic = True
        doc.add_paragraph()

        notice = doc.add_paragraph(
            '⚠ AI-generated — verify positions against source links before use in official documents.')
        notice.runs[0].font.size = Pt(9)
        notice.runs[0].italic = True
        _clr(notice.runs[0], '856404')
        doc.add_paragraph()

        # AI summary section
        if briefing_struct:
            b = briefing_struct
            doc.add_heading('AI Stakeholder Summary', 1)
            if b.get('summary'):
                doc.add_paragraph(b['summary'])
            if b.get('stated_positions'):
                doc.add_heading('Stated Positions', 2)
                for p in b['stated_positions']:
                    para = doc.add_paragraph(style='List Bullet')
                    run = para.add_run(p.get('org', '') + ': ')
                    run.bold = True
                    src = p.get('source_type', '')
                    dt = p.get('source_date', '')
                    meta = f' [{", ".join(filter(None, [src, dt]))}]' if src or dt else ''
                    para.add_run(p.get('position', '') + meta)
            if b.get('key_asks'):
                doc.add_heading('Key Asks from Stakeholders', 2)
                for ask in b['key_asks']:
                    doc.add_paragraph(ask, style='List Bullet')
            if b.get('coverage_note'):
                p = doc.add_paragraph(b['coverage_note'])
                p.runs[0].italic = True
                p.runs[0].font.size = Pt(9)
            doc.add_paragraph()

        def _add_item_table(heading, items, cols):
            """Add a simple table section for stakeholder results."""
            if not items:
                return
            doc.add_heading(heading, 1)
            tbl = doc.add_table(rows=1, cols=len(cols))
            tbl.style = 'Table Grid'
            for i, label in enumerate(cols):
                cell = tbl.rows[0].cells[i]
                cell.text = label
                cell.paragraphs[0].runs[0].bold = True
            for item in items:
                row = tbl.add_row().cells
                if cols[0] == 'Date':
                    row[0].text = item.get('published', '')
                    row[1].text = item.get('org_name', '') or item.get('outlet', '')
                    row[2].text = item.get('title', '')
                    if len(cols) > 3:
                        row[3].text = (item.get('summary', '') or '')[:200]
                else:
                    row[0].text = item.get('speaker_name', '')
                    row[1].text = item.get('hdate', '')
                    row[2].text = clean_body_text(item.get('body', ''))[:300]
            doc.add_paragraph()

        _add_item_table('News Media Coverage', sk_news,
                        ['Date', 'Outlet', 'Headline', 'Summary'])
        _add_item_table('Reports & Publications', sk_pubs,
                        ['Date', 'Organisation', 'Title', 'Summary'])
        _add_item_table('Social Media (Bluesky)', sk_social,
                        ['Date', 'Organisation', 'Post', 'Link'])
        _add_item_table('Parliamentary Mentions', sk_hansard,
                        ['Speaker', 'Date', 'Extract'])

        mem_doc = io.BytesIO()
        doc.save(mem_doc)
        mem_doc.seek(0)
        safe_topic = re.sub(r'[^\w\s-]', '', topic)[:40].strip()
        filename = f'Stakeholder Brief - {safe_topic} - {datetime.now().strftime("%Y%m%d")}.docx'
        return send_file(mem_doc, as_attachment=True, download_name=filename,
                         mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    # ── end stakeholder mode ────────────────────────────────────────

    SECTION_TITLES = {
        'wq': 'Written Questions & Answers',
        'oral': 'Oral Questions',
        'urgent': 'Urgent Questions',
        'statement': 'Ministerial Statements',
        'debate': 'Parliamentary Debates',
        'legislation': 'Legislation',
    }
    REL_LABELS = {'high': 'HIGH RELEVANCE', 'medium': 'MODERATE RELEVANCE', 'low': 'SESSION CONTEXT'}

    def _set_cell_bg(cell, hex_colour):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = _OxmlElement('w:shd')
        shd.set(_qn('w:val'), 'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'), hex_colour)
        tcPr.append(shd)

    def _set_run_colour(run, hex_colour):
        run.font.color.rgb = RGBColor(
            int(hex_colour[0:2], 16),
            int(hex_colour[2:4], 16),
            int(hex_colour[4:6], 16)
        )

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────
    for section_obj in doc.sections:
        section_obj.top_margin = Inches(0.9)
        section_obj.bottom_margin = Inches(0.9)
        section_obj.left_margin = Inches(1.1)
        section_obj.right_margin = Inches(1.1)

    # ── Title block ───────────────────────────────────────────────
    title_p = doc.add_heading('Parliamentary Research Brief', 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_lines = [f'Topic: {topic}']
    if depts:
        meta_lines.append(f'Department(s): {", ".join(depts)}')
    date_parts = []
    if start_date: date_parts.append(f'from {start_date}')
    if end_date: date_parts.append(f'to {end_date}')
    if date_parts:
        meta_lines.append('Date range: ' + ' '.join(date_parts))
    meta_lines.append(f'Generated: {datetime.now().strftime("%d %B %Y at %H:%M")}')
    meta_lines.append('Source: Westminster Brief — westminsterbrief.co.uk')

    for line in meta_lines:
        mp = doc.add_paragraph(line)
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mp.runs[0].font.size = Pt(10)
        mp.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    # ── Summary table ─────────────────────────────────────────────
    non_empty = [(s, SECTION_TITLES.get(s.get('type',''), s.get('type','').title()), s.get('items', []))
                 for s in sections if s.get('items')]
    if non_empty:
        doc.add_heading('Contents Summary', 1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, label in enumerate(['Section', 'Contributions', 'Notes']):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(hdr[i], '1C3E6E')
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for sec, title, items in non_empty:
            row = tbl.add_row().cells
            row[0].text = title
            row[1].text = str(len(items))
            ministers = sum(1 for r in items if r.get('is_minister'))
            row[2].text = f'{ministers} minister(s)' if ministers else '—'
        doc.add_paragraph()

    # ── AI Summary ────────────────────────────────────────────────
    if briefing_struct or briefing_text:
        doc.add_heading('AI Parliamentary Briefing', 1)
        notice = doc.add_paragraph(
            '⚠ AI-generated — review for accuracy and remove any non-neutral language before use.')
        notice.runs[0].font.size = Pt(9)
        notice.runs[0].italic = True
        _set_run_colour(notice.runs[0], '856404')

        if briefing_struct:
            b = briefing_struct
            # Topic summary
            if b.get('topic_summary'):
                doc.add_heading('Summary', 2)
                doc.add_paragraph(b['topic_summary'])

            # Government / Opposition positions side by side as plain sections
            if b.get('government_position'):
                doc.add_heading('Government Position', 2)
                doc.add_paragraph(b['government_position'])
            if b.get('opposition_position'):
                doc.add_heading('Opposition Position', 2)
                doc.add_paragraph(b['opposition_position'])

            # Government speakers
            govt_spks = b.get('government_speakers', [])
            if govt_spks:
                doc.add_heading('Government Speakers', 2)
                for s in govt_spks:
                    role = s.get('confirmed_role') or s.get('role', '')
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(s.get('name', '')).bold = True
                    if role:
                        p.add_run(f' ({role})')
                    if s.get('stance'):
                        p.add_run(f': {s["stance"]}')

            # Opposition speakers
            opp_spks = b.get('non_government_speakers', [])
            if opp_spks:
                doc.add_heading('Opposition / Non-Government Speakers', 2)
                for s in opp_spks:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(s.get('name', '')).bold = True
                    if s.get('role_or_party'):
                        p.add_run(f' ({s["role_or_party"]})')
                    if s.get('stance'):
                        p.add_run(f': {s["stance"]}')

            # Key opposition questions
            kqs = b.get('key_questions', [])
            if kqs:
                doc.add_heading('Key Opposition Questions', 2)
                doc.add_paragraph('Questions raised in Parliament — likely lines of challenge in future debates.').italic = True
                for q in kqs:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f'{q.get("speaker", "")}').bold = True
                    role_party = q.get('role_or_party', '')
                    date = q.get('date', '')
                    if role_party or date:
                        p.add_run(f' ({", ".join(filter(None, [role_party, date]))})')
                    p.add_run(f': {q.get("question", "")}')

            # Anticipated / harder questions
            aqs = b.get('anticipated_questions', [])
            if aqs:
                doc.add_heading('Anticipated Questions (AI-Generated)', 2)
                doc.add_paragraph(
                    'Questions not yet raised in Parliament but likely based on gaps in government answers '
                    'and opposition lines of attack. For ministerial preparation only.').italic = True
                for q in aqs:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(q.get('question', '')).bold = False
                    if q.get('rationale'):
                        rat = doc.add_paragraph(f'Why likely: {q["rationale"]}')
                        rat.runs[0].font.size = Pt(9)
                        rat.runs[0].italic = True
                        _set_run_colour(rat.runs[0], '666666')

            # Next steps
            if b.get('next_steps'):
                doc.add_heading('Next Steps', 2)
                doc.add_paragraph(b['next_steps'])

            if b.get('coverage_note'):
                p = doc.add_paragraph(f'Coverage note: {b["coverage_note"]}')
                p.runs[0].font.size = Pt(9)
                _set_run_colour(p.runs[0], '666666')

        elif briefing_text:
            # Fallback: render flat text if structured briefing not available
            for line in briefing_text.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('## '):
                    doc.add_heading(line[3:].strip(), level=2)
                elif line.startswith('* ') or line.startswith('- '):
                    doc.add_paragraph(line[2:].strip(), style='List Bullet')
                else:
                    doc.add_paragraph(line.replace('**', ''))

        doc.add_paragraph()

    # ── Sections ──────────────────────────────────────────────────
    SECTION_ORDER = ['debate', 'oral', 'urgent', 'statement', 'wq']
    sections_by_type = {s.get('type'): s for s in sections}

    for sec_type in SECTION_ORDER:
        section = sections_by_type.get(sec_type)
        if not section:
            continue
        items = section.get('items', [])
        if not items:
            continue
        title = SECTION_TITLES.get(sec_type, sec_type.title())
        doc.add_heading(f'{title} ({len(items)})', 1)

        if sec_type == 'wq':
            # Written Questions as a table
            tbl = doc.add_table(rows=1, cols=5)
            tbl.style = 'Table Grid'
            hdr = tbl.rows[0].cells
            for i, label in enumerate(['Status', 'Date', 'Department', 'Member (Party)', 'Question / Answer']):
                hdr[i].text = label
                hdr[i].paragraphs[0].runs[0].bold = True
                _set_cell_bg(hdr[i], '1C3E6E')
                hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            for q in items:
                row = tbl.add_row().cells
                status = 'Answered' if q.get('is_answered') else 'Unanswered'
                if q.get('is_withdrawn'): status = 'Withdrawn'
                if q.get('is_holding'): status = 'Holding'
                row[0].text = status
                row[1].text = q.get('date_tabled', '')
                row[2].text = q.get('dept', '')
                mp_info = q.get('asking_mp', '')
                party = q.get('party', '')
                row[3].text = f'{mp_info} ({party})' if party else mp_info

                qa_cell = row[4]
                qa_para = qa_cell.paragraphs[0]
                qa_para.add_run('Q: ').bold = True
                qa_para.add_run(q.get('question_text', ''))
                if q.get('answer_text'):
                    ans_para = qa_cell.add_paragraph()
                    ans_para.add_run('A: ').bold = True
                    ans_para.add_run(q.get('answer_text', ''))
                if q.get('url'):
                    link_para = qa_cell.add_paragraph()
                    _add_hyperlink(link_para, q['url'], '↗ Parliament.uk')
            doc.add_paragraph()

        else:
            # Group debate speeches under their debate heading
            # Build ordered list of unique debates preserving minister-first order
            seen_debates = {}
            ordered_debates = []
            for r in items:
                key = r.get('listurl', '')[:50] or r.get('debate_title', '')
                if key not in seen_debates:
                    seen_debates[key] = []
                    ordered_debates.append(key)
                seen_debates[key].append(r)

            for debate_key in ordered_debates:
                speeches = seen_debates[debate_key]
                first = speeches[0]

                # Debate header
                debate_hdr = doc.add_paragraph()
                rel = first.get('relevance_level', '')
                rel_label = REL_LABELS.get(rel, '')
                if rel_label:
                    rl_run = debate_hdr.add_run(f'[{rel_label}]  ')
                    rl_run.bold = True
                    rl_run.font.size = Pt(9)
                    colour = {'HIGH RELEVANCE': '166534', 'MODERATE RELEVANCE': '92400E', 'SESSION CONTEXT': '6B7280'}.get(rel_label, '444444')
                    _set_run_colour(rl_run, colour)
                src_run = debate_hdr.add_run(
                    f"{first.get('source_label', '')}  ·  {first.get('hdate', '')}  ·  {first.get('debate_type', '')}"
                )
                src_run.font.size = Pt(9)
                _set_run_colour(src_run, '555555')

                title_line = doc.add_paragraph(first.get('debate_title', ''))
                title_line.runs[0].bold = True
                title_line.runs[0].font.size = Pt(11)

                for r in speeches:
                    spk_p = doc.add_paragraph()
                    spk_run = spk_p.add_run(r.get('speaker_name', ''))
                    spk_run.bold = True
                    if r.get('is_minister'):
                        role_run = spk_p.add_run(f"  —  {r.get('minister_role', 'Minister')}")
                        role_run.font.size = Pt(9)
                        _set_run_colour(role_run, '1C3E6E')
                    elif r.get('speaker_party'):
                        party_run = spk_p.add_run(f"  ({r['speaker_party']})")
                        party_run.font.size = Pt(9)
                        _set_run_colour(party_run, '555555')

                    # Full text for ministers (body_export up to 3000 chars); snippet for others
                    body = r.get('body_export') or r.get('body_clean', '')
                    if not r.get('is_minister'):
                        body = body[:600]
                    if body:
                        body_p = doc.add_paragraph(body)
                        body_p.runs[0].font.size = Pt(10)
                        if len(r.get('body_export', r.get('body_clean', ''))) > len(body):
                            body_p.add_run('  […]').italic = True

                    url = r.get('listurl', '')
                    if url:
                        if not url.startswith('http'):
                            url = 'https://www.theyworkforyou.com' + url
                        lp = doc.add_paragraph()
                        _add_hyperlink(lp, url, '↗ View on Hansard')
                        if lp.runs:
                            lp.runs[0].font.size = Pt(9)

                doc.add_paragraph('─' * 72)
            doc.add_paragraph()

    # ── Debates — Sources appendix ────────────────────────────────────
    # Collect unique debate titles + links from all non-WQ sections
    all_debate_links = {}
    for sec_type in ['debate', 'oral', 'urgent', 'statement']:
        section = sections_by_type.get(sec_type)
        if not section:
            continue
        for r in section.get('items', []):
            title = r.get('debate_title', '') or ''
            url = r.get('listurl', '') or ''
            if url and not url.startswith('http'):
                url = 'https://www.theyworkforyou.com' + url
            date = r.get('hdate', '') or ''
            source = r.get('source_label', '') or ''
            key = url or title
            if key and key not in all_debate_links:
                all_debate_links[key] = {'title': title, 'url': url, 'date': date, 'source': source}

    if all_debate_links:
        doc.add_heading('Debates — Sources', 1)
        for entry in all_debate_links.values():
            p = doc.add_paragraph(style='List Bullet')
            label = '  ·  '.join(filter(None, [entry['source'], entry['date'], entry['title']]))
            if entry['url']:
                _add_hyperlink(p, entry['url'], label)
            else:
                p.add_run(label)
        doc.add_paragraph()

    mem_doc = io.BytesIO()
    doc.save(mem_doc)
    mem_doc.seek(0)
    safe_topic = re.sub(r'[^\w\s-]', '', topic)[:40].strip()
    filename = f"Research - {safe_topic} - {datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(mem_doc, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


# ==========================================
# DEBATE PREP — Helper functions
# ==========================================

def _prep_resolve_peer(peer_name):
    """Look up a peer's Parliament profile.
    Returns dict with parliament_id, display_name, party, house,
    biography_posts, interests — or None if not found."""
    # Build a list of name variants to try — full name first, then with
    # honorifics stripped (Parliament API searches on the name part, not title)
    _HONORIFICS = re.compile(
        r'^(The\s+)?(Lord|Baroness|Baron|Lady|Viscount|Viscountess|Earl|Countess|Duke|Duchess|Marquess|Prince|Princess)\s+',
        re.IGNORECASE
    )
    name_variants = [peer_name]
    stripped = _HONORIFICS.sub('', peer_name).strip()
    if stripped and stripped != peer_name:
        name_variants.append(stripped)

    try:
        items = []
        for name in name_variants:
            if items:
                break
            resp = requests.get(
                'https://members-api.parliament.uk/api/Members/Search',
                params={'Name': name, 'IsCurrentMember': 'true', 'take': 5},
                timeout=8
            )
            if resp.status_code == 200:
                items = resp.json().get('items', [])
            if not items:
                resp2 = requests.get(
                    'https://members-api.parliament.uk/api/Members/Search',
                    params={'Name': name, 'take': 5},
                    timeout=8
                )
                if resp2.status_code == 200:
                    items = resp2.json().get('items', [])
        if not items:
            return None

        v = items[0]['value']
        parliament_id = v['id']
        house_num = (v.get('latestHouseMembership') or {}).get('house', 1)
        house = 'Lords' if house_num == 2 else 'Commons'
        party = (v.get('latestParty') or {}).get('name', '')
        name = v.get('nameDisplayAs', peer_name)
        thumbnail = v.get('thumbnailUrl', '')

        bio_posts = []
        try:
            bio_resp = requests.get(
                f'https://members-api.parliament.uk/api/Members/{parliament_id}/Biography',
                timeout=8
            )
            if bio_resp.status_code == 200:
                bio = bio_resp.json().get('value', {})
                for category_key, cat_label in [
                    ('governmentPosts', 'Government Post'),
                    ('oppositionPosts', 'Opposition Post'),
                    ('committeeMemberships', 'Committee'),
                    ('representations', 'Career'),
                ]:
                    for post in (bio.get(category_key) or []):
                        bio_posts.append({
                            'category': cat_label,
                            'title': post.get('name', ''),
                            'start': (post.get('startDate') or '')[:10],
                            'end': (post.get('endDate') or '')[:10],
                            'is_current': not post.get('endDate'),
                        })
        except Exception:
            pass

        interests = []
        try:
            int_resp = requests.get(
                f'https://members-api.parliament.uk/api/Members/{parliament_id}/RegisteredInterests',
                timeout=8
            )
            if int_resp.status_code == 200:
                for category in (int_resp.json().get('value') or []):
                    cat_name = category.get('name', '')
                    cat_items = []
                    for item in (category.get('interests') or []):
                        desc = item.get('interest', '') or item.get('description', '')
                        if desc:
                            cat_items.append(desc[:200])
                    if cat_items:
                        interests.append({'category': cat_name, 'items': cat_items})
        except Exception:
            pass

        return {
            'parliament_id': parliament_id,
            'display_name': name,
            'party': party,
            'house': house,
            'thumbnail': thumbnail,
            'biography_posts': bio_posts,
            'interests': interests,
        }
    except Exception:
        return None


def _prep_one_pager(question_text, topic, question_date='', house='lords'):
    """Generate an AI one-pager briefing for an oral question.
    Returns dict with why_now, sector_context, major_criticisms, opposition_position
    or None on failure."""
    if not GEMINI_API_KEY and not CLAUDE_API_KEY:
        return None

    house_label = 'Commons oral question' if house == 'commons' else 'Lords oral question'
    prompt = (
        f"You are briefing a UK civil servant preparing for a {house_label}.\n\n"
        f"Question: {question_text}\n"
        f"Topic: {topic}\n"
        f"Date: {question_date or 'upcoming'}\n\n"
        "Generate a structured background briefing with these five sections:\n"
        "1. why_now: 2-3 sentences explaining why this question is being asked at this "
        "time — recent events, policy changes, sector pressures, or public debate driving it\n"
        "2. sector_context: 3-4 sentences of factual background on the sector or policy area\n"
        "3. major_criticisms: Array of 3-5 specific criticisms or concerns currently being "
        "raised about this area by campaigners, academics, media or opposition\n"
        "4. opposition_position: 2-3 sentences summarising the likely opposition angle — "
        "what they want the government to commit to, and what they may push back on\n"
        "5. sources: Array of 3-6 specific, real sources that a civil servant could consult — "
        "e.g. named government reports, parliamentary publications, think-tank research, "
        "named media outlets or academic bodies. Use the format 'Organisation/Publication — "
        "Title or description (Year if known)'. Only cite sources that are likely to genuinely "
        "exist; do not invent titles.\n\n"
        'Return as JSON only, no prose, no markdown fences:\n'
        '{"why_now": "...", "sector_context": "...", "major_criticisms": ["...", "..."], '
        '"opposition_position": "...", "sources": ["...", "..."]}'
    )

    if GEMINI_API_KEY:
        try:
            model_path = get_working_model(GEMINI_API_KEY)
            url = (f"https://generativelanguage.googleapis.com/v1beta/"
                   f"{model_path}:generateContent?key={GEMINI_API_KEY}")
            body = {"contents": [{"parts": [{"text": prompt}]}],
                    "generationConfig": {"temperature": 0.3, "maxOutputTokens": 1400}}
            resp = requests.post(url, json=body, timeout=30)
            if resp.status_code == 200:
                raw = resp.json()['candidates'][0]['content']['parts'][0]['text']
                result = _parse_ai_json(raw)
                if result:
                    return result
        except Exception:
            pass

    raw = _claude_fallback(prompt, max_tokens=1000)
    if raw:
        return _parse_ai_json(raw)
    return None


_NEGATIVE_SENTIMENT_WORDS = {
    'crisis', 'failure', 'fail', 'cut', 'cuts', 'closure', 'closures', 'lost',
    'concern', 'problem', 'criticism', 'shortage', 'scandal', 'row', 'backlash',
    'warning', 'threat', 'decline', 'collapse', 'protest', 'underfund', 'underfunded',
    'anger', 'outrage', 'catastrophe', 'catastrophic', 'devastating',
}
_POSITIVE_SENTIMENT_WORDS = {
    'invest', 'investment', 'growth', 'support', 'success', 'improve', 'improvement',
    'boost', 'achieve', 'achievement', 'celebrate', 'award', 'record', 'launch',
    'fund', 'expand', 'expansion', 'progress', 'reform',
}


def _classify_sentiment(title, summary):
    text = (title + ' ' + (summary or '')).lower()
    words = set(re.findall(r'\b\w+\b', text))
    neg_score = len(words & _NEGATIVE_SENTIMENT_WORDS)
    pos_score = len(words & _POSITIVE_SENTIMENT_WORDS)
    if neg_score > pos_score:
        return 'negative'
    if pos_score > neg_score:
        return 'positive'
    return 'neutral'


def _prep_media(topic, start_date='', end_date=''):
    """Search News API for media coverage of the topic.
    Returns list of {title, outlet, published, summary, link, sentiment}."""
    NEWS_API_KEY = os.environ.get('NEWS_API_KEY')
    if not NEWS_API_KEY:
        return []
    try:
        from newsapi import NewsApiClient
        newsapi = NewsApiClient(api_key=NEWS_API_KEY)
        kwargs = {'q': topic, 'language': 'en', 'sort_by': 'relevancy', 'page_size': 20}
        if start_date:
            kwargs['from_param'] = start_date
        if end_date:
            kwargs['to'] = end_date
        resp = newsapi.get_everything(**kwargs)
        results = []
        for art in resp.get('articles', []):
            title = art.get('title', '')
            summary = art.get('description') or art.get('title', '')
            results.append({
                'title': title,
                'link': art.get('url', ''),
                'published': (art.get('publishedAt') or '')[:10],
                'summary': summary[:200],
                'outlet': art.get('source', {}).get('name', ''),
                'sentiment': _classify_sentiment(title, summary),
            })
        return results
    except Exception:
        return []


def _prep_parl_lords(topic, date_range):
    """Lords debates for topic — called directly in the top-level executor."""
    try:
        rows = fetch_twfy_topic(topic, 'lords', date_range, num=50)
        return [r for r in rows if not r.get('_error')]
    except Exception:
        return []


def _prep_parl_commons(topic, date_range):
    """Commons + Westminster Hall debates for topic — called directly in the top-level executor."""
    try:
        rows = []
        rows.extend(fetch_twfy_topic(topic, 'commons', date_range, num=40))
        rows.extend(fetch_twfy_topic(topic, 'westminsterhall', date_range, num=30))
        return [r for r in rows if not r.get('_error')]
    except Exception:
        return []


def _prep_parl_wqs(topic, start_date, end_date):
    """Written questions for topic — called directly in the top-level executor."""
    try:
        wqs, _ = _fetch_topic_wqs(topic, start_date, end_date, [], limit=20)
        return wqs
    except Exception:
        return []


def _prep_parl_statements(topic, date_range):
    """Oral/written ministerial statements for topic — called directly in the top-level executor."""
    try:
        rows = fetch_twfy_topic(topic, 'wms', date_range, num=20)
        return [r for r in rows if not r.get('_error')]
    except Exception:
        return []


def _prep_peer_contributions(parliament_id, start_date='', end_date=''):
    """Fetch the peer's own recent speeches and tabled WPQs.
    Returns dict: {speeches, tabled_wqs}."""
    date_range = ''
    if start_date and end_date:
        date_range = f'{start_date.replace("-","")}..{end_date.replace("-","")}'
    elif start_date:
        date_range = f'{start_date.replace("-","")}..{datetime.now().strftime("%Y%m%d")}'

    def _speeches():
        return fetch_hansard_minister_topic(
            parliament_id, '', date_range,
            ['lords', 'commons', 'wms', 'westminsterhall'],
            num=20, is_lord=True
        )

    def _wqs():
        try:
            params = {
                'tabledBy': parliament_id, 'take': 20, 'skip': 0,
                'expandMember': 'false', 'orderBy': 'DateTabledDesc',
            }
            resp = requests.get(PARLIAMENT_WQ_API, params=params, timeout=15)
            if resp.status_code != 200:
                return []
            data = resp.json()
            results = []
            for item in data.get('results', []):
                val = item.get('value', {})
                raw_date = (val.get('dateTabled') or '').split('T')[0]
                uin = str(val.get('uin', ''))
                q_text = re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ',
                                val.get('questionText') or '')).strip()
                dept = val.get('answeringBodyName', '')
                results.append({
                    'uin': uin,
                    'question_text': q_text[:300],
                    'dept': dept,
                    'date_tabled': raw_date,
                    'url': (f"https://questions-statements.parliament.uk/written-questions"
                            f"/detail/{raw_date}/{uin}"),
                })
            return results
        except Exception:
            return []

    try:
        speeches = _speeches()
    except Exception:
        speeches = []
    try:
        tabled_wqs = _wqs()
    except Exception:
        tabled_wqs = []

    return {'speeches': speeches, 'tabled_wqs': tabled_wqs}


# ==========================================
# ROUTE 11: DEBATE PREP
# ==========================================
@debate_scanner_bp.route('/debate_prep', methods=['GET', 'POST'])
def debate_prep():
    import logging, traceback
    if request.method == 'GET':
        return render_template('debate_prep.html',
            is_post=False, peer_name='', question_date='', question_text='',
            media_start='', media_end='', house='lords', peer_info=None, one_pager=None,
            media_items=[], parl_sections={}, peer_contributions={}, error=None)

    try:
        return _debate_prep_post()
    except Exception as e:
        tb = traceback.format_exc()
        logging.error(f"[debate_prep POST] UNHANDLED: {e}\n{tb}")
        _h = request.form.get('house', 'lords')
        return render_template('debate_prep.html',
            is_post=True, peer_name=request.form.get('peer_name',''),
            question_date=request.form.get('question_date',''),
            question_text=request.form.get('question_text',''),
            media_start='', media_end='', house=_h, peer_info=None, one_pager=None,
            media_items=[], parl_sections={}, peer_contributions={},
            error=f'Server error: {e}'), 200


def _debate_prep_post():
    peer_name = request.form.get('peer_name', '').strip()
    question_date = request.form.get('question_date', '').strip()
    question_text = request.form.get('question_text', '').strip()
    media_start = request.form.get('media_start', '').strip()
    media_end = request.form.get('media_end', '').strip()
    house = request.form.get('house', 'lords').strip().lower()
    if house not in ('lords', 'commons'):
        house = 'lords'

    if not peer_name or not question_text:
        return render_template('debate_prep.html',
            is_post=True, peer_name=peer_name, question_date=question_date,
            question_text=question_text, media_start=media_start, media_end=media_end,
            house=house, peer_info=None, one_pager=None, media_items=[], parl_sections={},
            peer_contributions={}, error='Please provide a member name and question text.')

    # Derive topic from question text — use first sentence as search topic
    topic = re.split(r'[.?]', question_text)[0].strip()[:150]

    # Default media range: 60 days before question date
    if not media_start and question_date:
        try:
            from datetime import timedelta
            qd = datetime.strptime(question_date, '%Y-%m-%d')
            media_start = (qd - timedelta(days=60)).strftime('%Y-%m-%d')
            media_end = media_end or question_date
        except Exception:
            pass

    # Parliamentary range: 1 year look-back
    parl_start = ''
    parl_end = question_date or datetime.now().strftime('%Y-%m-%d')
    if question_date:
        try:
            from datetime import timedelta
            qd = datetime.strptime(question_date, '%Y-%m-%d')
            parl_start = (qd - timedelta(days=365)).strftime('%Y-%m-%d')
        except Exception:
            pass

    parl_date_range = get_twfy_date_range(parl_start, parl_end)

    peer_info = None
    one_pager = None
    media_items = []
    peer_contributions = {}

    # All tasks submitted to a single flat executor — no nested thread pools
    with concurrent.futures.ThreadPoolExecutor(max_workers=8) as executor:
        f_peer    = executor.submit(copy_current_request_context(_prep_resolve_peer), peer_name)
        f_op      = executor.submit(copy_current_request_context(_prep_one_pager), question_text, topic, question_date, house)
        f_media   = executor.submit(copy_current_request_context(_prep_media), topic, media_start, media_end)
        f_lords   = executor.submit(copy_current_request_context(_prep_parl_lords), topic, parl_date_range)
        f_commons = executor.submit(copy_current_request_context(_prep_parl_commons), topic, parl_date_range)
        f_wqs     = executor.submit(copy_current_request_context(_prep_parl_wqs), topic, parl_start, parl_end)
        f_stmts   = executor.submit(copy_current_request_context(_prep_parl_statements), topic, parl_date_range)

        try:
            peer_info = f_peer.result(timeout=15)
        except Exception:
            peer_info = None
        try:
            one_pager = f_op.result(timeout=35)
        except Exception:
            one_pager = None
        try:
            media_items = f_media.result(timeout=20)
        except Exception:
            media_items = []
        try:
            lords_debates = f_lords.result(timeout=25)
        except Exception:
            lords_debates = []
        try:
            commons_debates = f_commons.result(timeout=25)
        except Exception:
            commons_debates = []
        try:
            parl_wqs = f_wqs.result(timeout=25)
        except Exception:
            parl_wqs = []
        try:
            statements = f_stmts.result(timeout=25)
        except Exception:
            statements = []

    parl_sections = {
        'lords_debates': lords_debates,
        'commons_debates': commons_debates,
        'wqs': parl_wqs,
        'statements': statements,
    }

    # Peer contributions requires parliament_id — run sequentially after peer resolves
    if peer_info and peer_info.get('parliament_id'):
        try:
            peer_contributions = _prep_peer_contributions(
                peer_info['parliament_id'], parl_start, parl_end
            )
        except Exception:
            peer_contributions = {}

    return render_template('debate_prep.html',
        is_post=True,
        peer_name=peer_name,
        question_date=question_date,
        question_text=question_text,
        media_start=media_start,
        media_end=media_end,
        house=house,
        peer_info=peer_info,
        one_pager=one_pager,
        media_items=media_items or [],
        parl_sections=parl_sections or {},
        peer_contributions=peer_contributions or {},
        error=None)


# ==========================================
# ROUTE 12: DOWNLOAD DEBATE PREP BRIEF
# ==========================================
@debate_scanner_bp.route('/download_debate_prep_brief', methods=['POST'])
def download_debate_prep_brief():
    if not Document:
        return "Word library missing.", 500

    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    def _parse_field(field, default=None):
        val = request.form.get(field, '')
        if not val:
            return default if default is not None else {}
        try:
            return json.loads(val)
        except Exception:
            return default if default is not None else {}

    peer_name = request.form.get('peer_name', 'Unknown Peer')
    question_date = request.form.get('question_date', '')
    question_text = request.form.get('question_text', '')
    house = request.form.get('house', 'lords')
    member_label = 'MP' if house == 'commons' else 'Peer'
    peer_info = _parse_field('peer_info_json', {})
    one_pager = _parse_field('one_pager_json', {})
    media_items = _parse_field('media_json', [])
    parl_sections = _parse_field('parl_json', {})
    peer_contributions = _parse_field('peer_contrib_json', {})

    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.1)
    sec.right_margin = Inches(1.1)

    def _set_cell_bg(cell, hex_colour):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = _OxmlElement('w:shd')
        shd.set(_qn('w:val'), 'clear')
        shd.set(_qn('w:color'), 'auto')
        shd.set(_qn('w:fill'), hex_colour)
        tcPr.append(shd)

    def _clr(run, hex_str):
        r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)

    # Cover block
    title_p = doc.add_heading('Debate Preparation Briefing', 0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in filter(None, [
        f'{member_label}: {peer_name}',
        f'Question date: {question_date}' if question_date else None,
        f'Generated: {datetime.now().strftime("%d %b %Y %H:%M")}',
    ]):
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(11)

    if question_text:
        q_p = doc.add_paragraph()
        q_p.add_run('Question: ').bold = True
        q_p.add_run(question_text)
        q_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    notice = doc.add_paragraph(
        '⚠ AI-generated — verify against source links before use in official documents.')
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice.runs[0].font.size = Pt(9)
    notice.runs[0].italic = True
    _clr(notice.runs[0], '856404')
    doc.add_paragraph()

    # ── Section 1: One Pager ──────────────────────────────────────
    if one_pager:
        doc.add_heading('One Pager — Background Briefing', 1)
        for sub_key, sub_title in [
            ('why_now', 'Why Now'),
            ('sector_context', 'Sector Context'),
            ('opposition_position', 'Opposition Position'),
        ]:
            if one_pager.get(sub_key):
                doc.add_heading(sub_title, 2)
                doc.add_paragraph(one_pager[sub_key])
        if one_pager.get('major_criticisms'):
            doc.add_heading('Major Criticisms', 2)
            for crit in one_pager['major_criticisms']:
                doc.add_paragraph(crit, style='List Bullet')
        if one_pager.get('sources'):
            doc.add_heading('Suggested Sources', 2)
            for src in one_pager['sources']:
                doc.add_paragraph(src, style='List Bullet')
            p = doc.add_paragraph('AI-suggested — verify existence before citing')
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(9)
            _clr(p.runs[0], '888888')
        doc.add_paragraph()

    # ── Section 2: Media ──────────────────────────────────────────
    if media_items:
        doc.add_heading(f'Media Coverage ({len(media_items)} articles)', 1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = 'Table Grid'
        hdr = tbl.rows[0].cells
        for i, lbl in enumerate(['Date', 'Outlet', 'Sentiment', 'Headline & Summary']):
            hdr[i].text = lbl
            hdr[i].paragraphs[0].runs[0].bold = True
            _set_cell_bg(hdr[i], '1C3E6E')
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        sent_colours = {'positive': 'D4EDDA', 'negative': 'F8D7DA', 'neutral': 'F8F9FA'}
        for art in media_items:
            row = tbl.add_row().cells
            row[0].text = art.get('published', '')
            row[1].text = art.get('outlet', '')
            sentiment = art.get('sentiment', 'neutral')
            row[2].text = sentiment.capitalize()
            _set_cell_bg(row[2], sent_colours.get(sentiment, 'F8F9FA'))
            cell = row[3]
            p = cell.paragraphs[0]
            link = art.get('link', '')
            if link:
                _add_hyperlink(p, link, art.get('title', ''))
            else:
                p.add_run(art.get('title', ''))
            summary = art.get('summary', '')
            if summary and summary != art.get('title', ''):
                sp = cell.add_paragraph(summary)
                sp.runs[0].font.size = Pt(9)
        doc.add_paragraph()

    # ── Section 3: Parliamentary Record ───────────────────────────
    parl_label_map = [
        ('lords_debates', 'Lords Debates'),
        ('commons_debates', 'Commons Debates'),
        ('wqs', 'Written Questions'),
        ('statements', 'Oral Statements'),
    ]
    if any(parl_sections.get(k) for k, _ in parl_label_map):
        doc.add_heading('Parliamentary Record', 1)
        for key, label in parl_label_map:
            rows = parl_sections.get(key, [])
            if not rows:
                continue
            doc.add_heading(f'{label} ({len(rows)})', 2)
            if key == 'wqs':
                for q in rows[:15]:
                    p = doc.add_paragraph()
                    meta_run = p.add_run(
                        f"{q.get('date_tabled', '')}  ·  {q.get('dept', '')}  — ")
                    meta_run.font.size = Pt(9)
                    _clr(meta_run, '555555')
                    p.add_run(q.get('question_text', '')[:200])
                    if q.get('url'):
                        lp = doc.add_paragraph()
                        _add_hyperlink(lp, q['url'], '↗ Parliament.uk')
                        if lp.runs:
                            lp.runs[0].font.size = Pt(9)
            else:
                seen_d = set()
                for r in rows[:20]:
                    url = r.get('listurl', '')
                    if url and not url.startswith('http'):
                        url = 'https://www.theyworkforyou.com' + url
                    key_d = url or r.get('debate_title', '')
                    if key_d in seen_d:
                        continue
                    seen_d.add(key_d)
                    hp = doc.add_paragraph()
                    meta_run = hp.add_run(
                        f"{r.get('hdate', '')}  ·  {r.get('source_label', '')}  — ")
                    meta_run.font.size = Pt(9)
                    _clr(meta_run, '555555')
                    hp.add_run(r.get('debate_title', '')).bold = True
                    sp = doc.add_paragraph()
                    sp.add_run(f"{r.get('speaker_name', '')}:  ").bold = True
                    body_run = sp.add_run(r.get('body_clean', '')[:250])
                    body_run.font.size = Pt(10)
                    if url:
                        lp = doc.add_paragraph()
                        _add_hyperlink(lp, url, '↗ View on Hansard')
                        if lp.runs:
                            lp.runs[0].font.size = Pt(9)
                    doc.add_paragraph('─' * 60)
        doc.add_paragraph()

    # ── Section 4: Member Profile ─────────────────────────────────
    doc.add_heading(f'{member_label} Profile — {peer_name}', 1)
    if peer_info:
        meta = '  ·  '.join(filter(None, [peer_info.get('party', ''),
                                           peer_info.get('house', '')]))
        if meta:
            doc.add_paragraph(meta).runs[0].font.size = Pt(10)

        bio_posts = peer_info.get('biography_posts', [])
        if bio_posts:
            doc.add_heading('Career & Posts', 2)
            current = [p for p in bio_posts if p.get('is_current')]
            past = [p for p in bio_posts if not p.get('is_current')]
            for post in current[:5]:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"[Current] {post['category']}: {post['title']}")
                run.bold = True
                _clr(run, '1C3E6E')
            for post in past[:12]:
                p = doc.add_paragraph(style='List Bullet')
                yr = ''
                if post.get('start'):
                    yr = f" ({post['start'][:4]}"
                    yr += f"–{post['end'][:4]})" if post.get('end') else '–present)'
                p.add_run(f"{post['category']}: {post['title']}{yr}")

        interests = peer_info.get('interests', [])
        if interests:
            doc.add_heading('Register of Interests', 2)
            for cat in interests[:6]:
                doc.add_paragraph(cat['category']).runs[0].bold = True
                for item in cat['items'][:5]:
                    p = doc.add_paragraph(item, style='List Bullet')
                    p.runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(f'{member_label} not found in Parliament Members API.')

    speeches = (peer_contributions or {}).get('speeches', [])
    if speeches:
        doc.add_heading('Recent Spoken Contributions', 2)
        seen_s = set()
        for r in speeches[:12]:
            url = r.get('listurl', '')
            if url and not url.startswith('http'):
                url = 'https://www.theyworkforyou.com' + url
            key_s = url or r.get('debate_title', '')
            if key_s in seen_s:
                continue
            seen_s.add(key_s)
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{r.get('hdate', '')}  —  {r.get('debate_title', '')}")
            if url:
                lp = doc.add_paragraph()
                _add_hyperlink(lp, url, '↗ Hansard')
                if lp.runs:
                    lp.runs[0].font.size = Pt(9)

    tabled_wqs = (peer_contributions or {}).get('tabled_wqs', [])
    if tabled_wqs:
        doc.add_heading('Recent Written Questions Tabled', 2)
        for q in tabled_wqs[:12]:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(
                f"{q.get('date_tabled', '')}  ·  {q.get('dept', '')}  —  "
                f"{q.get('question_text', '')[:150]}")
            if q.get('url'):
                lp = doc.add_paragraph()
                _add_hyperlink(lp, q['url'], '↗ Parliament.uk')
                if lp.runs:
                    lp.runs[0].font.size = Pt(9)

    mem_doc = io.BytesIO()
    doc.save(mem_doc)
    mem_doc.seek(0)
    safe_name = re.sub(r'[^\w\s-]', '', peer_name)[:30].strip()
    filename = f"Debate Prep - {safe_name} - {datetime.now().strftime('%Y%m%d')}.docx"
    return send_file(mem_doc, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')