"""
Hansard Archive — public-facing routes (Phase 2A Week 3).

Blueprint: archive_bp, url_prefix=/archive

Routes:
  GET /archive                               browse/home — recent sessions + filters
  GET /archive/search                        FTS text search (noindex)
  GET /archive/debate/<date>/<slug>          session detail + full transcript
  GET /archive/debate/<date>/<slug>/word     Word export of session transcript
  GET /archive/date/<date_str>               all sessions on a date
  GET /archive/mp/<member_id>               sessions a member contributed to
  GET /archive/department/<slug>             sessions by answering department
  GET /archive/policy/<slug>                 sessions by GOV.UK policy area tag
  GET /archive/theme/<slug>                  sessions by specific theme tag
  GET /archive/pq/<uin>                      Written Question detail page

URL conventions (locked — do not change after indexing):
  date format: 22-july-2025  (day no leading zero, lowercase full month, 4-digit year)
  slug format: {title-slug}-{4-char-hex}
"""

import os
import re
import time
from collections import defaultdict
from datetime import date as date_type, timedelta
from io import BytesIO
from urllib.parse import urlencode

from docx import Document
from docx.shared import Pt, RGBColor
from flask import Blueprint, abort, redirect, render_template, make_response, request
from markupsafe import Markup
from sqlalchemy import func, text as sqla_text

from extensions import db
from hansard_archive.models import (
    HansardContribution,
    HansardSession,
    HansardSessionTheme,
    HaPQ,
    HaPQTheme,
    THEME_TYPE_POLICY_AREA,
    THEME_TYPE_SPECIFIC,
)

archive_bp = Blueprint("archive", __name__, url_prefix="/archive")

# ---------------------------------------------------------------------------
# Party colours + attribution parsing
# ---------------------------------------------------------------------------

_PARTY_COLOURS: dict[str, str] = {
    "Lab":        "#E4003B",
    "Lab/Co-op":  "#E4003B",
    "Lab/ Co-op": "#E4003B",
    "Lab Co-op":  "#E4003B",
    "Con":        "#0087DC",
    "LD":         "#FAA61A",
    "SNP":        "#c9a800",
    "PC":         "#005B54",
    "Green":      "#02A95B",
    "Reform":     "#12B6CF",
    "DUP":        "#CF1F25",
    "SDLP":       "#2AA82C",
    "UUP":        "#48A5EE",
    "Alliance":   "#c9960a",
    "TUV":        "#0C3A6A",
    "CB":         "#7a8a9a",
    "Ind":        "#7a8a9a",
    "Non-Afl":    "#7a8a9a",
}
_DEFAULT_PARTY_COLOUR = "#1a4a6e"

_NOT_PARTY = frozenset({
    "Maiden Speech", "Valedictory Speech", "Urgent Question", "Maiden",
    "Your Party", "Restore Britain",
})
_STRIP_PREFIXES = frozenset({"Mr", "Mrs", "Ms", "Miss", "Dr"})


def _parse_attribution(raw: str | None) -> dict:
    """
    Parse raw member_name into display-ready fields.

    Handles:
      "Name (Constituency) (Party)"  — standard Commons
      "Name (Party)"                 — standard Lords
      "The [Role] (Name)"            — ministerial proxy format
      "Name"                         — short reference form
    """
    if not raw:
        return {"name": "Speaker", "role": "", "party": None,
                "party_colour": _DEFAULT_PARTY_COLOUR}

    raw = raw.strip()

    # Ministerial proxy: "The Secretary of State for X (Name)"
    m = re.match(r'^(The\s+[^(]+?)\s*\(([^)]+)\)\s*$', raw)
    if m:
        role = m.group(1).strip()
        name = m.group(2).strip()
        return {"name": name, "role": role, "party": None,
                "party_colour": _DEFAULT_PARTY_COLOUR}

    # Standard: extract parenthetical groups
    groups = re.findall(r'\(([^)]+)\)', raw)
    base = re.sub(r'\s*\([^)]+\)', '', raw).strip()

    # Strip leading salutation (Mr/Mrs etc — not Lord/Baroness which are titles)
    parts = base.split(None, 1)
    if parts and parts[0] in _STRIP_PREFIXES:
        base = parts[1] if len(parts) > 1 else base

    party = None
    if len(groups) >= 2:
        candidate = groups[-1].strip()
        if candidate not in _NOT_PARTY:
            party = candidate

    colour = _PARTY_COLOURS.get(party, _DEFAULT_PARTY_COLOUR) if party else _DEFAULT_PARTY_COLOUR
    return {"name": base or raw, "role": "", "party": party, "party_colour": colour}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_MONTH_NAMES = [
    "", "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]

_DEBATE_TYPE_LABELS = {
    "oral_questions":       "Oral Questions",
    "pmqs":                 "Prime Minister's Questions",
    "westminster_hall":     "Westminster Hall",
    "debate":               "Debate",
    "ministerial_statement":"Ministerial Statement",
    "statutory_instrument": "Statutory Instrument",
    "committee_stage":      "Committee Stage",
    "petition":             "Petition",
    "other":                "Proceedings",
}


def _human_date(d) -> str:
    """date → '27 April 2026'"""
    return f"{d.day} {_MONTH_NAMES[d.month]} {d.year}"


def _url_date(d) -> str:
    """date → '27-april-2026'"""
    return f"{d.day}-{_MONTH_NAMES[d.month].lower()}-{d.year}"


def _parse_url_date(date_str: str) -> date_type | None:
    """Parse '27-april-2026' → date object, or None if invalid."""
    parts = date_str.split('-')
    if len(parts) != 3:
        return None
    day_str, month_str, year_str = parts
    try:
        day = int(day_str)
        year = int(year_str)
        month_idx = next(
            (i for i, m in enumerate(_MONTH_NAMES) if m.lower() == month_str.lower()),
            0,
        )
        if month_idx == 0:
            return None
        return date_type(year, month_idx, day)
    except (ValueError, IndexError):
        return None


def _slugify(s: str) -> str:
    """Convert a label to a URL slug (matches sitemap slugification)."""
    s = s.lower()
    s = re.sub(r"[^a-z0-9\s-]", "", s)
    s = re.sub(r"\s+", "-", s.strip())
    s = re.sub(r"-+", "-", s)
    return s


def _is_postgres() -> bool:
    """True when connected to Postgres (FTS available), False for SQLite."""
    try:
        return db.engine.dialect.name == "postgresql"
    except Exception:
        return False


def _build_session_items(sessions: list, contrib_counts: dict,
                         policy_areas: dict, specific_topics: dict) -> list:
    """Build template-ready item dicts from a list of HansardSession objects."""
    return [
        {
            "session":           s,
            "human_date":        _human_date(s.date),
            "url_date":          _url_date(s.date),
            "debate_type_label": _DEBATE_TYPE_LABELS.get(s.debate_type, "Proceedings"),
            "contrib_count":     contrib_counts.get(s.id, 0),
            "policy_areas":      sorted(policy_areas.get(s.id, [])),
            "specific_topics":   sorted(specific_topics.get(s.id, [])),
            "department":        s.department or "",
        }
        for s in sessions
    ]


def _batch_load_tags(session_ids: list) -> tuple[dict, dict]:
    """Batch-load policy area and specific topic tags for a list of session IDs."""
    policy: dict[int, list[str]] = defaultdict(list)
    specific: dict[int, list[str]] = defaultdict(list)
    if session_ids:
        for t in (
            db.session.query(HansardSessionTheme)
            .filter(HansardSessionTheme.session_id.in_(session_ids))
            .all()
        ):
            if t.theme_type == THEME_TYPE_POLICY_AREA:
                policy[t.session_id].append(t.theme)
            elif t.theme_type == THEME_TYPE_SPECIFIC:
                specific[t.session_id].append(t.theme)
    return policy, specific


def _batch_load_contrib_counts(session_ids: list) -> dict:
    """Batch-load speaker contribution counts for a list of session IDs."""
    if not session_ids:
        return {}
    return dict(
        db.session.query(
            HansardContribution.session_id,
            func.count(HansardContribution.id),
        )
        .filter(
            HansardContribution.session_id.in_(session_ids),
            HansardContribution.member_name.isnot(None),
        )
        .group_by(HansardContribution.session_id)
        .all()
    )


def _session_or_404(slug: str) -> HansardSession:
    """Fetch a non-container session by slug or 404."""
    session = (
        HansardSession.query
        .filter_by(slug=slug, is_container=False)
        .first()
    )
    if session is None:
        abort(404)
    return session


# ---------------------------------------------------------------------------
# Related sessions
# ---------------------------------------------------------------------------

def _normalise_title(title: str) -> str:
    """
    Normalise a session title for related-session matching.

    Handles two legislative patterns:
      - SI pattern: strips 'Draft ' prefix
        ("Draft Warm Home Discount Regulations" → "Warm Home Discount Regulations")
      - Lords Bill pattern: strips '[Lords]' suffix
        ("Finance Bill [Lords]" → "Finance Bill")
      - Whitespace: collapses double-spaces (seen in some Hansard API titles)
    """
    t = title.strip()
    t = re.sub(r'^\s*Draft\s+', '', t, flags=re.IGNORECASE)
    t = re.sub(r'\s*\[Lords\]\s*$', '', t, flags=re.IGNORECASE)
    t = re.sub(r'\s+', ' ', t).strip()
    return t.lower()


# Procedural titles that recur constantly with no relationship between instances.
# A denylist is more precise than a frequency threshold: high-frequency Bills
# (Crime and Policing Bill: 48 sessions) are genuine related stages and must
# NOT be suppressed. Only titles that are the same event repeated — not stages
# of one legislative item — belong here.
_PROCEDURAL_TITLE_NOISE: frozenset[str] = frozenset({
    "Topical Questions",
    "Arrangement of Business",
    "Points of Order",
    "Point of Order",
    "Business of the House",
    "Engagements",
    "Speaker's Statement",
    "Speaker’s Statement",   # curly apostrophe variant
    "Business without Debate",
    "Retirements of Members",
    "Retirement of a Member",
    "Oral Questions",
    "Written Statements",
})

_RELATED_PANEL_MAX = 6   # max cards shown inline; Bills with more stages get a count badge


def _is_procedural_noise(title: str) -> bool:
    return title.strip() in _PROCEDURAL_TITLE_NOISE


def _related_sessions(session: HansardSession) -> dict:
    """
    Find sessions related to the given one by normalised title match.

    Matching strategy (exact normalised title only — no fuzzy matching):
      - Exact match covers: Lords SI GC→Chamber (identical title), cross-house
        Bill stages where title is unchanged, topical recurrences.
      - 'Draft X' normalisation covers: SI draft laid before passage vs
        approved instrument (SI-specific pattern).
      - '[Lords]' normalisation covers: Lords-originated Bills whose Commons
        title drops the '[Lords]' suffix.

    Window: 365 days — Bill stages span Commons Second Reading (Oct) to Lords
    Third Reading (Mar); SIs typically resolve within days.

    Noise guard: denylist of known procedural titles that recur constantly with
    no relationship between instances (PMQs 'Engagements', 'Business of the
    House', etc.). Bills with 48+ stages are NOT noise — they're genuine related
    stages and must not be suppressed by a frequency threshold.

    Returns dict:
      items      — up to _RELATED_PANEL_MAX cards, sorted by date proximity
      total      — total related sessions found (for overflow badge)
    """
    if _is_procedural_noise(session.title):
        return {"items": [], "total": 0}

    norm = _normalise_title(session.title)
    if not norm:
        return {"items": [], "total": 0}

    date_min = session.date - timedelta(days=365)
    date_max = session.date + timedelta(days=365)

    # Broad DB filter using title prefix; Python post-filters to exact normalised match.
    prefix = norm[:20]
    candidates = (
        HansardSession.query
        .filter(
            HansardSession.is_container == False,
            HansardSession.id != session.id,
            HansardSession.date >= date_min,
            HansardSession.date <= date_max,
            HansardSession.title.ilike(f"%{prefix}%"),
        )
        .order_by(HansardSession.date)
        .limit(60)
        .all()
    )

    related = [
        c for c in candidates
        if _normalise_title(c.title) == norm
        and not _is_procedural_noise(c.title)
    ]

    if not related:
        return {"items": [], "total": 0}

    # Sort by date proximity, then Commons before Lords within the same date.
    related.sort(key=lambda c: (abs((c.date - session.date).days), 0 if c.house == "Commons" else 1))
    total = len(related)
    display = related[:_RELATED_PANEL_MAX]

    # Batch-load contrib counts (avoids N+1)
    display_ids = [c.id for c in display]
    counts = dict(
        db.session.query(
            HansardContribution.session_id,
            func.count(HansardContribution.id),
        )
        .filter(
            HansardContribution.session_id.in_(display_ids),
            HansardContribution.member_name.isnot(None),
        )
        .group_by(HansardContribution.session_id)
        .all()
    )

    items = [
        {
            "session":           c,
            "human_date":        _human_date(c.date),
            "url_date":          _url_date(c.date),
            "debate_type_label": _DEBATE_TYPE_LABELS.get(c.debate_type, "Proceedings"),
            "contrib_count":     counts.get(c.id, 0),
        }
        for c in display
    ]

    return {"items": items, "total": total}


def _day_navigation(session: HansardSession) -> dict:
    """
    Return the previous and next non-container session on the same date and
    in the same house, ordered by DB insertion id (which approximates chain
    order from the BFS walk).
    """
    base = dict(
        date=session.date,
        house=session.house,
        is_container=False,
    )
    prev_s = (
        HansardSession.query
        .filter_by(**base)
        .filter(HansardSession.id < session.id)
        .order_by(HansardSession.id.desc())
        .first()
    )
    next_s = (
        HansardSession.query
        .filter_by(**base)
        .filter(HansardSession.id > session.id)
        .order_by(HansardSession.id.asc())
        .first()
    )

    def _nav_item(s):
        return {
            "title": s.title,
            "url":   f"/archive/debate/{_url_date(s.date)}/{s.slug}",
            "dtype": _DEBATE_TYPE_LABELS.get(s.debate_type, "Proceedings"),
        }

    return {
        "prev": _nav_item(prev_s) if prev_s else None,
        "next": _nav_item(next_s) if next_s else None,
    }


def _session_context(session: HansardSession) -> dict:
    """Build common template context for a session."""
    raw_contribs = (
        session.contributions
        .filter(HansardContribution.member_name.isnot(None))
        .order_by(HansardContribution.speech_order)
        .all()
    )
    themes = session.themes.all()
    policy_areas = sorted(
        {t.theme for t in themes if t.theme_type == THEME_TYPE_POLICY_AREA}
    )
    specific_topics = sorted(
        {t.theme for t in themes if t.theme_type == THEME_TYPE_SPECIFIC}
    )

    contributions = []
    for c in raw_contribs:
        attr = _parse_attribution(c.member_name)
        party_name = attr["party"] or (c.party if c.party else None)
        contributions.append({
            "id":           c.id,
            "member_id":    c.member_id,
            "raw_name":     c.member_name,
            "name":         attr["name"],
            "role":         attr["role"],
            "party":        party_name,
            "party_colour": _PARTY_COLOURS.get(party_name, _DEFAULT_PARTY_COLOUR) if party_name else _DEFAULT_PARTY_COLOUR,
            "speech_text":  c.speech_text or "",
        })

    return {
        "session":           session,
        "contributions":     contributions,
        "policy_areas":      policy_areas,
        "specific_topics":   specific_topics,
        "human_date":        _human_date(session.date),
        "url_date":          _url_date(session.date),
        "debate_type_label": _DEBATE_TYPE_LABELS.get(session.debate_type, "Proceedings"),
        "department":        session.department or "",
        "related_sessions":  _related_sessions(session),
        "day_nav":           _day_navigation(session),
    }


# ---------------------------------------------------------------------------
# Archive home — search and browse
# ---------------------------------------------------------------------------

_PER_PAGE       = 25
_GROUPS_PER_PAGE = 7   # groups shown per page in grouped OQ / PMQs view


def _build_oq_group_header(dtype_filter: str, date, house: str, dept: str) -> str:
    ds = _human_date(date)
    if dtype_filter == "pmqs":
        return f"{ds} — Prime Minister's Questions"
    if house == "Lords":
        return f"{ds} — Lords Oral Questions"
    if dept:
        return f"{ds} — {dept} Oral Questions"
    return f"{ds} — Oral Questions"


def _last_ingested_label() -> str:
    """Human-readable label for the most recently ingested session date."""
    row = (
        db.session.query(func.max(HansardSession.date))
        .filter(HansardSession.is_container == False)
        .scalar()
    )
    if not row:
        return ""
    return _human_date(row)


_VOCAB_CACHE_TTL  = 300   # 5 minutes — refreshed after each ingest cycle
_RECENT_CACHE_TTL = 900   # 15 minutes — recent-additions widget
_policy_area_cache: tuple[list, float] | None = None
_dept_cache: tuple[list, float] | None = None
_recent_additions_cache: tuple[dict, float] | None = None


def _all_policy_areas() -> list[str]:
    """Sorted list of all policy area terms present in the corpus (cached 5 min)."""
    global _policy_area_cache
    now = time.monotonic()
    if _policy_area_cache and now - _policy_area_cache[1] < _VOCAB_CACHE_TTL:
        return _policy_area_cache[0]
    result = sorted(
        r[0] for r in db.session.query(
            func.distinct(HansardSessionTheme.theme)
        ).filter(HansardSessionTheme.theme_type == THEME_TYPE_POLICY_AREA).all()
    )
    _policy_area_cache = (result, now)
    return result


def _all_departments() -> list[str]:
    """Sorted list of departments that have attributed oral questions sessions (cached 5 min)."""
    global _dept_cache
    now = time.monotonic()
    if _dept_cache and now - _dept_cache[1] < _VOCAB_CACHE_TTL:
        return _dept_cache[0]
    result = sorted(
        r[0] for r in db.session.query(
            func.distinct(HansardSession.department)
        ).filter(HansardSession.department.isnot(None)).all()
    )
    _dept_cache = (result, now)
    return result


def _recent_additions() -> dict:
    """
    Count of non-container sessions ingested in the last 24 hours and the
    most recent ingestion timestamp. Cached for 15 minutes — no need for
    real-time freshness on the browse page.
    """
    global _recent_additions_cache
    now = time.monotonic()
    if _recent_additions_cache and now - _recent_additions_cache[1] < _RECENT_CACHE_TTL:
        return _recent_additions_cache[0]

    from datetime import datetime as dt, timedelta
    cutoff = dt.utcnow() - timedelta(hours=24)

    count_24h = (
        db.session.query(func.count(HansardSession.id))
        .filter(
            HansardSession.is_container == False,
            HansardSession.ingested_at >= cutoff,
        )
        .scalar()
    ) or 0

    last_at = db.session.query(func.max(HansardSession.ingested_at)).scalar()

    result = {"count_24h": count_24h, "last_ingested_at": last_at}
    _recent_additions_cache = (result, now)
    return result


@archive_bp.route("")
def archive_home():
    q = request.args.get("q", "").strip()
    # Text search lives at /archive/search (noindex); browse stays at /archive (indexed).
    if q:
        fwd: dict = {"q": q}
        for k in ("policy", "house", "dtype", "dept", "from", "to"):
            v = request.args.get(k, "").strip()
            if v:
                fwd[k] = v
        if request.args.get("title_only") == "1":
            fwd["title_only"] = "1"
        return redirect(f"/archive/search?{urlencode(fwd)}", code=302)

    house_filter  = request.args.get("house", "")
    dtype_filter  = request.args.get("dtype", "")
    dept_filter   = "" if house_filter == "Lords" else request.args.get("dept", "")
    policy_filter = request.args.get("policy", "")
    date_from     = request.args.get("from", "")
    date_to       = request.args.get("to", "")
    title_only    = request.args.get("title_only") == "1"
    try:
        page = max(1, int(request.args.get("page", 1)))
    except (ValueError, TypeError):
        page = 1

    stmt = HansardSession.query.filter_by(is_container=False)

    if house_filter in ("Commons", "Lords"):
        stmt = stmt.filter(HansardSession.house == house_filter)

    if dtype_filter and dtype_filter in _DEBATE_TYPE_LABELS:
        stmt = stmt.filter(HansardSession.debate_type == dtype_filter)

    if dept_filter:
        stmt = stmt.filter(HansardSession.department == dept_filter)

    if policy_filter:
        policy_sub = (
            db.session.query(HansardSessionTheme.session_id)
            .filter(
                HansardSessionTheme.theme == policy_filter,
                HansardSessionTheme.theme_type == THEME_TYPE_POLICY_AREA,
            )
            .subquery()
        )
        stmt = stmt.filter(HansardSession.id.in_(policy_sub))

    if date_from:
        try:
            stmt = stmt.filter(HansardSession.date >= date_type.fromisoformat(date_from))
        except ValueError:
            date_from = ""

    if date_to:
        try:
            stmt = stmt.filter(HansardSession.date <= date_type.fromisoformat(date_to))
        except ValueError:
            date_to = ""

    # Build base query string and flags (needed in both grouped and flat paths)
    filter_params = {k: v for k, v in {
        "house":      house_filter,
        "dtype":      dtype_filter,
        "dept":       dept_filter,
        "policy":     policy_filter,
        "from":       date_from,
        "to":         date_to,
        "title_only": "1" if title_only else "",
    }.items() if v}
    base_qs    = urlencode(filter_params)
    has_filters = bool(house_filter or dtype_filter or dept_filter or policy_filter or date_from or date_to)

    # --- Grouped view: oral_questions or pmqs ---
    grouped = dtype_filter in ("oral_questions", "pmqs")

    if grouped:
        # Lightweight pass: fetch only grouping fields for all matching sessions
        lite_rows = (
            stmt
            .with_entities(
                HansardSession.id,
                HansardSession.date,
                HansardSession.house,
                HansardSession.department,
            )
            .order_by(HansardSession.date.desc(), HansardSession.id)
            .all()
        )

        groups_dict: dict[tuple, list[int]] = defaultdict(list)
        for sid, row_date, row_house, row_dept in lite_rows:
            key = (row_date, row_house, row_dept or "")
            groups_dict[key].append(sid)

        # Sort groups: newest date first; within same date Commons before Lords
        sorted_keys = sorted(
            groups_dict,
            key=lambda k: (-k[0].toordinal(), 0 if k[1] == "Commons" else 1),
        )
        total_groups   = len(sorted_keys)
        total_sessions = sum(len(v) for v in groups_dict.values())
        total_pages    = max(1, (total_groups + _GROUPS_PER_PAGE - 1) // _GROUPS_PER_PAGE)
        page_keys      = sorted_keys[(page - 1) * _GROUPS_PER_PAGE : page * _GROUPS_PER_PAGE]

        page_session_ids = [sid for key in page_keys for sid in groups_dict[key]]

        sessions_full  = (
            HansardSession.query.filter(HansardSession.id.in_(page_session_ids)).all()
            if page_session_ids else []
        )
        sessions_by_id = {s.id: s for s in sessions_full}

        g_contrib_counts = _batch_load_contrib_counts(page_session_ids)
        g_policy_areas, g_specific_topics = _batch_load_tags(page_session_ids)

        def _make_item(s):
            return {
                "session":           s,
                "human_date":        _human_date(s.date),
                "url_date":          _url_date(s.date),
                "debate_type_label": _DEBATE_TYPE_LABELS.get(s.debate_type, "Proceedings"),
                "contrib_count":     g_contrib_counts.get(s.id, 0),
                "policy_areas":      sorted(g_policy_areas.get(s.id, [])),
                "specific_topics":   sorted(g_specific_topics.get(s.id, [])),
                "department":        s.department or "",
            }

        group_items = []
        for key in page_keys:
            row_date, row_house, row_dept = key
            sids = sorted(groups_dict[key])  # session ID order = Hansard chain order
            group_sessions = [sessions_by_id[sid] for sid in sids if sid in sessions_by_id]
            group_items.append({
                "date":          row_date,
                "house":         row_house,
                "dept":          row_dept,
                "header":        _build_oq_group_header(dtype_filter, row_date, row_house, row_dept),
                "session_count": len(sids),
                "sessions":      [_make_item(s) for s in group_sessions],
            })

        return render_template(
            "hansard_archive/archive_home.html",
            grouped=True,
            group_items=group_items,
            total_groups=total_groups,
            total_sessions=total_sessions,
            session_items=[],
            total=total_groups,
            total_pages=total_pages,
            per_page=_GROUPS_PER_PAGE,
            page=page,
            q="",
            house_filter=house_filter,
            dtype_filter=dtype_filter,
            dept_filter=dept_filter,
            policy_filter=policy_filter,
            date_from=date_from,
            date_to=date_to,
            title_only=title_only,
            all_policy_areas=_all_policy_areas(),
            all_departments=_all_departments(),
            debate_type_labels=_DEBATE_TYPE_LABELS,
            base_qs=base_qs,
            has_filters=has_filters,
            last_ingested=_last_ingested_label(),
            today_str=date_type.today().isoformat(),
            recent_additions=_recent_additions(),
        )

    # --- Flat list (default) ---
    stmt = stmt.order_by(HansardSession.date.desc())

    total       = stmt.count()
    sessions    = stmt.offset((page - 1) * _PER_PAGE).limit(_PER_PAGE).all()
    total_pages = max(1, (total + _PER_PAGE - 1) // _PER_PAGE)

    session_ids     = [s.id for s in sessions]
    contrib_counts  = _batch_load_contrib_counts(session_ids)
    policy_areas, specific_topics = _batch_load_tags(session_ids)
    session_items   = _build_session_items(sessions, contrib_counts, policy_areas, specific_topics)

    last_ingested = _last_ingested_label()
    today_str     = date_type.today().isoformat()

    return render_template(
        "hansard_archive/archive_home.html",
        grouped=False,
        group_items=[],
        session_items=session_items,
        q=q,
        house_filter=house_filter,
        dtype_filter=dtype_filter,
        dept_filter=dept_filter,
        policy_filter=policy_filter,
        date_from=date_from,
        date_to=date_to,
        title_only=title_only,
        page=page,
        total=total,
        total_pages=total_pages,
        per_page=_PER_PAGE,
        all_policy_areas=_all_policy_areas(),
        all_departments=_all_departments(),
        debate_type_labels=_DEBATE_TYPE_LABELS,
        base_qs=base_qs,
        has_filters=has_filters,
        last_ingested=last_ingested,
        today_str=today_str,
        recent_additions=_recent_additions(),
    )


# ---------------------------------------------------------------------------
# Session detail
# ---------------------------------------------------------------------------

@archive_bp.route("/debate/<string:date_str>/<string:slug>")
def session_detail(date_str: str, slug: str):
    session = _session_or_404(slug)
    if _url_date(session.date) != date_str:
        abort(404)
    ctx = _session_context(session)
    return render_template("hansard_archive/session_detail.html", **ctx)


# ---------------------------------------------------------------------------
# Word export — session transcript
# ---------------------------------------------------------------------------

@archive_bp.route("/debate/<string:date_str>/<string:slug>/word")
def session_word(date_str: str, slug: str):
    session = _session_or_404(slug)
    if _url_date(session.date) != date_str:
        abort(404)
    ctx = _session_context(session)

    doc = Document()

    # Title
    title_para = doc.add_heading(session.title, level=1)
    title_para.runs[0].font.size = Pt(16)

    # Metadata line
    meta = doc.add_paragraph()
    meta.add_run(f"{ctx['human_date']}  ·  {session.house}  ·  {ctx['debate_type_label']}").font.size = Pt(10)

    if ctx["policy_areas"]:
        pa = doc.add_paragraph()
        pa.add_run("Policy areas: ").bold = True
        pa.add_run(", ".join(ctx["policy_areas"])).font.size = Pt(10)

    if ctx["specific_topics"]:
        st = doc.add_paragraph()
        st.add_run("Topics: ").bold = True
        st.add_run(", ".join(ctx["specific_topics"])).font.size = Pt(10)

    if session.hansard_url:
        src = doc.add_paragraph()
        src.add_run("Source: ").bold = True
        src.add_run(session.hansard_url).font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Contributions
    for item in ctx["contributions"]:
        speaker = item["name"]
        if item["role"]:
            speaker += f" ({item['role']})"
        elif item["party"]:
            speaker += f" ({item['party']})"
        spk_para = doc.add_paragraph()
        spk_run = spk_para.add_run(speaker)
        spk_run.bold = True
        spk_run.font.size = Pt(11)
        spk_run.font.color.rgb = RGBColor(0x1A, 0x4A, 0x6E)

        for para_text in (item["speech_text"] or "").split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.runs[0].font.size = Pt(11)
        doc.add_paragraph()  # spacer between contributions

    # Footer
    footer = doc.add_paragraph()
    footer.add_run(
        f"Source: Hansard (Parliament Open Parliament Licence v3.0). "
        f"Exported from Westminster Brief — westminsterbrief.co.uk"
    ).font.size = Pt(9)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = "".join(c if c.isalnum() or c in " -" else "" for c in session.title)[:60]
    filename = f"{ctx['url_date']}-{safe_title.lower().replace(' ', '-')}.docx"

    response = make_response(buf.read())
    response.headers["Content-Type"] = (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response.headers["Content-Disposition"] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Date browse — /archive/date/<date_str>
# ---------------------------------------------------------------------------

@archive_bp.route("/date/<string:date_str>")
def archive_date(date_str: str):
    d = _parse_url_date(date_str)
    if d is None:
        abort(404)

    sessions = (
        HansardSession.query
        .filter_by(is_container=False, date=d)
        .order_by(HansardSession.id)   # Hansard chain order
        .all()
    )
    if not sessions:
        abort(404)

    session_ids    = [s.id for s in sessions]
    contrib_counts = _batch_load_contrib_counts(session_ids)
    policy_areas, specific_topics = _batch_load_tags(session_ids)
    items = _build_session_items(sessions, contrib_counts, policy_areas, specific_topics)

    human = _human_date(d)
    return render_template(
        "hansard_archive/archive_collection.html",
        page_type      = "date",
        heading        = human,
        subtitle       = f"{len(items)} session{'s' if len(items) != 1 else ''}",
        breadcrumb     = [("Hansard Archive", "/archive"), (human, None)],
        items          = items,
        page           = 1,
        total_pages    = 1,
        total          = len(items),
        per_page       = len(items),
        canonical_path = f"/archive/date/{date_str}",
        og_title       = f"{human} — Hansard Archive",
        meta_desc      = (
            f"Parliamentary debates from {human}. "
            f"{len(items)} Hansard session{'s' if len(items) != 1 else ''} including "
            f"Commons and Lords proceedings."
        ),
        json_ld_type   = "CollectionPage",
    )


# ---------------------------------------------------------------------------
# MP page — /archive/mp/<member_id>
# ---------------------------------------------------------------------------

@archive_bp.route("/mp/<int:member_id>")
def archive_mp(member_id: int):
    try:
        page = max(1, int(request.args.get("page", 1) or 1))
    except (ValueError, TypeError):
        page = 1

    # Representative name row
    name_row = (
        HansardContribution.query
        .filter_by(member_id=member_id)
        .filter(HansardContribution.member_name.isnot(None))
        .first()
    )
    if name_row is None:
        abort(404)
    member_attr = _parse_attribution(name_row.member_name)

    # Distinct sessions this member contributed to, newest first
    session_ids_q = (
        db.session.query(HansardContribution.session_id)
        .join(HansardSession, HansardSession.id == HansardContribution.session_id)
        .filter(
            HansardContribution.member_id == member_id,
            HansardContribution.member_name.isnot(None),
            HansardSession.is_container == False,
        )
        .distinct()
        .order_by(HansardSession.date.desc())
    )
    total = session_ids_q.count()
    if total == 0:
        abort(404)

    total_pages  = max(1, (total + _PER_PAGE - 1) // _PER_PAGE)
    page_sids    = [r[0] for r in session_ids_q.offset((page - 1) * _PER_PAGE).limit(_PER_PAGE).all()]

    sessions_by_id = {s.id: s for s in HansardSession.query.filter(HansardSession.id.in_(page_sids)).all()}

    # Contributions for this member on this page
    page_contribs = (
        HansardContribution.query
        .filter(
            HansardContribution.session_id.in_(page_sids),
            HansardContribution.member_id == member_id,
            HansardContribution.member_name.isnot(None),
        )
        .order_by(HansardContribution.speech_order)
        .all()
    )
    contribs_by_session: dict[int, list[str]] = defaultdict(list)
    for c in page_contribs:
        contribs_by_session[c.session_id].append(c.speech_text or "")

    items = [
        {
            "session":           sessions_by_id[sid],
            "human_date":        _human_date(sessions_by_id[sid].date),
            "url_date":          _url_date(sessions_by_id[sid].date),
            "debate_type_label": _DEBATE_TYPE_LABELS.get(sessions_by_id[sid].debate_type, "Proceedings"),
            "department":        sessions_by_id[sid].department or "",
            "speech_texts":      contribs_by_session.get(sid, []),
            "contrib_count":     len(contribs_by_session.get(sid, [])),
        }
        for sid in page_sids
        if sid in sessions_by_id
    ]

    name      = member_attr["name"]
    party     = member_attr["party"] or ""
    base_qs   = urlencode({"page": page}) if page > 1 else ""

    return render_template(
        "hansard_archive/archive_mp.html",
        member_id     = member_id,
        member_name   = name,
        member_party  = party,
        member_colour = member_attr["party_colour"],
        items         = items,
        total         = total,
        total_pages   = total_pages,
        per_page      = _PER_PAGE,
        page          = page,
        base_qs       = base_qs,
        og_title      = f"{name} — Hansard Archive",
        meta_desc     = (
            f"Parliamentary contributions by {name}{' (' + party + ')' if party else ''} "
            f"in the Hansard Archive. {total} session{'s' if total != 1 else ''} on record."
        ),
    )


# ---------------------------------------------------------------------------
# Department page — /archive/department/<slug>
# ---------------------------------------------------------------------------

@archive_bp.route("/department/<string:dept_slug>")
def archive_department(dept_slug: str):
    # Reverse-lookup: find department whose slugified name matches
    all_depts = _all_departments()
    dept_name = next((d for d in all_depts if _slugify(d) == dept_slug), None)
    if not dept_name:
        abort(404)

    try:
        page = max(1, int(request.args.get("page", 1) or 1))
    except (ValueError, TypeError):
        page = 1

    stmt = (
        HansardSession.query
        .filter_by(is_container=False, department=dept_name)
        .order_by(HansardSession.date.desc())
    )
    total       = stmt.count()
    sessions    = stmt.offset((page - 1) * _PER_PAGE).limit(_PER_PAGE).all()
    total_pages = max(1, (total + _PER_PAGE - 1) // _PER_PAGE)

    session_ids    = [s.id for s in sessions]
    contrib_counts = _batch_load_contrib_counts(session_ids)
    policy_areas, specific_topics = _batch_load_tags(session_ids)
    items          = _build_session_items(sessions, contrib_counts, policy_areas, specific_topics)

    base_qs = urlencode({"page": page}) if page > 1 else ""

    return render_template(
        "hansard_archive/archive_collection.html",
        page_type      = "department",
        heading        = dept_name,
        subtitle       = f"{total} oral questions session{'s' if total != 1 else ''}",
        breadcrumb     = [("Hansard Archive", "/archive"), (dept_name, None)],
        items          = items,
        page           = page,
        total_pages    = total_pages,
        total          = total,
        per_page       = _PER_PAGE,
        base_qs        = base_qs,
        canonical_path = f"/archive/department/{dept_slug}",
        og_title       = f"{dept_name} — Hansard Archive",
        meta_desc      = (
            f"Parliamentary questions and debates answered by the {dept_name}. "
            f"{total} Hansard session{'s' if total != 1 else ''} in the archive."
        ),
        json_ld_type   = "GovernmentOrganization",
    )


# ---------------------------------------------------------------------------
# Policy area page — /archive/policy/<slug>
# ---------------------------------------------------------------------------

@archive_bp.route("/policy/<string:policy_slug>")
def archive_policy(policy_slug: str):
    # Reverse-lookup: find policy area whose slugified name matches
    all_policies = _all_policy_areas()
    policy_name  = next((p for p in all_policies if _slugify(p) == policy_slug), None)
    if not policy_name:
        abort(404)

    try:
        page = max(1, int(request.args.get("page", 1) or 1))
    except (ValueError, TypeError):
        page = 1

    policy_sub = (
        db.session.query(HansardSessionTheme.session_id)
        .filter(
            HansardSessionTheme.theme == policy_name,
            HansardSessionTheme.theme_type == THEME_TYPE_POLICY_AREA,
        )
        .subquery()
    )
    stmt = (
        HansardSession.query
        .filter_by(is_container=False)
        .filter(HansardSession.id.in_(policy_sub))
        .order_by(HansardSession.date.desc())
    )
    total       = stmt.count()
    sessions    = stmt.offset((page - 1) * _PER_PAGE).limit(_PER_PAGE).all()
    total_pages = max(1, (total + _PER_PAGE - 1) // _PER_PAGE)

    session_ids    = [s.id for s in sessions]
    contrib_counts = _batch_load_contrib_counts(session_ids)
    policy_areas, specific_topics = _batch_load_tags(session_ids)
    items          = _build_session_items(sessions, contrib_counts, policy_areas, specific_topics)

    base_qs = urlencode({"page": page}) if page > 1 else ""

    return render_template(
        "hansard_archive/archive_collection.html",
        page_type      = "policy",
        heading        = policy_name,
        subtitle       = f"{total} session{'s' if total != 1 else ''} tagged with this policy area",
        breadcrumb     = [("Hansard Archive", "/archive"), (policy_name, None)],
        items          = items,
        page           = page,
        total_pages    = total_pages,
        total          = total,
        per_page       = _PER_PAGE,
        base_qs        = base_qs,
        canonical_path = f"/archive/policy/{policy_slug}",
        og_title       = f"{policy_name} — Hansard Archive",
        meta_desc      = (
            f"UK parliamentary debates on {policy_name}. "
            f"{total} Hansard session{'s' if total != 1 else ''} tagged with this GOV.UK policy area."
        ),
        json_ld_type   = "CollectionPage",
    )


# ---------------------------------------------------------------------------
# Specific theme page — /archive/theme/<slug>
# ---------------------------------------------------------------------------

@archive_bp.route("/theme/<string:theme_slug>")
def archive_theme(theme_slug: str):
    # Reverse-lookup: find theme whose slugified name matches (themes with ≥5 sessions only)
    rows = (
        db.session.query(
            HansardSessionTheme.theme,
            func.count(HansardSessionTheme.session_id),
        )
        .filter(HansardSessionTheme.theme_type == THEME_TYPE_SPECIFIC)
        .group_by(HansardSessionTheme.theme)
        .having(func.count(HansardSessionTheme.session_id) >= 5)
        .all()
    )
    theme_name = next((r[0] for r in rows if _slugify(r[0]) == theme_slug), None)
    if not theme_name:
        abort(404)

    try:
        page = max(1, int(request.args.get("page", 1) or 1))
    except (ValueError, TypeError):
        page = 1

    theme_sub = (
        db.session.query(HansardSessionTheme.session_id)
        .filter(
            HansardSessionTheme.theme == theme_name,
            HansardSessionTheme.theme_type == THEME_TYPE_SPECIFIC,
        )
        .subquery()
    )
    stmt = (
        HansardSession.query
        .filter_by(is_container=False)
        .filter(HansardSession.id.in_(theme_sub))
        .order_by(HansardSession.date.desc())
    )
    total       = stmt.count()
    sessions    = stmt.offset((page - 1) * _PER_PAGE).limit(_PER_PAGE).all()
    total_pages = max(1, (total + _PER_PAGE - 1) // _PER_PAGE)

    session_ids    = [s.id for s in sessions]
    contrib_counts = _batch_load_contrib_counts(session_ids)
    policy_areas, specific_topics = _batch_load_tags(session_ids)
    items          = _build_session_items(sessions, contrib_counts, policy_areas, specific_topics)

    base_qs = urlencode({"page": page}) if page > 1 else ""

    return render_template(
        "hansard_archive/archive_collection.html",
        page_type      = "theme",
        heading        = theme_name,
        subtitle       = f"{total} session{'s' if total != 1 else ''} on this topic",
        breadcrumb     = [("Hansard Archive", "/archive"), (theme_name, None)],
        items          = items,
        page           = page,
        total_pages    = total_pages,
        total          = total,
        per_page       = _PER_PAGE,
        base_qs        = base_qs,
        canonical_path = f"/archive/theme/{theme_slug}",
        og_title       = f"{theme_name} — Hansard Archive",
        meta_desc      = (
            f"UK parliamentary debates on {theme_name}. "
            f"{total} Hansard session{'s' if total != 1 else ''} tagged with this topic."
        ),
        json_ld_type   = "CollectionPage",
    )


# ---------------------------------------------------------------------------
# FTS search — /archive/search  (noindex)
# ---------------------------------------------------------------------------

_FTS_HEADLINE_OPTS = (
    "MaxFragments=1,StartSel=<mark>,StopSel=</mark>,MaxWords=35,MinWords=15"
)


def _fts_search(
    q_raw: str,
    page: int,
    policy_filter: str = "",
    house_filter: str = "",
    dtype_filter: str = "",
    dept_filter: str = "",
    date_from: str = "",
    date_to: str = "",
    title_only: bool = False,
) -> tuple[list, int]:
    """
    Full-text search using Postgres tsvectors. Returns (results, total_count).

    Each result dict contains:
      session_id, title, date, house, debate_type, slug, department,
      snippet (Markup — safe HTML with <mark> highlights), final_rank
    """
    # Detect phrase query (user wrapped in double quotes)
    if q_raw.startswith('"') and q_raw.endswith('"') and len(q_raw) > 2:
        ts_func = "phraseto_tsquery"
        q_clean = q_raw[1:-1]
    else:
        ts_func = "plainto_tsquery"
        q_clean = q_raw

    offset = (page - 1) * _PER_PAGE

    policy_join = (
        "JOIN ha_session_theme sth ON sth.session_id = s.id"
        "  AND sth.theme = :policy AND sth.theme_type = 'policy_area'"
        if policy_filter else ""
    )

    # Build optional WHERE conditions for extra filters
    extra_where_parts = []
    extra_params: dict = {}

    if house_filter in ("Commons", "Lords"):
        extra_where_parts.append("AND s.house = :house")
        extra_params["house"] = house_filter

    if dtype_filter and dtype_filter in _DEBATE_TYPE_LABELS:
        extra_where_parts.append("AND s.debate_type = :dtype")
        extra_params["dtype"] = dtype_filter

    if dept_filter:
        extra_where_parts.append("AND s.department ILIKE :dept")
        extra_params["dept"] = f"%{dept_filter}%"

    if date_from:
        extra_where_parts.append("AND s.date >= :date_from")
        extra_params["date_from"] = date_from

    if date_to:
        extra_where_parts.append("AND s.date <= :date_to")
        extra_params["date_to"] = date_to

    extra_where_sql = " ".join(extra_where_parts)

    body_match_clause = "" if title_only else f"OR c.id IS NOT NULL"

    count_sql = sqla_text(f"""
        SELECT COUNT(DISTINCT s.id)
        FROM ha_session s
        {policy_join}
        LEFT JOIN ha_contribution c ON c.session_id = s.id
            AND c.speech_tsv @@ {ts_func}('english', :q)
        WHERE s.is_container = false
          AND (
              s.title_tsv @@ {ts_func}('english', :q)
              {body_match_clause}
          )
          {extra_where_sql}
    """)
    rows_sql = sqla_text(f"""
        WITH best_contrib AS (
            SELECT DISTINCT ON (session_id)
                session_id,
                ts_rank(speech_tsv, {ts_func}('english', :q))      AS body_rank,
                ts_headline('english', speech_text,
                            {ts_func}('english', :q),
                            :hl_opts)                               AS snippet
            FROM ha_contribution
            WHERE speech_tsv @@ {ts_func}('english', :q)
            ORDER BY session_id, body_rank DESC
        ),
        title_match AS (
            SELECT id AS session_id,
                   ts_rank(title_tsv, {ts_func}('english', :q))    AS title_rank
            FROM ha_session
            WHERE title_tsv @@ {ts_func}('english', :q)
              AND is_container = false
        )
        SELECT
            s.id, s.title, s.date, s.house, s.debate_type, s.slug, s.department,
            COALESCE(tm.title_rank, 0.0) * 3
                + COALESCE(bc.body_rank, 0.0)                       AS final_rank,
            bc.snippet
        FROM ha_session s
        {policy_join}
        LEFT JOIN best_contrib bc ON bc.session_id = s.id
        LEFT JOIN title_match  tm ON tm.session_id = s.id
        WHERE s.is_container = false
          AND ({f"tm.session_id IS NOT NULL" if title_only else "bc.session_id IS NOT NULL OR tm.session_id IS NOT NULL"})
          {extra_where_sql}
        ORDER BY final_rank DESC
        LIMIT :lim OFFSET :off
    """)

    count_params = {"q": q_clean}
    if policy_filter:
        count_params["policy"] = policy_filter
    count_params.update(extra_params)

    params = {"q": q_clean, "hl_opts": _FTS_HEADLINE_OPTS,
              "lim": _PER_PAGE, "off": offset}
    if policy_filter:
        params["policy"] = policy_filter
    params.update(extra_params)

    total   = db.session.execute(count_sql, count_params).scalar() or 0
    rows    = db.session.execute(rows_sql, params).fetchall()

    results = []
    for row in rows:
        sid, title, d, house, dtype, slug, dept, rank, snippet = row
        results.append({
            "session_id":        sid,
            "title":             title,
            "date":              d,
            "house":             house,
            "debate_type":       dtype,
            "debate_type_label": _DEBATE_TYPE_LABELS.get(dtype, "Proceedings"),
            "slug":              slug,
            "department":        dept or "",
            "human_date":        _human_date(d),
            "url_date":          _url_date(d),
            "snippet":           Markup(snippet) if snippet else None,
        })
    return results, total


def _pq_fts_search(q_raw: str, limit: int = 20, policy_filter: str = "") -> list:
    """
    Full-text search over ha_pq using question_tsv. Returns up to `limit` results.

    Each result dict:
      result_type, uin, heading, asking_member, answering_body,
      tabled_date, is_answered, url, human_date, snippet
    """
    if q_raw.startswith('"') and q_raw.endswith('"') and len(q_raw) > 2:
        ts_func = "phraseto_tsquery"
        q_clean = q_raw[1:-1]
    else:
        ts_func = "plainto_tsquery"
        q_clean = q_raw

    policy_join = (
        "JOIN ha_pq_theme pth ON pth.pq_id = p.id"
        "  AND pth.theme = :policy AND pth.theme_type = 'policy_area'"
        if policy_filter else ""
    )

    sql = sqla_text(f"""
        SELECT
            p.id, p.uin, p.heading, p.asking_member, p.answering_body,
            p.tabled_date, p.is_answered,
            ts_rank(p.question_tsv, {ts_func}('english', :q)) AS rank,
            ts_headline('english',
                coalesce(p.heading, '') || ' ' || coalesce(p.question_text, ''),
                {ts_func}('english', :q),
                :hl_opts) AS snippet
        FROM ha_pq p
        {policy_join}
        WHERE p.question_tsv @@ {ts_func}('english', :q)
        ORDER BY rank DESC
        LIMIT :lim
    """)

    sql_params = {"q": q_clean, "hl_opts": _FTS_HEADLINE_OPTS, "lim": limit}
    if policy_filter:
        sql_params["policy"] = policy_filter

    rows = db.session.execute(sql, sql_params).fetchall()

    results = []
    for row in rows:
        pq_id, uin, heading, asking, answering, tabled, is_answered, rank, snippet = row
        results.append({
            "result_type":   "pq",
            "uin":           uin,
            "heading":       heading or uin,
            "asking_member": asking or "",
            "answering_body":answering or "",
            "tabled_date":   tabled,
            "is_answered":   is_answered,
            "url":           f"/archive/pq/{uin}",
            "human_date":    _human_date(tabled) if tabled else "",
            "snippet":       Markup(snippet) if snippet else None,
        })
    return results


def _ilike_search(q_raw: str, page: int) -> tuple[list, int]:
    """ilike fallback for SQLite local development."""
    stmt = (
        HansardSession.query
        .filter_by(is_container=False)
        .filter(HansardSession.title.ilike(f"%{q_raw}%"))
        .order_by(HansardSession.date.desc())
    )
    total    = stmt.count()
    sessions = stmt.offset((page - 1) * _PER_PAGE).limit(_PER_PAGE).all()
    results  = [
        {
            "session_id":        s.id,
            "title":             s.title,
            "date":              s.date,
            "house":             s.house,
            "debate_type":       s.debate_type,
            "debate_type_label": _DEBATE_TYPE_LABELS.get(s.debate_type, "Proceedings"),
            "slug":              s.slug,
            "department":        s.department or "",
            "human_date":        _human_date(s.date),
            "url_date":          _url_date(s.date),
            "snippet":           None,
        }
        for s in sessions
    ]
    return results, total


@archive_bp.route("/search")
def archive_search():
    q             = request.args.get("q", "").strip()
    policy_filter = request.args.get("policy", "").strip()
    house_filter  = request.args.get("house", "").strip()
    dtype_filter  = request.args.get("dtype", "").strip()
    dept_filter   = "" if house_filter == "Lords" else request.args.get("dept", "").strip()
    date_from     = request.args.get("from", "").strip()
    date_to       = request.args.get("to", "").strip()
    title_only    = request.args.get("title_only") == "1"
    try:
        page = max(1, int(request.args.get("page", 1) or 1))
    except (ValueError, TypeError):
        page = 1

    results, total, total_pages = [], 0, 1
    pq_results = []
    error_msg = ""

    if q:
        try:
            if _is_postgres():
                results, total = _fts_search(
                    q, page,
                    policy_filter=policy_filter,
                    house_filter=house_filter,
                    dtype_filter=dtype_filter,
                    dept_filter=dept_filter,
                    date_from=date_from,
                    date_to=date_to,
                    title_only=title_only,
                )
                pq_results = _pq_fts_search(q, limit=20)
            else:
                results, total = _ilike_search(q, page)
            total_pages = max(1, (total + _PER_PAGE - 1) // _PER_PAGE)
        except Exception as exc:
            error_msg = "Search is temporarily unavailable."
            print(f"[archive_search] error: {exc}", flush=True)

    qs_parts: dict = {"q": q}
    for k, v in [("policy", policy_filter), ("house", house_filter),
                 ("dtype", dtype_filter), ("dept", dept_filter),
                 ("from", date_from), ("to", date_to)]:
        if v:
            qs_parts[k] = v
    if title_only:
        qs_parts["title_only"] = "1"
    if page > 1:
        qs_parts["page"] = page
    base_qs = urlencode(qs_parts)

    has_filters = any([house_filter, dtype_filter, dept_filter,
                       policy_filter, date_from, date_to, title_only])

    resp = make_response(render_template(
        "hansard_archive/archive_search.html",
        q                   = q,
        policy_filter       = policy_filter,
        house_filter        = house_filter,
        dtype_filter        = dtype_filter,
        dept_filter         = dept_filter,
        date_from           = date_from,
        date_to             = date_to,
        title_only          = title_only,
        has_filters         = has_filters,
        all_policy_areas    = _all_policy_areas(),
        all_departments     = _all_departments(),
        debate_type_labels  = _DEBATE_TYPE_LABELS,
        results             = results,
        pq_results          = pq_results,
        total               = total,
        total_pages         = total_pages,
        per_page            = _PER_PAGE,
        page                = page,
        base_qs             = base_qs,
        error_msg           = error_msg,
        is_postgres         = _is_postgres(),
    ))
    resp.headers["X-Robots-Tag"] = "noindex, nofollow"
    return resp


# ---------------------------------------------------------------------------
# Written Question detail — /archive/pq/<uin>
# ---------------------------------------------------------------------------

@archive_bp.route("/pq/<string:uin>")
def pq_detail(uin: str):
    pq = HaPQ.query.filter_by(uin=uin.upper()).first()
    if pq is None:
        # Try as-is (some UIDs may not be uppercase)
        pq = HaPQ.query.filter_by(uin=uin).first()
    if pq is None:
        abort(404)

    themes = HaPQTheme.query.filter_by(pq_id=pq.id).all()
    policy_areas = sorted({t.theme for t in themes if t.theme_type == THEME_TYPE_POLICY_AREA})
    specific_topics = sorted({t.theme for t in themes if t.theme_type == THEME_TYPE_SPECIFIC})

    seo_title = f"{pq.heading or pq.uin} — {pq.uin} — Westminster Brief"

    return render_template(
        "hansard_archive/archive_pq_detail.html",
        pq              = pq,
        policy_areas    = policy_areas,
        specific_topics = specific_topics,
        human_tabled    = _human_date(pq.tabled_date) if pq.tabled_date else "",
        human_answered  = _human_date(pq.answer_date) if pq.answer_date else "",
        human_updated   = _human_date(pq.updated_at.date()) if pq.updated_at else "",
        seo_title       = seo_title,
        meta_desc       = (
            f"Written Question {pq.uin} — {pq.heading or 'Written Question'} "
            f"tabled by {pq.asking_member or 'an MP'} to {pq.answering_body or 'a department'}."
        ),
        canonical_path  = f"/archive/pq/{pq.uin}",
    )
