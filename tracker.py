import requests, os, json, re, concurrent.futures, io
from flask import Blueprint, render_template, request, send_file
from datetime import datetime, timedelta
from cache_models import CachedMember

try:
    import docx
    from docx import Document
    from docx.shared import RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None

tracker_bp = Blueprint('tracker', __name__)
api_key = os.environ.get("GEMINI_API_KEY")

DEPARTMENTS = {
    "All Departments": "", "Department for Education": "60", "Department of Health and Social Care": "17",
    "HM Treasury": "14", "Home Office": "1", "Ministry of Defence": "11", "Ministry of Justice": "54",
    "Department for Science, Innovation and Technology": "216", "Cabinet Office": "53"
}
MEMBER_CACHE = {}

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    text_element = OxmlElement('w:t')
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

_GEMINI_MODEL_CACHE = {}


def _gemini_generate(api_key, prompt):
    """Call Gemini REST API. Auto-detects a working model and endpoint version."""
    global _GEMINI_MODEL_CACHE

    if api_key not in _GEMINI_MODEL_CACHE:
        try:
            resp = requests.get(
                f"https://generativelanguage.googleapis.com/v1beta/models?key={api_key}",
                timeout=5
            )
            if resp.status_code == 200:
                available = [m['name'] for m in resp.json().get('models', [])
                             if 'generateContent' in m.get('supportedGenerationMethods', [])]
                # Prefer 2.5-flash-lite (same price as 2.0, available to all keys)
                # 2.0-flash and 2.0-flash-lite deprecated Feb 2026, hard shutdown Jun 2026
                for prefix in ['models/gemini-2.5-flash-lite', 'models/gemini-2.5-flash',
                               'models/gemini-flash-latest']:
                    match = next((m for m in available if m.startswith(prefix)), None)
                    if match:
                        _GEMINI_MODEL_CACHE[api_key] = match.removeprefix('models/')
                        break
                else:
                    first = available[0] if available else None
                    _GEMINI_MODEL_CACHE[api_key] = first.removeprefix('models/') if first else 'gemini-2.5-flash-lite'
        except Exception:
            _GEMINI_MODEL_CACHE[api_key] = 'gemini-2.5-flash-lite'

    model = _GEMINI_MODEL_CACHE.get(api_key, 'gemini-2.5-flash-lite')
    payload = {"contents": [{"parts": [{"text": prompt}]}]}

    for version in ('v1', 'v1beta'):
        url = f"https://generativelanguage.googleapis.com/{version}/models/{model}:generateContent?key={api_key}"
        try:
            r = requests.post(url, json=payload, timeout=90)
            if r.status_code == 200:
                return r.json()['candidates'][0]['content']['parts'][0]['text']
            if r.status_code not in (404, 400):
                raise Exception(f"HTTP {r.status_code}")
        except Exception:
            continue

    raise Exception(f"All endpoints failed for {model}")

def get_member_name(member_id):
    if not member_id: return "Unknown Member"
    if member_id in MEMBER_CACHE: return MEMBER_CACHE[member_id]

    # Check DB cache first
    cached = CachedMember.get(member_id)
    if cached:
        MEMBER_CACHE[member_id] = cached.name
        return cached.name

    try:
        url = f"https://members-api.parliament.uk/api/Members/{member_id}"
        resp = requests.get(url, timeout=3)
        if resp.status_code == 200:
            data = resp.json().get('value', {})
            name = data.get('nameDisplayAs', 'Unknown Member')
            party = (data.get('latestParty') or {}).get('name', '')
            membership = data.get('latestHouseMembership') or {}
            constituency = membership.get('membershipFrom', '')
            house = "Lords" if membership.get('house') == 2 else "Commons"
            image_url = data.get('thumbnailUrl', '')
            CachedMember.store(member_id, name, party, constituency, house, image_url)
            MEMBER_CACHE[member_id] = name
            return name
    except: pass
    return "Unknown Member"

WQ_URL = "https://questions-statements-api.parliament.uk/api/writtenquestions/questions"


def _fetch_unanswered_wqs(dept_id: str, target_date: str) -> list:
    """Fetch all unanswered WQs for dept_id on target_date (YYYY-MM-DD).

    Uses correct API params verified against the OpenAPI spec April 2026.
    See CLAUDE.md "WQ API constraints" for full parameter reference.
    answeringBodies is safe with a date anchor (~2s response); without one it
    causes full-table scans and timeouts.
    """
    params = {
        'tabledWhenFrom': target_date,
        'tabledWhenTo': target_date,
        'answered': 'Unanswered',
        'take': 1000,
    }
    if dept_id:
        params['answeringBodies'] = int(dept_id)

    resp = requests.get(WQ_URL, params=params, timeout=30)
    if resp.status_code != 200:
        return []

    payload = resp.json()
    results = payload.get('results') or []
    total = payload.get('totalResults', 0)

    # Paginate on heavy days (post-recess returns can exceed 1000 for large depts)
    skip = 1000
    while len(results) < total and skip < 5000:
        page = requests.get(WQ_URL, params={**params, 'skip': skip}, timeout=30)
        if page.status_code != 200:
            break
        batch = page.json().get('results') or []
        if not batch:
            break
        results.extend(batch)
        skip += 1000

    return results


@tracker_bp.route('/tracker', methods=['GET', 'POST'])
def morning_tracker():
    sorted_grouped_results = {}
    error_message = None
    selected_dept = ""
    sitting_day_used = None
    api_failed = False

    if request.method == 'POST':
        selected_dept = request.form.get('department', '').strip()
        results = []

        try:
            # Walk back up to 5 days to find the most recent sitting day that has
            # unanswered questions for this department. Handles weekends, bank
            # holidays, and recess gracefully without arbitrary date-window hacks.
            data = []
            for days_back in range(1, 6):
                target_date = (datetime.now() - timedelta(days=days_back)).strftime('%Y-%m-%d')
                data = _fetch_unanswered_wqs(selected_dept, target_date)
                if data:
                    sitting_day_used = target_date
                    break

            m_ids = {(item.get('value') or {}).get('askingMemberId')
                     for item in data if (item.get('value') or {}).get('askingMemberId')}
            with concurrent.futures.ThreadPoolExecutor(max_workers=15) as executor:
                executor.map(get_member_name, m_ids)

            for item in data:
                val = item.get('value') or {}
                member_id = val.get('askingMemberId')
                tabled_date_str = (val.get('dateTabled') or '').split('T')[0]
                try:
                    date_obj = datetime.fromisoformat(tabled_date_str)
                    f_date = f"{date_obj.day} {date_obj.strftime('%B %Y')}"
                except Exception:
                    f_date = "N/A"

                # Derive question type from API fields (verified against live data Apr 2026):
                #   house == 'Lords' → LORDS (14 calendar day deadline)
                #   isNamedDay == True → NAMED_DAY (deadline set by MP, min 2 sitting days)
                #   otherwise → ORDINARY (conventional 7 sitting day deadline)
                if val.get('house') == 'Lords':
                    question_type = 'LORDS'
                elif val.get('isNamedDay'):
                    question_type = 'NAMED_DAY'
                else:
                    question_type = 'ORDINARY'

                date_for_answer = (val.get('dateForAnswer') or '').split('T')[0]

                results.append({
                    'dept': val.get('answeringBodyName'),
                    'uin': str(val.get('uin')),
                    'member': get_member_name(member_id),
                    'member_id': member_id,
                    'house': val.get('house', 'Commons'),
                    'text': val.get('questionText', '').replace('<p>', '').replace('</p>', ''),
                    'raw_date': tabled_date_str,
                    'date_asked': f_date,
                    'due_date': tabled_date_str or 'TBC',
                    'date_for_answer': date_for_answer,
                    'question_type': question_type,
                    'is_answered': False,
                    'status': 'UNANSWERED',
                })

            categories = {}
            ai_status = ""

            if results and api_key:
                try:
                    questions_data = [{"uin": r['uin'], "text": r['text'][:200]} for r in results]
                    prompt = (
                        "Categorize these UK Parliamentary questions into broad team-level policy themes "
                        "that a single policy team would own (e.g., 'SEND', 'Early Years', 'Higher Education Finance', "
                        "'Disabled Children\\'s Social Care', 'School Standards'). "
                        "Use short, team-level labels — do NOT add sub-categories or qualifiers after a dash. "
                        "Return ONLY a valid JSON dictionary where keys are the UIN strings and values are the Themes. "
                        f"Data: {json.dumps(questions_data)}"
                    )
                    raw_text = _gemini_generate(api_key, prompt)
                    match = re.search(r'\{.*\}|\[.*\]', raw_text.replace('\n', ' '), re.DOTALL)
                    if match:
                        parsed = json.loads(match.group(0))
                        if isinstance(parsed, list):
                            for item in parsed:
                                k = str(item.get('uin', item.get('id', '')))
                                v = str(item.get('theme', 'Uncategorized'))
                                if k: categories[k] = v
                        elif isinstance(parsed, dict):
                            for k, v in parsed.items():
                                categories[str(k)] = str(v)
                    else:
                        ai_status = "(AI: unexpected response format)"
                except Exception as e:
                    ai_status = f"(AI Error: {str(e)[:60]})"

            if ai_status:
                print(f"[tracker] categorisation issue: {ai_status}")

            if results:
                temp_group = {}
                for r in results:
                    r_date = r['raw_date']
                    theme = categories.get(r['uin']) or "Uncategorized"
                    if ' - ' in theme:
                        theme = theme.split(' - ')[0].strip()
                    if r_date not in temp_group:
                        temp_group[r_date] = {'display_date': r['date_asked'], 'themes': {}}
                    if theme not in temp_group[r_date]['themes']:
                        temp_group[r_date]['themes'][theme] = []
                    temp_group[r_date]['themes'][theme].append(r)

                for date_key in sorted(temp_group.keys(), reverse=True):
                    sorted_grouped_results[date_key] = temp_group[date_key]

        except requests.exceptions.Timeout:
            error_message = "Parliament API timed out — try again in a moment."
            api_failed = True
        except Exception as e:
            error_message = f"Search error: {str(e)}"

    return render_template(
        'tracker.html',
        sorted_grouped_results=sorted_grouped_results,
        error_message=error_message,
        departments=DEPARTMENTS,
        selected_dept=selected_dept,
        is_post=(request.method == 'POST'),
        sitting_day_used=sitting_day_used,
        api_failed=api_failed,
    )

@tracker_bp.route('/download_tracker_word', methods=['POST'])
def download_tracker_word():
    if not Document:
        return "Word generator library missing! Please run 'pip install python-docx' in your PythonAnywhere bash console.", 500

    export_data_str = request.form.get('export_data')
    include_history = request.form.get('include_history') == 'true'
    include_ai_context = request.form.get('include_ai_context') == 'true'
    selected_dept = request.form.get('selected_dept', '').strip()

    if not export_data_str: return "No data provided to download.", 400
    export_data = json.loads(export_data_str)
    
    ai_context_dict = {}
    if include_ai_context and api_key:
        flat_questions = []
        for section in export_data:
            for q in section['questions']:
                flat_questions.append({"uin": q['uin'], "member": q['member'], "text": q['text']})
        
        try:
            prompt = (
                "You are an expert UK political analyst. For each of the following parliamentary questions, "
                "provide a brief 2-sentence explanation of WHY this specific MP might be asking this right now. "
                "Consider local constituency issues, party political campaigns, or recent national news. "
                "Return a pure JSON dictionary where the keys are the UIN strings and the values are your analysis. "
                f"Data: {json.dumps(flat_questions)}"
            )
            raw_text = _gemini_generate(api_key, prompt)
            match = re.search(r'\{.*\}', raw_text.replace('\n', ' '), re.DOTALL)
            if match:
                ai_context_dict = json.loads(match.group(0))
        except Exception as e:
            print(f"AI Context Error: {e}")

    doc = Document()
    doc.add_heading('Today’s PQs (Enhanced Briefing)', 0)
    three_months_ago = (datetime.now() - timedelta(days=90)).strftime('%Y-%m-%d')
    
    ignore_words = {'school', 'schools', 'education', 'student', 'students', 'university', 'universities', 'college', 'health', 'nhs', 'funding', 'fund', 'policy', 'department', 'support', 'review', 'system', 'provision', 'england'}

    for section in export_data:
        theme = section['theme']
        doc.add_heading(f"📅 {section['date']} - 🏷️ {theme}", level=1)
        
        for q in section['questions']:
            p = doc.add_paragraph()
            # NEW: Print dynamic status (Answered or Unanswered)
            p.add_run(f"[{q.get('status', 'UNANSWERED')}] {q['member']} (UIN: {q['uin']})\n").bold = True
            p.add_run(f"\"{q['text']}\"\n").italic = True
            p.add_run(f"Due: {q['due_date']}")
            
            if include_ai_context:
                analysis = ai_context_dict.get(str(q['uin']))
                if analysis:
                    ai_p = doc.add_paragraph()
                    ai_run = ai_p.add_run(f"   🤖 AI Political Context: {analysis}")
                    ai_run.font.color.rgb = RGBColor(100, 100, 100)
            
            if include_history and q.get('member_id') and str(q['member_id']).isdigit():
                try:
                    hist_params = {'askingMemberId': int(q['member_id']), 'tabledWhenFrom': three_months_ago, 'take': 50}
                    if selected_dept: hist_params['answeringBodies'] = [int(selected_dept)]

                    hist_resp = requests.get("https://questions-statements-api.parliament.uk/api/writtenquestions/questions", params=hist_params, timeout=10)
                    
                    if hist_resp.status_code == 200:
                        similar_questions_data = []
                        keywords = [w.lower() for w in re.findall(r'\w+', theme) if len(w) > 3 and w.lower() not in ignore_words and w.lower() != 'uncategorized']
                        if not keywords: keywords = [w.lower() for w in re.findall(r'\w+', q['text']) if len(w) > 4 and w.lower() not in ignore_words]

                        for h_item in hist_resp.json().get('results', []):
                            h_val = h_item.get('value', {})
                            h_uin = str(h_val.get('uin'))
                            if h_uin == q['uin']: continue 
                            h_text = h_val.get('questionText', '').replace('<p>','').replace('</p>','')
                            
                            if keywords and any(kw in h_text.lower() for kw in keywords):
                                h_date = (h_val.get('dateTabled') or '').split('T')[0]
                                h_link = f"https://questions-statements.parliament.uk/written-questions?SearchTerm={h_uin}"
                                similar_questions_data.append({'uin': h_uin, 'date': h_date, 'text': h_text, 'link': h_link})
                        
                        if similar_questions_data:
                            hp = doc.add_paragraph()
                            hp.add_run(f"   ↳ 🔍 Previous questions by {q['member']} on this topic (Last 3 Months):").bold = True
                            for sq in similar_questions_data:
                                sq_p = doc.add_paragraph()
                                sq_p.add_run(f"      • [UIN: {sq['uin']} | {sq['date']}] \"{sq['text']}\"\n").italic = True
                                sq_p.add_run(f"         Link: ").italic = True
                                add_hyperlink(sq_p, sq['link'], sq['link'])
                except Exception as e: print(f"History fetch error: {e}")
            
    mem_doc = io.BytesIO()
    doc.save(mem_doc)
    mem_doc.seek(0)
    
    return send_file(
        mem_doc,
        as_attachment=True,
        download_name=f"Today_PQs_Briefing_{datetime.now().strftime('%Y%m%d')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )