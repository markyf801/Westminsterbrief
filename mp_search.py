import requests, os, json, io, time, re, logging
from concurrent.futures import ThreadPoolExecutor
from flask import Blueprint, render_template, request, send_file, jsonify, copy_current_request_context
from datetime import datetime, timedelta

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

mp_search_bp = Blueprint('mp_search', __name__)

TWFY_API_KEY = os.environ.get("TWFY_API_KEY")

DEPARTMENTS = {
    "All Departments": "",
    "Cabinet Office": "53",
    "Department for Culture, Media and Sport": "47",
    "Department for Education": "60",
    "Department for Energy Security and Net Zero": "202",
    "Department for Environment, Food and Rural Affairs": "13",
    "Department for Science, Innovation and Technology": "216",
    "Department for Transport": "21",
    "Department for Work and Pensions": "29",
    "Department of Health and Social Care": "17",
    "Foreign, Commonwealth and Development Office": "208",
    "HM Treasury": "14",
    "Home Office": "1",
    "Ministry of Defence": "11",
    "Ministry of Housing, Communities and Local Government": "7",
    "Ministry of Justice": "54",
}

# ── Module-level caches {key: (data, timestamp)} ─────────────────────────────

_profile_cache: dict = {}   # keyed on "header_{id}" or "profile_{id}"
_speech_cache: dict = {}    # keyed on member_id (int)
_twfy_id_cache: dict = {}   # keyed on member_id (int)

_TTL_24H = 86400.0
_TTL_30D = 86400.0 * 30
_TWFY_NONE = '__none__'     # sentinel for "resolved but not found"


def _cache_get(cache: dict, key, ttl: float):
    entry = cache.get(key)
    if entry and (time.monotonic() - entry[1] < ttl):
        return entry[0]
    return None


def _cache_set(cache: dict, key, value):
    cache[key] = (value, time.monotonic())


# ── Word hyperlink helper ─────────────────────────────────────────────────────

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0000FF'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


# ── Data fetchers ─────────────────────────────────────────────────────────────

def _fetch_member_header(member_id: int) -> dict:
    """Name, party, role, photo, house. Cached 24h."""
    cache_key = f'header_{member_id}'
    cached = _cache_get(_profile_cache, cache_key, _TTL_24H)
    if cached is not None:
        return cached
    try:
        resp = requests.get(
            f"https://members-api.parliament.uk/api/Members/{member_id}",
            timeout=15
        )
        if resp.status_code == 200:
            val = resp.json().get('value', {})
            membership = val.get('latestHouseMembership') or {}
            is_lord = membership.get('house') == 2
            constituency = membership.get('membershipFrom', '')
            result = {
                'id': member_id,
                'name': val.get('nameDisplayAs', ''),
                'party': (val.get('latestParty') or {}).get('name', ''),
                'role': 'Life Peer' if is_lord else f"MP for {constituency}",
                'photo_url': val.get('thumbnailUrl', ''),
                'is_lord': is_lord,
                'house': 'Lords' if is_lord else 'Commons',
                'constituency': constituency,
            }
            _cache_set(_profile_cache, cache_key, result)
            return result
    except Exception as e:
        logger.warning('[mp_research] member header fetch failed: %s', e)
    return {}


def _fetch_parliament_bio(member_id: int) -> dict:
    try:
        resp = requests.get(
            f"https://members-api.parliament.uk/api/Members/{member_id}/Biography",
            timeout=15
        )
        if resp.status_code == 200:
            return resp.json().get('value', {})
    except Exception:
        pass
    return {}


def _fetch_interests_structured(member_id: int) -> list:
    try:
        resp = requests.get(
            f"https://members-api.parliament.uk/api/Members/{member_id}/RegisteredInterests",
            timeout=15
        )
        if resp.status_code == 200:
            result = []
            for cat in (resp.json().get('value') or []):
                cat_name = cat.get('name', '')
                for interest in (cat.get('interests') or []):
                    desc = (interest.get('interest') or '').strip()
                    if desc:
                        result.append({'category': cat_name, 'description': desc[:300]})
            return result
    except Exception:
        pass
    return []


def _fetch_profile_data(member_id: int) -> dict | None:
    """Career history + interests for Profile tab. Cached 24h."""
    cache_key = f'profile_{member_id}'
    cached = _cache_get(_profile_cache, cache_key, _TTL_24H)
    if cached is not None:
        return cached
    try:
        with ThreadPoolExecutor(max_workers=2) as ex:
            f_bio = ex.submit(_fetch_parliament_bio, member_id)
            f_int = ex.submit(_fetch_interests_structured, member_id)
            bio = f_bio.result(timeout=20)
            interests = f_int.result(timeout=20)
        result = {
            'govt_posts': bio.get('governmentPosts') or [],
            'opp_posts': bio.get('oppositionPosts') or [],
            'committees': bio.get('committeeMemberships') or [],
            'interests': interests,
        }
        _cache_set(_profile_cache, cache_key, result)
        return result
    except Exception as e:
        logger.warning('[mp_research] profile fetch failed: %s', e)
        return None


def _resolve_twfy_id(member_id: int, member_name: str, is_lord: bool) -> str | None:
    """Parliament member ID → TWFY person_id. Cached 30 days."""
    cached = _cache_get(_twfy_id_cache, member_id, _TTL_30D)
    if cached is not None:
        return None if cached == _TWFY_NONE else cached
    if not TWFY_API_KEY:
        return None
    endpoint = 'getLords' if is_lord else 'getMPs'
    try:
        resp = requests.get(
            f"https://www.theyworkforyou.com/api/{endpoint}",
            params={'key': TWFY_API_KEY, 'search': member_name, 'output': 'js'},
            timeout=15
        )
        if resp.status_code == 200:
            data = resp.json()
            items = data if isinstance(data, list) else data.get('rows', [])
            if items:
                twfy_id = str(items[0].get('person_id', '')).strip()
                if twfy_id:
                    _cache_set(_twfy_id_cache, member_id, twfy_id)
                    return twfy_id
    except Exception as e:
        logger.warning('[mp_research] TWFY ID resolution failed for %s: %s', member_name, e)
    _cache_set(_twfy_id_cache, member_id, _TWFY_NONE)
    return None


def _fetch_speeches(member_id: int, member_name: str, is_lord: bool) -> list:
    """Recent TWFY debate contributions, filtered to ≥50 words. Cached 24h."""
    cached = _cache_get(_speech_cache, member_id, _TTL_24H)
    if cached is not None:
        return cached
    if not TWFY_API_KEY:
        return []
    twfy_id = _resolve_twfy_id(member_id, member_name, is_lord)
    if not twfy_id:
        return []
    debate_types = ['lords'] if is_lord else ['commons', 'westminhall', 'debates']
    raw = []
    try:
        for dtype in debate_types:
            resp = requests.get(
                "https://www.theyworkforyou.com/api/getDebates",
                params={'key': TWFY_API_KEY, 'person': twfy_id,
                        'type': dtype, 'num': 50, 'output': 'js'},
                timeout=15
            )
            if resp.status_code == 200:
                data = resp.json()
                rows = data if isinstance(data, list) else data.get('rows', [])
                raw.extend(rows)
    except Exception as e:
        logger.warning('[mp_research] TWFY speech fetch failed for %s: %s', member_name, e)

    # Filter to substantive contributions (≥50 words) if body field is present
    if raw and 'body' in raw[0]:
        raw = [s for s in raw
               if len(re.sub(r'<[^>]+>', '', s.get('body', '')).split()) >= 50]

    # Normalise and sort
    result = []
    for s in raw:
        url = s.get('listurl', '')
        if url and url.startswith('/'):
            url = f"https://www.theyworkforyou.com{url}"
        elif url and url.startswith('?'):
            url = f"https://www.theyworkforyou.com/debates/{url}"
        if not url:
            gid = s.get('gid', '')
            if gid:
                url = f"https://www.theyworkforyou.com/debates/?id={gid}"
        # Debate title lives in parent.body (may contain HTML tags)
        parent_body = (s.get('parent') or {}).get('body', '') or ''
        subject = re.sub(r'<[^>]+>', '', parent_body).strip()
        if not subject:
            subject = re.sub(r'<[^>]+>', '', s.get('subsection_name', '') or '').strip()
        if not subject:
            subject = re.sub(r'<[^>]+>', '', s.get('section_name', '') or '').strip()
        if not subject:
            subject = 'Untitled contribution'
        result.append({
            'date': s.get('hdate', ''),
            'subject': subject,
            'chamber': 'Lords' if is_lord else 'Commons',
            'url': url,
        })

    result.sort(key=lambda x: x['date'], reverse=True)
    result = result[:25]
    _cache_set(_speech_cache, member_id, result)
    return result


def _fetch_pqs(member_id: int, dept: str, start_date: str, end_date: str) -> list:
    """Written Questions for this member. No caching — user expects fresh results."""
    params = {'askingMemberId': member_id, 'take': 500}
    if dept:
        params['answeringBodies'] = [int(dept)]
    if start_date:
        params['tabledWhenFrom'] = start_date
    if end_date:
        params['tabledWhenTo'] = end_date
    try:
        resp = requests.get(
            "https://questions-statements-api.parliament.uk/api/writtenquestions/questions",
            params=params, timeout=30
        )
        if resp.status_code != 200:
            return []
        results = []
        for item in (resp.json().get('results') or []):
            val = item.get('value', {})
            # Client-side dept filter as belt-and-braces
            if dept and str(val.get('answeringBodyId', '')) != dept:
                continue
            raw_date = val.get('dateTabled', '')
            date_str = raw_date.split('T')[0] if raw_date else ''
            is_answered = bool(val.get('answerText') or val.get('dateAnswered'))
            uin = str(val.get('uin', ''))
            results.append({
                'uin': uin,
                'dept': val.get('answeringBodyName', 'Unknown Dept'),
                'text': re.sub(r'<[^>]+>', '', val.get('questionText', '')).strip(),
                'date': date_str,
                'status': 'ANSWERED' if is_answered else 'UNANSWERED',
                'link': f"https://questions-statements.parliament.uk/written-questions?SearchTerm={uin}",
            })
        return results
    except Exception as e:
        logger.warning('[mp_research] PQ fetch failed: %s', e)
        return []


# ── Routes ────────────────────────────────────────────────────────────────────

@mp_search_bp.route('/mp_search', methods=['GET', 'POST'])
def search_mp_pqs():
    member_header = {}
    profile_data = None
    profile_error = None
    pq_results = []
    pq_error = None
    speeches = []
    speeches_error = None
    selected_dept = ''
    start_date = ''
    end_date = ''
    member_id_str = ''

    if request.method == 'POST':
        member_id_str = (request.form.get('member_id') or '').strip()
        selected_dept = (request.form.get('department') or '').strip()
        start_date = (request.form.get('start_date') or '').strip()
        end_date = (request.form.get('end_date') or '').strip()

        if not member_id_str:
            pq_error = 'Please select a member.'
        else:
            try:
                member_id = int(member_id_str)
                member_header = _fetch_member_header(member_id)
                if not member_header:
                    pq_error = 'Could not load member details — please try again.'
                else:
                    is_lord = member_header.get('is_lord', False)
                    member_name = member_header.get('name', '')

                    with ThreadPoolExecutor(max_workers=3) as ex:
                        f_profile = ex.submit(
                            copy_current_request_context(_fetch_profile_data), member_id)
                        f_pqs = ex.submit(
                            copy_current_request_context(_fetch_pqs),
                            member_id, selected_dept, start_date, end_date)
                        f_speeches = ex.submit(
                            copy_current_request_context(_fetch_speeches),
                            member_id, member_name, is_lord)

                        try:
                            profile_data = f_profile.result(timeout=25)
                        except Exception as e:
                            profile_error = 'Could not load profile data.'
                            logger.warning('[mp_research] profile tab error: %s', e)

                        try:
                            pq_results = f_pqs.result(timeout=35)
                        except Exception as e:
                            pq_error = 'Could not load written questions.'
                            logger.warning('[mp_research] PQ tab error: %s', e)

                        try:
                            speeches = f_speeches.result(timeout=25)
                        except Exception as e:
                            speeches_error = 'Could not load speeches.'
                            logger.warning('[mp_research] speeches tab error: %s', e)

            except ValueError:
                pq_error = 'Invalid member selection.'

    return render_template(
        'mp_search.html',
        member_header=member_header,
        profile_data=profile_data, profile_error=profile_error,
        pq_results=pq_results, pq_error=pq_error,
        speeches=speeches, speeches_error=speeches_error,
        departments=DEPARTMENTS,
        selected_dept=selected_dept,
        start_date=start_date, end_date=end_date,
        member_id=member_id_str,
        is_post=(request.method == 'POST'),
    )


@mp_search_bp.route('/api/mp_pqs')
def api_mp_pqs():
    """Legacy JSON endpoint — used by Debate Scanner 'View PQs' widget."""
    name = request.args.get('name', '').strip()
    topic = request.args.get('topic', '').strip()
    dept = request.args.get('dept', '').strip()
    if not name:
        return jsonify({'error': 'Name required'}), 400
    try:
        s_resp = requests.get(
            f"https://members-api.parliament.uk/api/Members/Search?Name={name}&IsCurrentMember=true&take=5",
            timeout=5
        )
        if not (s_resp.status_code == 200 and s_resp.json().get('items')):
            return jsonify({'error': f'MP not found: {name}'}), 404

        member_data = s_resp.json()['items'][0]['value']
        member_id = member_data['id']
        display_name = member_data.get('nameDisplayAs', name)
        party = (member_data.get('latestParty') or {}).get('name', '')
        membership = member_data.get('latestHouseMembership') or {}
        house = 'Lords' if membership.get('house') == 2 else 'Commons'
        constituency = membership.get('membershipFrom', '')
        role = 'Life Peer' if house == 'Lords' else f"MP for {constituency}"

        six_months_ago = (datetime.now() - timedelta(days=180)).strftime('%Y-%m-%d')
        q_resp = requests.get(
            "https://questions-statements-api.parliament.uk/api/writtenquestions/questions",
            params={'askingMemberId': member_id, 'tabledWhenFrom': six_months_ago, 'take': 50},
            timeout=10
        )
        all_pqs = []
        if q_resp.status_code == 200:
            for item in q_resp.json().get('results', []):
                val = item.get('value', {})
                raw_date = (val.get('dateTabled') or '').split('T')[0]
                is_answered = bool(val.get('answerText') or val.get('dateAnswered'))
                uin = str(val.get('uin', ''))
                all_pqs.append({
                    'uin': uin,
                    'dept': val.get('answeringBodyName', ''),
                    'text': re.sub(r'<[^>]+>', '', val.get('questionText', '')).strip(),
                    'date': raw_date,
                    'status': 'ANSWERED' if is_answered else 'UNANSWERED',
                    'link': f"https://questions-statements.parliament.uk/written-questions/detail/{raw_date}/{uin}"
                })

        stop_words = {'the', 'and', 'for', 'that', 'this', 'with', 'from', 'are', 'have'}
        pqs = all_pqs
        topic_filtered = False
        if topic and all_pqs:
            keywords = [w.lower().strip('"\' ') for w in topic.split()
                        if len(w) > 3 and w.lower() not in stop_words]
            if keywords:
                filtered = [q for q in all_pqs if any(kw in q['text'].lower() for kw in keywords)]
                if filtered:
                    pqs = filtered
                    topic_filtered = True
                elif dept:
                    pqs = [q for q in all_pqs if dept.lower() in q['dept'].lower()]
                else:
                    pqs = []

        return jsonify({
            'name': display_name, 'party': party, 'role': role,
            'pqs': pqs[:15], 'topic_filtered': topic_filtered, 'total_pqs': len(all_pqs)
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@mp_search_bp.route('/download_mp_word', methods=['POST'])
def download_mp_word():
    if not Document:
        return 'Word library missing.', 500

    export_data_str = request.form.get('export_data')
    if not export_data_str:
        return 'No data.', 400

    payload = json.loads(export_data_str)
    header = payload.get('member_header') or payload.get('mp_details') or {}
    profile = payload.get('profile_data')
    questions = payload.get('questions', [])
    speech_list = payload.get('speeches', [])
    ai_profile = (payload.get('ai_profile') or '').strip()

    name = header.get('name', 'Unknown Member')
    party = header.get('party', '')
    role = header.get('role', '')
    photo_url = header.get('photo_url', '')

    doc = Document()

    def compact(para, before=0, after=3):
        pPr = para._p.get_or_add_pPr()
        ex = pPr.find(qn('w:spacing'))
        if ex is not None:
            pPr.remove(ex)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), str(before * 20))
        sp.set(qn('w:after'), str(after * 20))
        pPr.append(sp)

    def remove_borders(table):
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            borders.append(el)
        tblPr.append(borders)

    def section_heading(text):
        p = doc.add_heading(text, level=1)
        compact(p, before=12, after=4)

    # ── 1. Header ─────────────────────────────────────────────────────────────
    img_stream = None
    if photo_url:
        try:
            img_resp = requests.get(photo_url, timeout=5)
            if img_resp.status_code == 200:
                img_stream = io.BytesIO(img_resp.content)
        except Exception:
            pass

    ht = doc.add_table(rows=1, cols=2)
    remove_borders(ht)
    pc, dc = ht.cell(0, 0), ht.cell(0, 1)
    pc.width = Inches(1.5)
    dc.width = Inches(5.0)
    if img_stream:
        pc.paragraphs[0].add_run().add_picture(img_stream, width=Inches(1.3))

    np_ = dc.paragraphs[0]
    r = np_.add_run(name)
    r.bold = True; r.font.size = Pt(16)
    compact(np_, before=0, after=4)

    mp = dc.add_paragraph()
    mp.add_run('  ·  '.join(p for p in [party, role] if p))
    compact(mp, before=0, after=3)

    dp = dc.add_paragraph()
    dr = dp.add_run(f'Generated {datetime.now().strftime("%d %B %Y")}')
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(128, 128, 128)
    compact(dp, before=0, after=0)
    doc.add_paragraph('─' * 80)

    # ── 2. Profile section ────────────────────────────────────────────────────
    if profile:
        section_heading('Profile')

        def posts_table(posts, label):
            if not posts:
                return
            doc.add_heading(label, level=2)
            t = doc.add_table(rows=1, cols=3)
            t.style = 'Table Grid'
            hdr = t.rows[0].cells
            for i, h in enumerate(['Post', 'From', 'To']):
                hdr[i].text = h
                hdr[i].paragraphs[0].runs[0].bold = True
            for post in sorted(posts, key=lambda x: x.get('startDate') or '', reverse=True):
                row = t.add_row().cells
                row[0].text = post.get('name', '')
                row[1].text = (post.get('startDate') or '').split('T')[0][:7]
                end = (post.get('endDate') or '').split('T')[0][:7]
                row[2].text = end if end else 'present'

        posts_table(profile.get('govt_posts', []), 'Government Roles')
        posts_table(profile.get('opp_posts', []), 'Opposition Roles')
        posts_table(profile.get('committees', []), 'Committee Memberships')

        interests = profile.get('interests', [])
        if interests:
            doc.add_heading('Registered Interests', level=2)
            t = doc.add_table(rows=1, cols=2)
            t.style = 'Table Grid'
            hdr = t.rows[0].cells
            hdr[0].text = 'Category'; hdr[0].paragraphs[0].runs[0].bold = True
            hdr[1].text = 'Interest'; hdr[1].paragraphs[0].runs[0].bold = True
            for interest in interests[:40]:
                row = t.add_row().cells
                row[0].text = interest.get('category', '')
                row[1].text = interest.get('description', '')

    # ── 3. AI Profile (only if generated) ────────────────────────────────────
    if ai_profile:
        section_heading('AI-Generated Profile')
        heading_re = re.compile(r'^#{1,3}\s+(.+)$')
        bullet_re = re.compile(r'^[-*]\s+(.+)$')
        for line in ai_profile.split('\n'):
            line = line.rstrip()
            if not line:
                continue
            hm = heading_re.match(line)
            if hm:
                p = doc.add_heading(hm.group(1).strip(), level=2)
                compact(p, before=8, after=2)
                continue
            bm = bullet_re.match(line)
            if bm:
                p = doc.add_paragraph(bm.group(1).strip(), style='List Bullet')
                compact(p, before=0, after=2)
                continue
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            if clean:
                p = doc.add_paragraph(clean)
                compact(p, before=0, after=3)

    # ── 4. Written Questions section ──────────────────────────────────────────
    if questions:
        section_heading(f'Written Questions ({len(questions)})')
        for q in questions:
            qp = doc.add_paragraph()
            qp.add_run(f"[{q['status']}] {q['date']} | To: {q['dept']}\n").bold = True
            qp.add_run(f"\"{q['text']}\"\n").italic = True
            qp.add_run('Link: ')
            add_hyperlink(qp, q['link'], q['link'])
            compact(qp, before=0, after=8)

    # ── 5. Speeches section ───────────────────────────────────────────────────
    if speech_list:
        section_heading(f'Recent Speeches ({len(speech_list)})')
        house_note = 'Lords' if header.get('is_lord') else 'Commons + Westminster Hall'
        np2 = doc.add_paragraph(f'Source: {house_note} contributions via TheyWorkForYou')
        np2.runs[0].font.size = Pt(9)
        np2.runs[0].font.color.rgb = RGBColor(128, 128, 128)
        compact(np2, before=0, after=6)

        t = doc.add_table(rows=1, cols=3)
        t.style = 'Table Grid'
        hdr = t.rows[0].cells
        for i, h in enumerate(['Date', 'Debate / Subject', 'Chamber']):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for s in speech_list:
            row = t.add_row().cells
            row[0].text = s.get('date', '')
            sp = row[1].paragraphs[0]
            if s.get('url'):
                add_hyperlink(sp, s['url'], s.get('subject', ''))
            else:
                sp.add_run(s.get('subject', ''))
            row[2].text = s.get('chamber', '')

    # ── 6. Footer ─────────────────────────────────────────────────────────────
    doc.add_paragraph('─' * 80)
    fp = doc.add_paragraph()
    fp.add_run('Data sources: Parliament Members API · Parliament Written Questions API · TheyWorkForYou\n')
    fp.add_run(f'Generated {datetime.now().strftime("%d %B %Y %H:%M")} · Westminster Brief')
    fp.runs[0].font.size = Pt(8)
    fp.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    safe_name = re.sub(r'[^\w\s-]', '', name)[:40].strip()
    return send_file(
        mem, as_attachment=True,
        download_name=f"MP Research - {safe_name}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
