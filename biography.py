import os, requests, io, json, re, concurrent.futures, logging
from flask import Blueprint, render_template, request, jsonify, send_file
from datetime import datetime, timedelta
from cache_models import CachedMember
from extensions import limiter

try:
    import wikipedia
except ImportError:
    wikipedia = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None

biography_bp = Blueprint('biography', __name__)

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
_MODEL_CACHE = None


def _gemini_model():
    global _MODEL_CACHE
    if _MODEL_CACHE:
        return _MODEL_CACHE
    if not GEMINI_API_KEY:
        return None
    try:
        resp = requests.get(
            f"https://generativelanguage.googleapis.com/v1beta/models?key={GEMINI_API_KEY}",
            timeout=5
        )
        if resp.status_code == 200:
            available = [m['name'] for m in resp.json().get('models', [])
                         if 'generateContent' in m.get('supportedGenerationMethods', [])]
            for pref in ['models/gemini-2.5-flash-lite', 'models/gemini-2.5-flash',
                         'models/gemini-flash-latest']:
                match = next((m for m in available if m.startswith(pref)), None)
                if match:
                    _MODEL_CACHE = match
                    return _MODEL_CACHE
            if available:
                _MODEL_CACHE = available[0]
                return _MODEL_CACHE
    except Exception:
        pass
    _MODEL_CACHE = "models/gemini-2.5-flash-lite"
    return _MODEL_CACHE


def _add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0000FF'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ── Data fetchers ─────────────────────────────────────────────────────────────

def _fetch_parliament_biography(member_id):
    try:
        resp = requests.get(
            f"https://members-api.parliament.uk/api/Members/{member_id}/Biography",
            timeout=8
        )
        if resp.status_code == 200:
            return resp.json().get('value', {})
    except Exception:
        pass
    return {}


def _fetch_recent_pqs(member_id, limit=20):
    try:
        six_months_ago = (datetime.now() - timedelta(days=180)).strftime('%Y-%m-%d')
        resp = requests.get(
            "https://questions-statements-api.parliament.uk/api/writtenquestions/questions",
            params={'askingMemberId': member_id, 'tabledWhenFrom': six_months_ago, 'take': limit},
            timeout=10
        )
        if resp.status_code == 200:
            results = []
            for item in resp.json().get('results', []):
                val = item.get('value', {})
                text = re.sub(r'\s+', ' ', re.sub(r'<[^>]+>', ' ', val.get('questionText') or '')).strip()
                dept = val.get('answeringBodyName', '')
                if text:
                    results.append(f"[{dept}] {text[:200]}")
            return results
    except Exception:
        pass
    return []


def _fetch_registered_interests(member_id):
    try:
        resp = requests.get(
            f"https://members-api.parliament.uk/api/Members/{member_id}/RegisteredInterests",
            timeout=8
        )
        if resp.status_code == 200:
            categories = resp.json().get('value', []) or []
            lines = []
            for cat in categories:
                cat_name = cat.get('name', '')
                for interest in (cat.get('interests') or []):
                    desc = (interest.get('interest') or '').strip()
                    if desc:
                        lines.append(f"  [{cat_name}] {desc[:200]}")
            return lines
    except Exception:
        pass
    return []


def _fetch_wikipedia(mp_name):
    if not wikipedia:
        return ""
    try:
        search = wikipedia.search(f"{mp_name} British politician")
        if not search:
            search = wikipedia.search(mp_name)
        if search:
            try:
                page = wikipedia.page(search[0], auto_suggest=False)
            except wikipedia.DisambiguationError as e:
                options = [o for o in e.options
                           if any(k in o.lower() for k in ('politician', ' mp', 'member of parliament', 'lord', 'baroness'))]
                page = wikipedia.page(options[0] if options else e.options[0], auto_suggest=False)
            return page.content[:4000]
    except Exception:
        pass
    return ""


def _format_posts(posts, label):
    if not posts:
        return ""
    lines = [f"\n{label}:"]
    for p in sorted(posts, key=lambda x: x.get('startDate') or '', reverse=True):
        name = p.get('name', '')
        start = (p.get('startDate') or '').split('T')[0][:7]
        end = (p.get('endDate') or '').split('T')[0][:7]
        date_str = f"{start}–{end}" if end else f"{start}–present"
        if name:
            lines.append(f"  - {name} ({date_str})")
    return '\n'.join(lines)


def _format_committees(committees):
    if not committees:
        return ""
    lines = ["\nCommittee Memberships (current and historical):"]
    for c in sorted(committees, key=lambda x: x.get('startDate') or '', reverse=True):
        name = c.get('name', '')
        start = (c.get('startDate') or '').split('T')[0][:7]
        end = (c.get('endDate') or '').split('T')[0][:7]
        date_str = f"{start}–{end}" if end else f"{start}–present"
        if name:
            lines.append(f"  - {name} ({date_str})")
    return '\n'.join(lines)


# ── Main biography generator ──────────────────────────────────────────────────

def generate_mp_biography(mp_name, member_id):
    # Parallel fetch: Parliament biography API + registered interests + recent PQs + Wikipedia
    with concurrent.futures.ThreadPoolExecutor(max_workers=4) as ex:
        f_bio = ex.submit(_fetch_parliament_biography, member_id)
        f_interests = ex.submit(_fetch_registered_interests, member_id)
        f_pqs = ex.submit(_fetch_recent_pqs, member_id)
        f_wiki = ex.submit(_fetch_wikipedia, mp_name)
        parl_bio = f_bio.result()
        interests = f_interests.result()
        pqs = f_pqs.result()
        wiki_text = f_wiki.result()

    # Build Parliament context block
    parl_context = ""
    parl_context += _format_posts(parl_bio.get('governmentPosts', []), "Government Roles")
    parl_context += _format_posts(parl_bio.get('oppositionPosts', []), "Opposition Roles")
    parl_context += _format_committees(parl_bio.get('committeeMemberships', []))

    interests_context = ""
    if interests:
        cap = 30
        interests_context = "\nRegistered Interests:\n" + "\n".join(f"  • {i}" for i in interests[:cap])
        if len(interests) > cap:
            interests_context += f"\n  ... and {len(interests) - cap} further interests (see Parliament.uk for full register)"

    pq_context = ""
    if pqs:
        pq_context = "\nRecent Written Questions (sample):\n" + "\n".join(f"  • {q}" for q in pqs[:12])

    prompt = f"""You are a senior UK parliamentary researcher writing a professional briefing note.

Write a structured profile for: {mp_name}

Use ALL of the source data below. Clearly distinguish between their pre-parliamentary career
(from Wikipedia) and their parliamentary record (from the Parliament API data).

--- WIKIPEDIA EXTRACT (use for background, education, pre-parliament career) ---
{wiki_text[:3500] if wiki_text else "No Wikipedia data available."}

--- PARLIAMENT API DATA ---
{parl_context if parl_context else "No parliamentary posts or committee data available."}
{interests_context}
{pq_context}

---

Write the profile using these exact markdown headings:

## Career Before Parliament
Background, profession, education — draw on Wikipedia. If no data available, say so briefly.

## Parliamentary Career
When elected/appointed, party, constituency or peerage. Key milestones.

## Government & Opposition Roles
All ministerial or shadow ministerial roles held, with dates. Distinguish clearly between government and opposition roles. If none, say so.

## Committee Memberships
List all select committees and other parliamentary committees, with dates (include historical ones). If none on record, say so.

## Registered Interests
Summarise their declared financial and other interests by category. If none declared, say so.

## Policy Interests & Focus Areas
Based on their written questions and stated interests — be specific, not generic.

## Key Context for Meetings
2–3 sentences a civil servant would want to know before a ministerial meeting or correspondence.

Keep each section concise and factual. Only include information present in the source data above — do not supplement from your training data or general knowledge. If a section has no source data, say so briefly rather than inferring."""

    model = _gemini_model()
    if not model or not GEMINI_API_KEY:
        return "AI service unavailable — please check your GEMINI_API_KEY."

    try:
        resp = requests.post(
            f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={GEMINI_API_KEY}",
            json={"contents": [{"parts": [{"text": prompt}]}]},
            timeout=45
        )
        resp.raise_for_status()
        return resp.json()['candidates'][0]['content']['parts'][0]['text']
    except Exception as e:
        logging.warning(f"[biography] AI generation failed for {mp_name} (model={model}): {type(e).__name__}: {e}")
        return "AI summary unavailable — please try again."


# ── Routes ────────────────────────────────────────────────────────────────────

@biography_bp.route('/biography', methods=['GET', 'POST'])
def biography_home():
    mp_data, error = None, None
    if request.method == 'POST':
        member_id = request.form.get('member_id')
        if member_id:
            cached = CachedMember.get(member_id)
            if cached:
                mp_data = {
                    'id': cached.member_id, 'name': cached.name,
                    'party': cached.party, 'constituency': cached.constituency,
                    'image_url': cached.image_url
                }
            else:
                try:
                    resp = requests.get(
                        f"https://members-api.parliament.uk/api/Members/{member_id}",
                        timeout=8
                    ).json().get('value', {})
                    name = resp.get('nameDisplayAs')
                    party = (resp.get('latestParty') or {}).get('name')
                    membership = resp.get('latestHouseMembership') or {}
                    constituency = membership.get('membershipFrom')
                    house = "Lords" if membership.get('house') == 2 else "Commons"
                    image_url = resp.get('thumbnailUrl')
                    CachedMember.store(member_id, name, party, constituency, house, image_url)
                    mp_data = {
                        'id': resp.get('id'), 'name': name, 'party': party,
                        'constituency': constituency, 'image_url': image_url
                    }
                except Exception:
                    error = "Could not load member details."
    return render_template('biography.html', mp_data=mp_data, error_message=error)


@biography_bp.route('/api/search_members')
def api_search_members():
    term = request.args.get('q', '')
    if len(term) < 3:
        return jsonify({"results": []})
    try:
        items = requests.get(
            "https://members-api.parliament.uk/api/Members/Search",
            params={'Name': term, 'take': 15},
            timeout=5
        ).json().get('items') or []
        return jsonify({
            "results": [{"id": i['value']['id'], "text": i['value']['nameDisplayAs']} for i in items]
        })
    except Exception:
        return jsonify({"results": []})


@biography_bp.route("/api/biography", methods=["POST"])
@limiter.limit("10 per minute; 100 per day")
def api_biography():
    data = request.get_json()
    bio = generate_mp_biography(data.get("mp_name"), data.get("member_id"))
    return jsonify({"biography": bio})


@biography_bp.route('/export_biography_word', methods=['POST'])
def export_biography_word():
    if not Document:
        return "Word library missing.", 500

    from docx.shared import Inches, Pt
    mp_name = request.form.get('mp_name', 'Unknown Member')
    party = request.form.get('party', '')
    constituency = request.form.get('constituency', '')
    bio_text = request.form.get('bio_text', '')
    image_url = request.form.get('image_url', '')
    member_id = request.form.get('member_id', '')

    doc = Document()

    # Helper: set compact paragraph spacing (values in points)
    def compact(para, before=0, after=3):
        pPr = para._p.get_or_add_pPr()
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), str(before * 20))
        sp.set(qn('w:after'), str(after * 20))
        pPr.append(sp)

    # Helper: remove all borders from a table
    def remove_table_borders(table):
        tblPr = table._tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            table._tbl.insert(0, tblPr)
        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'none')
            borders.append(el)
        tblPr.append(borders)

    # Try to download photo
    img_stream = None
    if image_url:
        try:
            img_resp = requests.get(image_url, timeout=5)
            if img_resp.status_code == 200:
                img_stream = io.BytesIO(img_resp.content)
        except Exception:
            pass

    # ── Header: photo left, name/details right ───────────────────────────────
    header_table = doc.add_table(rows=1, cols=2)
    remove_table_borders(header_table)
    photo_cell = header_table.cell(0, 0)
    detail_cell = header_table.cell(0, 1)

    # Set column widths
    photo_cell.width = Inches(1.5)
    detail_cell.width = Inches(5.0)

    if img_stream:
        photo_cell.paragraphs[0].add_run().add_picture(img_stream, width=Inches(1.3))
    else:
        photo_cell.paragraphs[0].add_run('')

    # Name
    name_para = detail_cell.paragraphs[0]
    name_run = name_para.add_run(mp_name)
    name_run.bold = True
    name_run.font.size = Pt(16)
    compact(name_para, before=0, after=4)

    # Party · Constituency
    meta_para = detail_cell.add_paragraph()
    parts = [p for p in [party, constituency] if p]
    meta_para.add_run('  ·  '.join(parts))
    compact(meta_para, before=0, after=3)

    # Generated date
    date_para = detail_cell.add_paragraph()
    date_run = date_para.add_run(f'Profile generated {datetime.now().strftime("%d %B %Y")}')
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = RGBColor(128, 128, 128)
    compact(date_para, before=0, after=0)

    # Divider line after header
    doc.add_paragraph('─' * 80)

    # ── Biography content ─────────────────────────────────────────────────────
    heading_re = re.compile(r'^#{1,3}\s+(.+)$')
    bullet_re = re.compile(r'^[-*]\s+(.+)$')

    for line in bio_text.split('\n'):
        line = line.rstrip()
        if not line:
            continue  # skip blank lines — use spacing instead
        hm = heading_re.match(line)
        if hm:
            level = min(line.count('#', 0, 4), 3)
            p = doc.add_heading(hm.group(1).strip(), level=level)
            compact(p, before=10, after=2)
            continue
        bm = bullet_re.match(line)
        if bm:
            p = doc.add_paragraph(bm.group(1).strip(), style='List Bullet')
            compact(p, before=0, after=2)
            if p.runs:
                p.runs[0].font.size = Pt(10)
            continue
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
        if clean:
            p = doc.add_paragraph(clean)
            compact(p, before=0, after=3)
            if p.runs:
                p.runs[0].font.size = Pt(10)

    # ── Activity links ────────────────────────────────────────────────────────
    if member_id:
        p = doc.add_heading('Parliamentary Activity Links', level=2)
        compact(p, before=10, after=2)
        lp = doc.add_paragraph()
        lp.add_run('Hansard contributions: ').bold = True
        _add_hyperlink(lp, f"https://hansard.parliament.uk/search/MemberContributions?memberId={member_id}", "View on Hansard")
        compact(lp, before=0, after=2)
        lp2 = doc.add_paragraph()
        lp2.add_run('Written questions: ').bold = True
        _add_hyperlink(lp2, f"https://members.parliament.uk/member/{member_id}/writtenquestions", "View on Parliament.uk")
        compact(lp2, before=0, after=2)

    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)
    safe_name = re.sub(r'[^\w\s-]', '', mp_name)[:40].strip()
    return send_file(
        mem, as_attachment=True,
        download_name=f"Profile - {safe_name}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
