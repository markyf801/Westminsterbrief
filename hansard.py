import re
import json
import html as _html_mod
import requests, io, csv
from flask import Blueprint, render_template, request, make_response
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from cache_models import CachedQuestion, CachedMember

hansard_bp = Blueprint('hansard', __name__)

DEPARTMENTS = {
    "All Departments": "", "Department for Education": "60", "Department of Health and Social Care": "17",
    "HM Treasury": "14", "Home Office": "1", "Ministry of Defence": "11", "Ministry of Justice": "54",
    "Department for Science, Innovation and Technology": "216", "Cabinet Office": "53"
}

PARTY_COLOURS = {
    'Labour': '#E4003B',
    'Conservative': '#0087DC',
    'Liberal Democrat': '#FAA61A',
    'SNP': '#005EB8',
    'Green Party': '#00B140',
    'Reform UK': '#12B6CF',
    'Plaid Cymru': '#3F8428',
    'DUP': '#D46A4C',
    'Alliance': '#F6CB2F',
}

WQ_API_URL = "https://questions-statements-api.parliament.uk/api/writtenquestions/questions"
WQ_PAGE_SIZE = 400
WQ_MAX_RESULTS = 1200

MEMBER_CACHE = {}


def strip_html(raw):
    """Strip HTML tags and decode entities from Parliament API text."""
    if not raw:
        return ''
    text = re.sub(r'<[^>]+>', ' ', raw)
    text = _html_mod.unescape(text)
    return ' '.join(text.split())


def _word_root(w):
    """Strip common English suffixes so 'repayments' matches 'repayment' in text."""
    for suffix in ('ments', 'ment', 'tions', 'tion', 'ings', 'ing', 'ers', 'es', 's'):
        if w.endswith(suffix) and len(w) - len(suffix) >= 4:
            return w[:-len(suffix)]
    return w


def fetch_wq_pages(params_base):
    """Fetch up to WQ_MAX_RESULTS written questions via parallel pagination."""
    all_items, total_available = [], 0
    r = requests.get(WQ_API_URL, params={**params_base, 'take': WQ_PAGE_SIZE, 'skip': 0}, timeout=60)
    if r.status_code != 200:
        return [], 0
    data = r.json()
    total_available = data.get('totalResults', 0)
    all_items.extend(data.get('results') or [])

    skips = list(range(WQ_PAGE_SIZE, min(total_available, WQ_MAX_RESULTS), WQ_PAGE_SIZE))
    if not skips:
        return all_items, total_available

    def _fetch(skip):
        resp = requests.get(WQ_API_URL, params={**params_base, 'take': WQ_PAGE_SIZE, 'skip': skip}, timeout=60)
        return resp.json().get('results') or [] if resp.status_code == 200 else []

    with ThreadPoolExecutor(max_workers=4) as ex:
        for f in as_completed([ex.submit(_fetch, s) for s in skips]):
            try:
                all_items.extend(f.result())
            except Exception:
                pass

    return all_items, total_available


def prefetch_members(member_ids):
    """Parallel-fetch member details for all IDs not already cached."""
    unknown = [mid for mid in member_ids if mid and mid not in MEMBER_CACHE and not CachedMember.get(mid)]
    if not unknown:
        return
    with ThreadPoolExecutor(max_workers=10) as ex:
        futures = {ex.submit(get_member_details, mid, 'Commons'): mid for mid in unknown}
        for f in as_completed(futures):
            try:
                f.result()
            except Exception:
                pass


def get_member_details(member_id, fallback_house):
    if not member_id:
        return "Unknown", "No Party", "Unknown", fallback_house
    if member_id in MEMBER_CACHE:
        return MEMBER_CACHE[member_id]

    cached = CachedMember.get(member_id)
    if cached:
        result = (cached.name, cached.party, cached.constituency, cached.house)
        MEMBER_CACHE[member_id] = result
        return result

    try:
        url = f"https://members-api.parliament.uk/api/Members/{member_id}"
        resp = requests.get(url, timeout=3)
        if resp.status_code == 200:
            data = resp.json().get('value') or {}
            membership = data.get('latestHouseMembership') or {}
            name = data.get('nameDisplayAs') or 'Unknown'
            party = (data.get('latestParty') or {}).get('name') or 'No Party'
            constituency = membership.get('membershipFrom') or 'Unknown'
            house = "Lords" if membership.get('house') == 2 else "Commons"
            image_url = data.get('thumbnailUrl') or ''
            CachedMember.store(member_id, name, party, constituency, house, image_url)
            MEMBER_CACHE[member_id] = (name, party, constituency, house)
            return MEMBER_CACHE[member_id]
    except Exception:
        pass
    return "Member", "Party Info Pending", "Unknown", fallback_house


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


@hansard_bp.route('/questions', methods=['GET', 'POST'])
def index():
    results, error_message = [], None
    selected_dept_id = ""
    selected_house = "All"
    status_filter = "all"
    subject, start_date, end_date = "", "", ""
    total_available = 0
    pre_filter_count = 0
    grouped_results = []

    if request.method == 'POST':
        subject = request.form.get('subject', '').strip()
        start_date = request.form.get('start_date', '').strip()
        end_date = request.form.get('end_date', '').strip()
        selected_dept_id = request.form.get('department', '').strip()
        selected_house = request.form.get('house_filter', 'All')
        status_filter = request.form.get('status_filter', 'all')
        action = request.form.get('action', 'search')

        subjects = [s.strip() for s in subject.split(',') if s.strip()] or ['']
        all_raw_results = []
        seen_uins = set()

        try:
            # Build shared API params — house pushed to API to cut wasted fetches.
            # Correct param names per OpenAPI spec (questions-statements-api.parliament.uk/index.html).
            # tabledWhenFrom/tabledWhenTo work reliably; answeringBodies is safe when a
            # searchTerm narrows the result set. See CLAUDE.md "WQ API constraints".
            base_params = {}
            if start_date: base_params['tabledWhenFrom'] = start_date
            if end_date: base_params['tabledWhenTo'] = end_date
            if selected_dept_id: base_params['answeringBodies'] = [int(selected_dept_id)]
            if selected_house == 'Commons': base_params['house'] = 'Commons'
            elif selected_house == 'Lords': base_params['house'] = 'Lords'

            def _fetch_subject(subj):
                p = dict(base_params)
                if subj:
                    p['searchTerm'] = subj
                return fetch_wq_pages(p)

            # Parallel multi-subject fetches
            fetch_errors = []
            with ThreadPoolExecutor(max_workers=min(len(subjects), 4)) as ex:
                futures = [ex.submit(_fetch_subject, s) for s in subjects]
                for f in as_completed(futures):
                    try:
                        batch, avail = f.result()
                        total_available = max(total_available, avail)
                        for item in batch:
                            uin = (item.get('value') or {}).get('uin')
                            if uin and uin not in seen_uins:
                                seen_uins.add(uin)
                                all_raw_results.append(item)
                    except Exception as e:
                        fetch_errors.append(str(e))

            if fetch_errors and not all_raw_results:
                raise Exception(f"Parliament API unavailable — please try again in a moment. ({fetch_errors[0]})")

            pre_filter_count = len(all_raw_results)

            # Pre-warm member cache in parallel
            all_member_ids = list({
                mid
                for item in all_raw_results
                for mid in [
                    item.get('value', {}).get('askingMemberId'),
                    item.get('value', {}).get('answeringMemberId'),
                ]
                if mid
            })
            prefetch_members(all_member_ids)

            # Build per-subject root sets for relevance filtering
            # Strip common suffixes so "repayments" matches "repayment" in text
            subject_word_sets = []
            for subj in subjects:
                roots = {_word_root(w.lower()) for w in subj.split() if len(w) >= 3}
                if roots:
                    subject_word_sets.append(roots)

            # Python-level filtering and row building
            for item in all_raw_results:
                val = item.get('value') or {}

                if selected_dept_id and str(val.get('answeringBodyId')) != selected_dept_id:
                    continue

                raw_date_full = val.get('dateTabled') or ''
                raw_date_str = raw_date_full.split('T')[0]

                if start_date and raw_date_str < start_date: continue
                if end_date and raw_date_str > end_date: continue

                # Relevance filter: at least 2 of N root words must appear (matching debate scanner logic).
                # Requiring ALL words is too strict for multi-word topics — "student loan repayments"
                # drops questions that say "loan repayment" without "student", etc.
                if subject_word_sets:
                    q_lower = (
                        strip_html(val.get('questionText', '')) + ' ' +
                        (val.get('heading') or '') + ' ' +
                        strip_html(val.get('answerText') or '')
                    ).lower()
                    matched = False
                    for word_set in subject_word_sets:
                        min_m = 2 if len(word_set) >= 2 else 1
                        if sum(1 for w in word_set if w in q_lower) >= min_m:
                            matched = True
                            break
                    if not matched:
                        continue

                val_member_id = val.get('askingMemberId')
                house_raw = val.get('house', 'Commons')
                name, party, constituency, actual_house = get_member_details(val_member_id, house_raw)

                if selected_house != "All" and actual_house != selected_house:
                    continue

                is_answered = bool((val.get('answerText') or '').strip() or val.get('dateAnswered'))
                is_holding = val.get('answerIsHolding', False)
                is_withdrawn = val.get('isWithdrawn', False)

                # Answer status filter
                if status_filter == 'unanswered' and (is_answered or is_holding or is_withdrawn):
                    continue
                elif status_filter == 'answered' and not (is_answered and not is_holding and not is_withdrawn):
                    continue
                elif status_filter == 'holding' and not is_holding:
                    continue
                elif status_filter == 'withdrawn' and not is_withdrawn:
                    continue

                try:
                    date_obj = datetime.fromisoformat(raw_date_str)
                    f_date = f"{date_obj.day} {date_obj.strftime('%B %Y')}"
                except Exception:
                    f_date = "N/A"

                uin = val.get('uin', '')
                question_url = f"https://questions-statements.parliament.uk/written-questions/detail/{raw_date_str}/{uin}"
                question_text = strip_html(val.get('questionText', ''))

                answering_member_id = val.get('answeringMemberId')
                ans_name = None
                if answering_member_id:
                    ans_name, _, _, _ = get_member_details(answering_member_id, 'Commons')

                answer_text = strip_html(val.get('answerText') or '')

                raw_answered = (val.get('dateAnswered') or '').split('T')[0]
                try:
                    da_obj = datetime.fromisoformat(raw_answered)
                    date_answered = f"{da_obj.day} {da_obj.strftime('%B %Y')}"
                except Exception:
                    date_answered = ''

                if uin and CachedQuestion.is_cacheable(raw_date_str):
                    try:
                        if not CachedQuestion.get(uin):
                            CachedQuestion.store(
                                uin=uin, member_name=name, party=party,
                                department_id=val.get('answeringBodyId', ''),
                                department_name=val.get('answeringBodyName', ''),
                                question_text=question_text,
                                answer_text=val.get('answerText'),
                                date_tabled=raw_date_str, url=question_url
                            )
                    except Exception:
                        pass

                results.append({
                    'uin': uin,
                    'url': question_url,
                    'dept': val.get('answeringBodyName'),
                    'name': name,
                    'party': party,
                    'role': "Life Peer" if actual_house == "Lords" else f"MP for {constituency}",
                    'date': f_date,
                    'text': question_text,
                    'heading': val.get('heading', ''),
                    'party_colour': PARTY_COLOURS.get(party, '#888888'),
                    'answered': is_answered,
                    'is_holding': is_holding,
                    'is_withdrawn': is_withdrawn,
                    'answer_text': answer_text,
                    'answering_minister': ans_name,
                    'date_answered': date_answered,
                    '_sort_date': raw_date_str,
                })

            results.sort(key=lambda r: r.get('_sort_date', ''), reverse=True)
            for r in results:
                r.pop('_sort_date', None)

            # Group by Parliament heading for topic view
            heading_groups = {}
            for r in results:
                key = r.get('heading') or 'Other'
                heading_groups.setdefault(key, []).append(r)
            grouped_results = sorted(heading_groups.items(), key=lambda x: len(x[1]), reverse=True)

            # Export metadata (shared by Word and CSV)
            dept_label = next((k for k, v in DEPARTMENTS.items() if v == selected_dept_id), 'All Departments')
            filename_ts = datetime.now().strftime('%Y%m%d_%H%M')
            meta_lines = [
                f"Search: {subject or '(all)'}",
                f"Department: {dept_label}",
                f"House: {selected_house}",
                f"Date range: {start_date or 'any'} to {end_date or 'any'}",
                f"Status filter: {status_filter}",
                f"Results: {len(results)} shown",
                f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}",
            ]

            if action == 'word' and results:
                doc = Document()
                doc.add_heading('Parliamentary Written Questions', 0)
                for line in meta_lines:
                    p = doc.add_paragraph(line)
                    p.runs[0].italic = True
                doc.add_paragraph()
                for r in results:
                    p = doc.add_paragraph()
                    status_label = ('WITHDRAWN' if r['is_withdrawn'] else
                                    'HOLDING ANSWER' if r['is_holding'] else
                                    'ANSWERED' if r['answered'] else 'UNANSWERED')
                    p.add_run(f"[{status_label}] [{r['dept']}] {r['name']} ({r['party']}, {r['role']})\n").bold = True
                    meta_run = f"UIN: {r['uin']}  ·  Date Asked: {r['date']}"
                    if r['answering_minister']:
                        meta_run += f"  ·  Answered by: {r['answering_minister']}"
                    if r['date_answered']:
                        meta_run += f" ({r['date_answered']})"
                    p.add_run(meta_run + "\n")
                    p.add_run(f"Question: {r['text']}\n")
                    if r['answer_text']:
                        p.add_run(f"Answer: {r['answer_text']}\n")
                    p.add_run("Link: ")
                    add_hyperlink(p, r['url'], r['url'])
                b = io.BytesIO()
                doc.save(b)
                b.seek(0)
                resp = make_response(b.getvalue())
                resp.headers["Content-Disposition"] = f"attachment; filename=WQ_{filename_ts}.docx"
                return resp

            elif action == 'csv' and results:
                si = io.StringIO()
                si.write('\ufeff')  # UTF-8 BOM for Excel on Windows
                cw = csv.writer(si)
                for line in meta_lines:
                    cw.writerow([line])
                cw.writerow([])
                cw.writerow(['UIN', 'Status', 'Department', 'Member', 'Party', 'Role',
                              'Date Asked', 'Answering Minister', 'Date Answered',
                              'Question', 'Answer', 'URL'])
                for r in results:
                    status_label = ('WITHDRAWN' if r['is_withdrawn'] else
                                    'HOLDING ANSWER' if r['is_holding'] else
                                    'ANSWERED' if r['answered'] else 'UNANSWERED')
                    cw.writerow([r['uin'], status_label, r['dept'], r['name'], r['party'],
                                 r['role'], r['date'], r['answering_minister'] or '',
                                 r['date_answered'], r['text'], r['answer_text'], r['url']])
                resp = make_response(si.getvalue())
                resp.headers["Content-Disposition"] = f"attachment; filename=WQ_{filename_ts}.csv"
                resp.headers["Content-type"] = "text/csv; charset=utf-8"
                return resp

        except Exception as e:
            error_message = f"Search error: {str(e)}"

    return render_template('index.html',
                           results=results,
                           grouped_results=grouped_results,
                           error_message=error_message,
                           departments=DEPARTMENTS,
                           selected_dept=selected_dept_id,
                           selected_house=selected_house,
                           status_filter=status_filter,
                           subject=subject,
                           start_date=start_date,
                           end_date=end_date,
                           total_available=total_available,
                           pre_filter_count=pre_filter_count,
                           results_cap=WQ_MAX_RESULTS)


@hansard_bp.route('/questions/download_selected', methods=['POST'])
def download_selected():
    try:
        items = json.loads(request.form.get('items_json', '[]'))
    except Exception:
        items = []
    if not items:
        return "No items selected", 400

    fmt = request.form.get('format', 'word')
    filename_ts = datetime.now().strftime('%Y%m%d_%H%M')

    if fmt == 'csv':
        si = io.StringIO()
        si.write('﻿')
        cw = csv.writer(si)
        cw.writerow(['UIN', 'Status', 'Department', 'Member', 'Party', 'Role',
                     'Date Asked', 'Answering Minister', 'Date Answered',
                     'Question', 'Answer', 'URL'])
        for r in items:
            status_label = ('WITHDRAWN' if r.get('is_withdrawn') else
                            'HOLDING ANSWER' if r.get('is_holding') else
                            'ANSWERED' if r.get('answered') else 'UNANSWERED')
            cw.writerow([r.get('uin', ''), status_label, r.get('dept', ''),
                         r.get('name', ''), r.get('party', ''), r.get('role', ''),
                         r.get('date', ''), r.get('answering_minister') or '',
                         r.get('date_answered', ''), r.get('text', ''),
                         r.get('answer_text', ''), r.get('url', '')])
        resp = make_response(si.getvalue())
        resp.headers["Content-Disposition"] = f"attachment; filename=WQ_selected_{filename_ts}.csv"
        resp.headers["Content-type"] = "text/csv; charset=utf-8"
        return resp

    doc = Document()
    doc.add_heading('Selected Parliamentary Written Questions', 0)
    p = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %b %Y %H:%M')}  ·  {len(items)} question{'s' if len(items) != 1 else ''}")
    p.runs[0].italic = True
    doc.add_paragraph()
    for r in items:
        p = doc.add_paragraph()
        status_label = ('WITHDRAWN' if r.get('is_withdrawn') else
                        'HOLDING ANSWER' if r.get('is_holding') else
                        'ANSWERED' if r.get('answered') else 'UNANSWERED')
        p.add_run(f"[{status_label}] [{r.get('dept', '')}] {r.get('name', '')} ({r.get('party', '')}, {r.get('role', '')})\n").bold = True
        meta = f"UIN: {r.get('uin', '')}  ·  Date Asked: {r.get('date', '')}"
        if r.get('answering_minister'):
            meta += f"  ·  Answered by: {r['answering_minister']}"
        if r.get('date_answered'):
            meta += f" ({r['date_answered']})"
        p.add_run(meta + "\n")
        p.add_run(f"Question: {r.get('text', '')}\n")
        if r.get('answer_text'):
            p.add_run(f"Answer: {r['answer_text']}\n")
        p.add_run("Link: ")
        add_hyperlink(p, r.get('url', ''), r.get('url', ''))
    b = io.BytesIO()
    doc.save(b)
    b.seek(0)
    resp = make_response(b.getvalue())
    resp.headers["Content-Disposition"] = f"attachment; filename=WQ_selected_{filename_ts}.docx"
    return resp
